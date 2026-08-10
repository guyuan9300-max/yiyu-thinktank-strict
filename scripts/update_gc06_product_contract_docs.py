from __future__ import annotations

from datetime import datetime
from pathlib import Path
import shutil

from docx import Document


DESKTOP = Path("/Users/guyuanyuan/Desktop")
REPO = Path("/Users/guyuanyuan/Documents/New project/projects/yiyu-thinktank-strict")
BACKUP_DIR = REPO / "output" / "doc" / "gc06-product-contract-20260808"
MARKER = "2026-08-08｜GC-06、任务计划 Agent 与成长陪伴补充裁决"


DOCUMENTS: dict[str, list[tuple[str, list[str]]]] = {
    "益语智库AI新版_88逻辑对象权威与15条黄金链前端功能账本_20260803.docx": [
        (
            MARKER,
            [
                "GC-06 可见组织计划只消费严格 planning_cycles 与 decision_actions；不得以旧业务快照、临时空数组或第二套组织模型替代。后台刷新期间继续显示最后确认投影，并明确显示更新状态。",
                "planning_cycles 可选绑定一个 client_id；一旦绑定，所属 decision_actions 默认继承同一项目。跨项目工作应拆成不同计划；未绑定项目的组织性计划不得虚构项目背景。",
                "decision_actions 即计划中的可执行步骤，不新增步骤表。步骤转任务必须保留 action_id、planning_cycle_id、client_id 与 CAS 版本；正式保存同时形成 tasks 记录和 decision_actions.task_id 关联，不能留下任务已创建但步骤未挂接的半完成状态。",
                "任务计划 Agent 在用户从步骤生成任务时，可读取该项目已确认知识并生成可编辑的任务标题、说明和背景草稿；用户确认前不得写 tasks。基础模式至少预填步骤原文、计划和项目。",
                "周复盘切周以 ISO 周一至周日为准，所有请求以 sandbox、membership、weekLabel、perspective、departmentId 和 requestSeq 隔离；迟到回包不得把界面切回本周。",
                "部门信号是 GC-06 的只读派生消费者。响应字段缺失、无数据或生成失败只能在局部面板显示准确状态，严禁触发任务与日程模块级或桌面根级崩溃。",
                "生成周复盘由任务计划 Agent 消费本周获授权的任务、会议、计划行动、正式项目知识和用户允许的个人偏好，先形成可编辑草稿与真实来源清单；用户提交后才形成正式 weekly_review_version。",
                "GC-13 成长陪伴 Agent 弱可见地消费正式任务、会议、复盘和成长证据，为经验墙、能力成长、徽章与概览生成 growth_read_models；取消独立‘个人成长权威’页和逐条确认门槛，保留就地纠正、排除与重新计算。",
            ],
        )
    ],
    "益语智库AI新版_腾讯云智能体记忆能力移植与产品裁决_20260805.docx": [
        (
            MARKER,
            [
                "任务计划 Agent 的计划/复盘能力必须调用共享项目知识与成员允许的记忆，不复制第二套 Memory/Wiki。计划绑定项目后，步骤生成任务时只召回该 client 的正式知识和明确可用的个人偏好。",
                "周复盘草稿必须区分正式业务事实、项目共享知识、个人偏好和模型推断；模型推断不得冒充事实。Agent 运行失败时保留确定性基础草稿并显示 failed_retryable，不阻断人工复盘。",
                "成长陪伴 Agent 不设置独立聊天或权威管理页。成长证据与读模型直接服务成长中心既有栏目；用户纠正或排除证据会使相关读模型失效并触发重算。",
                "Agent 身份只在运行进度、来源、更新时间和失败状态中弱显示；不得让用户为了获得正常成长展示而逐条审批 Agent 候选。",
            ],
        )
    ],
    "益语智库AI新版_权威事实与知识消费者合同_20260806.docx": [
        (
            MARKER,
            [
                "CC-PLAN-01 补充：计划可选绑定单一 client，decision_actions 作为步骤继承该归属。步骤生成任务的正式命令必须原子保证 tasks 与 decision_actions.task_id 一致；失败时不得只留下其中一侧。",
                "CC-PLAN-03 补充：任务计划 Agent 生成的是可编辑周复盘草稿，不是正式事实。直接消费者为周复盘界面；提交后 weekly_review_versions、成长陪伴和组织计划才消费正式版本。",
                "CC-AGENT-01B 补充：任务计划 Agent 负责项目背景辅助、步骤任务草稿和周复盘草稿；不得决定负责人、优先级、截止日期，不得绕过用户提交修改任务或正式复盘。",
                "CC-AGENT-01F 补充：成长陪伴 Agent 自动加工获授权的正式工作证据并更新成长中心全部既有栏目。用户可就地纠正、排除并重算，但无需逐条确认候选；Agent 不单独占用前端标签页。",
                "消费者状态规则补充：组织计划、周复盘和部门信号刷新时继续显示最后确认版本及‘更新中’状态；迟到回包、字段缺失和局部失败不得清空正确界面或升级为桌面全屏错误。",
            ],
        )
    ],
}


def append_section(path: Path, sections: list[tuple[str, list[str]]]) -> bool:
    doc = Document(path)
    all_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    if MARKER in all_text:
        return False
    doc.add_page_break()
    for heading, bullets in sections:
        doc.add_heading(heading, level=1)
        for item in bullets:
            doc.add_paragraph(item, style="List Bullet")
    doc.save(path)
    return True


def main() -> None:
    BACKUP_DIR.mkdir(parents=True, exist_ok=True)
    changed: list[str] = []
    for filename, sections in DOCUMENTS.items():
        path = DESKTOP / filename
        if not path.exists():
            raise FileNotFoundError(path)
        backup = BACKUP_DIR / filename
        if not backup.exists():
            shutil.copy2(path, backup)
        if append_section(path, sections):
            changed.append(filename)
    print(f"updated={len(changed)} checked={len(DOCUMENTS)} at={datetime.now().isoformat(timespec='seconds')}")


if __name__ == "__main__":
    main()
