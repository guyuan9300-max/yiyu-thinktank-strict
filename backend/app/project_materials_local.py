"""Device-local project material storage and smart-import workflow.

Source files, raw smart-import text, and draft parsing state stay inside the
member's local managed files with strict ``object_manifests`` receipts.  This
module deliberately has no dependency on the cloud service package.
"""

from __future__ import annotations

import hashlib
import json
import math
import mimetypes
import re
import zipfile
from collections import Counter
from datetime import datetime, timedelta, timezone
from io import BytesIO
from pathlib import Path
from typing import Any, Iterable, Mapping
from xml.etree import ElementTree
from xml.sax.saxutils import escape

from strict_common.ids import canonical_json, new_id, utc_now

from .runtime import LocalRuntimeError, WorkspaceRuntime


def _bounded_project_lease(value: Any) -> str:
    """Defensively reject a cloud projection lease beyond the 24h contract."""
    now = datetime.now(timezone.utc)
    maximum = now + timedelta(hours=24)
    raw = str(value or "").strip()
    try:
        candidate = datetime.fromisoformat(raw.replace("Z", "+00:00"))
        if candidate.tzinfo is None:
            candidate = candidate.replace(tzinfo=timezone.utc)
        candidate = candidate.astimezone(timezone.utc)
    except ValueError:
        candidate = maximum
    return min(candidate, maximum).isoformat(timespec="milliseconds").replace(
        "+00:00", "Z"
    )


class LocalProjectMaterialsRepository:
    DOCX_MEDIA_TYPE = (
        "application/vnd.openxmlformats-officedocument."
        "wordprocessingml.document"
    )
    SUMMARY_MEDIA_TYPE = "application/vnd.yiyu.project-knowledge-summary+json"
    SMART_SESSION_MEDIA_TYPE = "application/vnd.yiyu.smart-import-session+json"
    PROJECT_STATE_MEDIA_TYPE = "application/vnd.yiyu.project-local-state+json"
    IMPORT_OPERATION_MEDIA_TYPE = (
        "application/vnd.yiyu.project-material-import-operation+json"
    )
    WIKI_DOCUMENT_MEDIA_TYPE = "text/vnd.yiyu.local-wiki-document"
    WIKI_CHUNK_MEDIA_TYPE = "text/vnd.yiyu.local-wiki-chunk"
    WIKI_SEARCH_INDEX_MEDIA_TYPE = (
        "application/vnd.yiyu.local-wiki-hybrid-index+json"
    )
    WIKI_GENERATOR_VERSION = (
        "tencentdb-agent-memory-wiki-b44c6db5+yiyu-gc10-p04"
    )
    WIKI_CHUNK_TARGET_CHARS = 12_000
    WIKI_CHUNK_OVERLAP_CHARS = 400
    WIKI_RETRIEVAL_GENERATOR_VERSION = (
        "tencentdb-agent-memory-rrf+yiyu-gc10-p05b"
    )
    WIKI_SPARSE_VECTOR_MODEL = "yiyu-local-hashed-subword-v2"
    WIKI_SPARSE_VECTOR_DIMENSIONS = 65_536
    _DOCX_PLACEHOLDER_PATTERN = re.compile(
        r"\{\{\s*([^{}]{1,120}?)\s*\}\}"
    )
    _DOCX_EMPTY_MARKERS = {
        "",
        "____",
        "待填写",
        "待补充",
        "待完善",
        "tbd",
        "todo",
    }

    def __init__(self, runtime: WorkspaceRuntime, context_provider: Any = None):
        self.runtime = runtime
        self.data_root = Path(runtime.database_path).resolve().parent
        self.context_provider = context_provider

    @staticmethod
    def _stable_segment(value: str) -> str:
        return hashlib.sha256(value.encode("utf-8")).hexdigest()[:32]

    @staticmethod
    def _safe_name(value: str) -> str:
        normalized = re.sub(r"[^0-9A-Za-z\u4e00-\u9fff._-]+", "_", value)
        return normalized.strip("._")[:160] or "material"

    @staticmethod
    def select_relevant_excerpt(
        content: str,
        query: str,
        *,
        max_chars: int = 40_000,
    ) -> str:
        """Keep prompt-relevant windows instead of blindly truncating the file head."""
        normalized = str(content or "").strip()
        if len(normalized) <= max_chars:
            return normalized

        raw_terms = re.findall(
            r"[A-Za-z0-9_]{2,}|[\u4e00-\u9fff]{2,24}",
            str(query or ""),
        )
        terms: list[str] = []
        for raw_term in raw_terms:
            terms.append(raw_term.casefold())
            if re.fullmatch(r"[\u4e00-\u9fff]{4,}", raw_term):
                terms.extend(
                    raw_term[index : index + 3]
                    for index in range(max(0, len(raw_term) - 2))
                )
        lowered_content = normalized.casefold()
        positions = sorted(
            {
                position
                for term in terms
                if term
                for position in [lowered_content.find(term)]
                if position >= 0
            }
        )
        if not positions:
            half = max_chars // 2
            return (
                normalized[:half].rstrip()
                + "\n\n[…中间内容未命中问题关键词，已省略…]\n\n"
                + normalized[-half:].lstrip()
            )[:max_chars]

        radius = 4_000
        spans: list[tuple[int, int]] = []
        for position in positions:
            start = max(0, position - radius)
            end = min(len(normalized), position + radius)
            if spans and start <= spans[-1][1]:
                spans[-1] = (spans[-1][0], max(spans[-1][1], end))
            else:
                spans.append((start, end))

        excerpts: list[str] = []
        used = 0
        for start, end in spans:
            remaining = max_chars - used
            if remaining <= 0:
                break
            excerpt = normalized[start:end].strip()
            if not excerpt:
                continue
            separator = "\n\n[…]\n\n" if excerpts else ""
            excerpt = excerpt[: max(0, remaining - len(separator))]
            excerpts.append(separator + excerpt)
            used += len(separator) + len(excerpt)
        return "".join(excerpts)[:max_chars]

    @classmethod
    def _chunk_wiki_text(cls, content: str) -> list[dict[str, Any]]:
        """Create deterministic local chunks with exact character locators.

        The target and overlap follow the imported TencentDB Agent Memory Wiki
        chunker.  The Python adapter keeps exact offsets so later citations can
        reopen the member-local original without uploading chunk text.
        """
        normalized = content.replace("\r\n", "\n").replace("\r", "\n").strip()
        if not normalized:
            return []
        target = cls.WIKI_CHUNK_TARGET_CHARS
        overlap = cls.WIKI_CHUNK_OVERLAP_CHARS
        chunks: list[dict[str, Any]] = []
        start = 0
        ordinal = 0
        while start < len(normalized):
            hard_end = min(len(normalized), start + target)
            end = hard_end
            if hard_end < len(normalized):
                boundary = normalized.rfind(
                    "\n\n",
                    start + max(1_000, target // 2),
                    hard_end,
                )
                if boundary > start:
                    end = boundary
            raw = normalized[start:end]
            leading = len(raw) - len(raw.lstrip())
            trailing = len(raw) - len(raw.rstrip())
            content_start = start + leading
            content_end = end - trailing
            chunk = normalized[content_start:content_end]
            if chunk:
                chunks.append(
                    {
                        "ordinal": ordinal,
                        "content": chunk,
                        "start": content_start,
                        "end": content_end,
                    }
                )
                ordinal += 1
            if end >= len(normalized):
                break
            next_start = max(start + 1, content_end - overlap)
            start = next_start
        return chunks

    @staticmethod
    def _retrieval_tokens(value: str) -> list[str]:
        """Tokenize Chinese/Latin text for local keyword and sparse-vector recall.

        The result is deliberately rebuildable and stays outside SQLite.  Chinese
        bi/tri-grams make non-identical short phrasings comparable without
        pretending that a remote embedding model was called.
        """
        normalized = re.sub(r"\s+", " ", str(value or "").casefold()).strip()
        for stop_phrase in (
            "为什么",
            "是什么",
            "怎么样",
            "怎么办",
            "如何",
            "哪些",
            "什么",
            "是否",
            "请问",
            "帮我找",
            "想了解",
            "关于",
        ):
            normalized = normalized.replace(stop_phrase, " ")
        tokens: list[str] = []
        for match in re.finditer(r"[a-z0-9_]{2,}|[\u4e00-\u9fff]+", normalized):
            segment = match.group(0)
            if re.fullmatch(r"[\u4e00-\u9fff]+", segment):
                if len(segment) <= 12:
                    tokens.append(segment)
                for width in (2, 3):
                    tokens.extend(
                        segment[index : index + width]
                        for index in range(max(0, len(segment) - width + 1))
                    )
            else:
                tokens.append(segment)
        return [token for token in tokens if token]

    @staticmethod
    def _fact_candidates(
        content: str,
        *,
        document_offset: int,
    ) -> list[dict[str, Any]]:
        """Create local candidate facts with exact chunk/document locators."""
        candidates: list[dict[str, Any]] = []
        for match in re.finditer(r"[^。！？!?；;\n]+[。！？!?；;]?", content):
            raw = match.group(0)
            leading = len(raw) - len(raw.lstrip())
            trailing = len(raw) - len(raw.rstrip())
            start = match.start() + leading
            end = match.end() - trailing
            value = content[start:end].strip()
            if len(value) < 8 or not re.search(r"[A-Za-z0-9\u4e00-\u9fff]", value):
                continue
            for block_start in range(0, len(value), 800):
                block = value[block_start : block_start + 800].strip()
                if len(block) < 8:
                    continue
                local_start = start + block_start
                local_end = local_start + len(block)
                candidates.append(
                    {
                        "content": block,
                        "chunkStart": local_start,
                        "chunkEnd": local_end,
                        "documentStart": document_offset + local_start,
                        "documentEnd": document_offset + local_end,
                    }
                )
        return candidates

    @classmethod
    def _sparse_vector(cls, tokens: Iterable[str]) -> dict[int, float]:
        counts = Counter(str(token) for token in tokens if str(token))
        vector: dict[int, float] = {}
        for token, count in counts.items():
            bucket = int(hashlib.sha256(token.encode("utf-8")).hexdigest()[:8], 16)
            bucket %= cls.WIKI_SPARSE_VECTOR_DIMENSIONS
            vector[bucket] = vector.get(bucket, 0.0) + 1.0 + math.log(count)
        norm = math.sqrt(sum(value * value for value in vector.values()))
        return (
            {key: value / norm for key, value in vector.items()}
            if norm > 0
            else {}
        )

    @staticmethod
    def _sparse_cosine(
        left: Mapping[int, float],
        right: Mapping[int, float],
    ) -> float:
        if not left or not right:
            return 0.0
        if len(left) > len(right):
            left, right = right, left
        return sum(value * right.get(key, 0.0) for key, value in left.items())

    def _read_local_manifest_text(self, manifest_id: str) -> str:
        context = self._context()
        with self.runtime._connection() as connection:
            scope_id = self.runtime._local_object_scope_id(
                connection,
                context.sandbox_id,
            )
            row = connection.execute(
                """
                SELECT storage_key, content_hash, byte_size, availability_state
                FROM object_manifests
                WHERE id=? AND scope_id=? AND holder_role='sandbox'
                  AND holder_instance_id=? AND lifecycle_state='active'
                """,
                (manifest_id, scope_id, context.sandbox_id),
            ).fetchone()
        if row is None or str(row["availability_state"] or "") != "ready":
            raise LocalRuntimeError(
                409,
                "local_wiki_object_unavailable",
                "本机知识对象当前不可读取",
            )
        path = self._managed_path(str(row["storage_key"] or ""))
        try:
            data = path.read_bytes()
        except OSError as exc:
            raise LocalRuntimeError(
                409,
                "local_wiki_object_missing",
                "本机知识对象文件缺失",
            ) from exc
        if (
            len(data) != int(row["byte_size"] or 0)
            or hashlib.sha256(data).hexdigest() != str(row["content_hash"] or "")
        ):
            raise LocalRuntimeError(
                409,
                "local_wiki_object_corrupt",
                "本机知识对象校验失败",
            )
        try:
            return data.decode("utf-8")
        except UnicodeDecodeError as exc:
            raise LocalRuntimeError(
                409,
                "local_wiki_object_encoding_invalid",
                "本机知识对象编码无效",
            ) from exc

    def _ensure_wiki_object(
        self,
        *,
        object_id: str,
        storage_key: str,
        media_type: str,
        content: str,
    ) -> dict[str, Any]:
        context = self._context()
        data = content.encode("utf-8")
        content_hash = hashlib.sha256(data).hexdigest()
        current = self.runtime.local_storage_object_get(
            sandbox_id=context.sandbox_id,
            object_id=object_id,
        )
        if (
            current is not None
            and str(current.get("lifecycle_state") or "") == "active"
            and str(current.get("content_hash") or "") == content_hash
            and str(current.get("storage_key") or "") == storage_key
        ):
            return {
                "objectId": object_id,
                "manifestId": str(current["manifest_id"]),
                "storageKey": storage_key,
                "contentHash": content_hash,
                "byteSize": len(data),
                "version": int(current["version"]),
            }
        stored = self._upsert_object(
            sandbox_id=context.sandbox_id,
            object_id=object_id,
            storage_key=storage_key,
            media_type=media_type,
            data=data,
            expected_version=int((current or {}).get("version") or 0),
        )
        manifest = self.runtime.local_storage_object_get(
            sandbox_id=context.sandbox_id,
            object_id=object_id,
        )
        if manifest is None:
            raise LocalRuntimeError(
                500,
                "local_wiki_manifest_missing",
                "本机知识对象写入后缺少清单回执",
            )
        return {**stored, "manifestId": str(manifest["manifest_id"])}

    def _context(self) -> Any:
        if callable(self.context_provider):
            return self.context_provider()
        return self.runtime._current_context(require_ready=True)

    def _managed_path(self, storage_key: str) -> Path:
        candidate = (self.data_root / storage_key).resolve()
        if self.data_root not in candidate.parents:
            raise LocalRuntimeError(
                422,
                "local_storage_path_invalid",
                "本机受管路径越界",
            )
        return candidate

    def _upsert_object(
        self,
        *,
        sandbox_id: str,
        object_id: str,
        storage_key: str,
        media_type: str,
        data: bytes,
        expected_version: int | None = None,
    ) -> dict[str, Any]:
        path = self._managed_path(storage_key)
        content_hash = hashlib.sha256(data).hexdigest()
        with self.runtime.local_storage_object_lock(
            sandbox_id=sandbox_id,
            object_id=object_id,
        ):
            current = self.runtime.local_storage_object_get(
                sandbox_id=sandbox_id,
                object_id=object_id,
            )
            current_version = int(current["version"]) if current is not None else 0
            if (
                expected_version is not None
                and current_version != expected_version
            ):
                raise LocalRuntimeError(
                    409,
                    "local_storage_version_conflict",
                    "本机对象版本已变化，请刷新后重试",
                )
            path.parent.mkdir(parents=True, exist_ok=True)
            temporary = path.with_name(f".{path.name}.{new_id()}.tmp")
            temporary.write_bytes(data)
            temporary.replace(path)
            stored = self.runtime.local_storage_object_put(
                sandbox_id=sandbox_id,
                object_id=object_id,
                storage_key=storage_key,
                content_hash=content_hash,
                media_type=media_type,
                byte_size=len(data),
                expected_version=current_version,
            )
        return {
            "objectId": object_id,
            "storageKey": storage_key,
            "contentHash": content_hash,
            "mediaType": media_type,
            "byteSize": len(data),
            "version": stored["version"],
            "updatedAt": stored["updatedAt"],
            "path": str(path),
        }

    def _project_state_key(self, sandbox_id: str, project_id: str) -> str:
        return (
            "local-project-state/"
            f"{self._stable_segment(sandbox_id)}/"
            f"{self._stable_segment(project_id)}.json"
        )

    def _import_operation_identity(
        self,
        sandbox_id: str,
        idempotency_key: str,
    ) -> tuple[str, str]:
        operation_segment = self._stable_segment(idempotency_key)
        return (
            f"project-material-import:{operation_segment}",
            "local-project-import-operations/"
            f"{self._stable_segment(sandbox_id)}/{operation_segment}.json",
        )

    def _load_import_operation(
        self,
        *,
        sandbox_id: str,
        object_id: str,
        storage_key: str,
        request_fingerprint: str,
    ) -> dict[str, Any] | None:
        row = self.runtime.local_storage_object_get(
            sandbox_id=sandbox_id,
            object_id=object_id,
            storage_key=storage_key,
        )
        if row is None:
            return None
        if (
            str(row.get("media_type") or "") != self.IMPORT_OPERATION_MEDIA_TYPE
            or str(row.get("lifecycle_state") or "") != "active"
        ):
            raise LocalRuntimeError(
                409,
                "local_import_receipt_invalid",
                "本机导入操作回执无效",
            )
        path = self._managed_path(storage_key)
        try:
            data = path.read_bytes()
        except OSError as exc:
            raise LocalRuntimeError(
                409,
                "local_import_receipt_missing",
                "本机导入操作回执缺失",
            ) from exc
        if (
            len(data) != int(row["byte_size"])
            or hashlib.sha256(data).hexdigest() != str(row["content_hash"])
        ):
            raise LocalRuntimeError(
                409,
                "local_import_receipt_corrupt",
                "本机导入操作回执校验失败",
            )
        try:
            receipt = json.loads(data.decode("utf-8"))
        except (UnicodeDecodeError, json.JSONDecodeError) as exc:
            raise LocalRuntimeError(
                409,
                "local_import_receipt_corrupt",
                "本机导入操作回执无法读取",
            ) from exc
        if (
            not isinstance(receipt, Mapping)
            or str(receipt.get("requestFingerprint") or "")
            != request_fingerprint
        ):
            raise LocalRuntimeError(
                409,
                "local_import_idempotency_conflict",
                "导入操作标识已用于不同请求",
            )
        result = receipt.get("result")
        if not isinstance(result, Mapping):
            raise LocalRuntimeError(
                409,
                "local_import_receipt_corrupt",
                "本机导入操作回执没有结果",
            )
        return dict(result)

    def _write_import_operation(
        self,
        *,
        sandbox_id: str,
        object_id: str,
        storage_key: str,
        project_id: str,
        operation_kind: str,
        request_fingerprint: str,
        result: Mapping[str, Any],
    ) -> None:
        payload = {
            "schema": "yiyu.project-material-import-operation.v1",
            "projectId": project_id,
            "operationKind": operation_kind,
            "requestFingerprint": request_fingerprint,
            "state": "local_prepared",
            "result": dict(result),
            "updatedAt": utc_now(),
        }
        self._upsert_object(
            sandbox_id=sandbox_id,
            object_id=object_id,
            storage_key=storage_key,
            media_type=self.IMPORT_OPERATION_MEDIA_TYPE,
            data=canonical_json(payload).encode("utf-8"),
            expected_version=0,
        )

    def _empty_project_state(
        self,
        project_id: str,
        *,
        sandbox_id: str,
    ) -> dict[str, Any]:
        return {
            "schema": "yiyu.project-local-state.v1",
            "projectId": project_id,
            "folders": [],
            "documents": {},
            "pendingCloudDeletes": {},
            "duplicateResolutions": {},
            "linkImportRuns": {},
            "templateFillRuns": {},
            "meetings": {},
            "reportDrafts": {},
            "updatedAt": utc_now(),
            "_localStorageVersion": 0,
            "_localSandboxId": sandbox_id,
        }

    def _load_project_state(self, project_id: str) -> dict[str, Any]:
        context = self._context()
        object_id = f"project-state:{self._stable_segment(project_id)}"
        storage_key = self._project_state_key(context.sandbox_id, project_id)
        row = self.runtime.local_storage_object_get(
            sandbox_id=context.sandbox_id,
            object_id=object_id,
            storage_key=storage_key,
        )
        if row is None:
            return self._empty_project_state(
                project_id,
                sandbox_id=context.sandbox_id,
            )
        if str(row["sandbox_id"]) != context.sandbox_id:
            raise LocalRuntimeError(
                409,
                "local_storage_sandbox_changed",
                "本机工作空间已切换，请重试",
            )
        path = self._managed_path(storage_key)
        try:
            data = path.read_bytes()
        except OSError as exc:
            raise LocalRuntimeError(
                409,
                "local_project_state_missing",
                "本机项目资料状态文件缺失",
            ) from exc
        if (
            len(data) != int(row["byte_size"])
            or hashlib.sha256(data).hexdigest() != str(row["content_hash"])
        ):
            raise LocalRuntimeError(
                409,
                "local_project_state_corrupt",
                "本机项目资料状态校验失败",
            )
        try:
            state = json.loads(data.decode("utf-8"))
        except (UnicodeDecodeError, json.JSONDecodeError) as exc:
            raise LocalRuntimeError(
                409,
                "local_project_state_corrupt",
                "本机项目资料状态无法读取",
            ) from exc
        if (
            not isinstance(state, dict)
            or str(state.get("projectId") or "") != project_id
        ):
            raise LocalRuntimeError(
                409,
                "local_project_state_mismatch",
                "本机项目资料状态不属于当前项目",
            )
        state["_localStorageVersion"] = int(row["version"])
        state["_localSandboxId"] = context.sandbox_id
        return state

    def _write_project_state(
        self,
        project_id: str,
        state: Mapping[str, Any],
    ) -> dict[str, Any]:
        sandbox_id = str(state.get("_localSandboxId") or "")
        if not sandbox_id:
            sandbox_id = self._context().sandbox_id
        expected_version = int(state.get("_localStorageVersion") or 0)
        payload = {
            **{
                key: value
                for key, value in dict(state).items()
                if not str(key).startswith("_")
            },
            "schema": "yiyu.project-local-state.v1",
            "projectId": project_id,
            "updatedAt": utc_now(),
        }
        self._upsert_object(
            sandbox_id=sandbox_id,
            object_id=f"project-state:{self._stable_segment(project_id)}",
            storage_key=self._project_state_key(
                sandbox_id,
                project_id,
            ),
            media_type=self.PROJECT_STATE_MEDIA_TYPE,
            data=canonical_json(payload).encode("utf-8"),
            expected_version=expected_version,
        )
        payload["_localStorageVersion"] = expected_version + 1
        payload["_localSandboxId"] = sandbox_id
        return payload

    def bind_cloud_documents(
        self,
        *,
        project_id: str,
        local_materials: Iterable[Mapping[str, Any]],
        cloud_documents: Iterable[Mapping[str, Any]],
    ) -> None:
        by_source = {
            str(item.get("localSourceId") or ""): dict(item)
            for item in local_materials
            if str(item.get("localSourceId") or "")
        }
        state = self._load_project_state(project_id)
        documents = dict(state.get("documents") or {})
        for cloud in cloud_documents:
            source_id = str(cloud.get("localSourceId") or "")
            document_id = str(cloud.get("documentId") or "")
            local = by_source.get(source_id)
            if not local or not document_id:
                continue
            for existing_id, existing in list(documents.items()):
                if (
                    existing_id != document_id
                    and isinstance(existing, Mapping)
                    and str(existing.get("localSourceId") or "") == source_id
                ):
                    documents.pop(existing_id, None)
            documents[document_id] = {
                "documentId": document_id,
                "cloudDocumentId": (
                    document_id
                    if not document_id.startswith("local-pending:")
                    else None
                ),
                "localSourceId": source_id,
                "localSummaryId": local.get("localSummaryId"),
                "fileName": local.get("fileName"),
                "title": local.get("title") or local.get("fileName"),
                "mediaType": local.get("mediaType"),
                "contentHash": local.get("contentHash"),
                "byteSize": int(local.get("byteSize") or 0),
                "managedPath": local.get("managedPath"),
                "originalSourcePath": local.get("originalSourcePath"),
                "folderId": None,
                "cloudMetadataState": (
                    "pending"
                    if document_id.startswith("local-pending:")
                    else "ready"
                ),
                "version": int(cloud.get("version") or 1),
                "updatedAt": local.get("updatedAt") or utc_now(),
            }
        state["documents"] = documents
        self._write_project_state(project_id, state)

    def bind_pending_materials(
        self,
        *,
        project_id: str,
        local_materials: Iterable[Mapping[str, Any]],
    ) -> list[dict[str, Any]]:
        materials = [dict(item) for item in local_materials]
        self.bind_cloud_documents(
            project_id=project_id,
            local_materials=materials,
            cloud_documents=[
                {
                    "localSourceId": item.get("localSourceId"),
                    "documentId": (
                        "local-pending:"
                        + str(item.get("localSourceId") or "")
                    ),
                }
                for item in materials
            ],
        )
        return self.pending_cloud_materials(project_id)

    def bind_meeting_materials(
        self,
        *,
        project_id: str,
        meeting_id: str,
        local_materials: Iterable[Mapping[str, Any]],
    ) -> dict[str, Any]:
        """Relate imported local project files to one meeting via a source set.

        The files remain ordinary project materials and therefore appear in the
        workbench.  The source set adds a many-to-many meeting consumption edge
        without changing the source asset's identity or inventing an attachment
        authority table.
        """
        context = self._context()
        materials = [dict(item) for item in local_materials]
        state_documents = dict(self._load_project_state(project_id).get("documents") or {})
        entries: list[tuple[str, int]] = []
        for material in materials:
            local_source_id = str(material.get("localSourceId") or "").strip()
            entry = next(
                (
                    dict(raw)
                    for raw in state_documents.values()
                    if isinstance(raw, Mapping)
                    and str(raw.get("localSourceId") or "") == local_source_id
                ),
                None,
            )
            if entry is None:
                continue
            source_asset_id = self._ensure_local_source_asset(
                project_id=project_id,
                entry=entry,
            )
            entries.append((source_asset_id, int(entry.get("version") or 1)))
        now = utc_now()
        source_set_id = "meeting_materials_" + self._stable_segment(
            f"{context.sandbox_id}|{meeting_id}"
        )
        with self.runtime._connection() as connection:
            scope_id = self.runtime._local_object_scope_id(
                connection,
                context.sandbox_id,
            )
            meeting = connection.execute(
                "SELECT id FROM meetings WHERE id=? AND scope_id=? AND client_id=? "
                "AND lifecycle_state='active'",
                (meeting_id, scope_id, project_id),
            ).fetchone()
            if meeting is None:
                raise LocalRuntimeError(
                    409,
                    "meeting_material_project_mismatch",
                    "会议与当前项目不一致，附件未建立会议关联",
                )
            connection.execute("BEGIN IMMEDIATE")
            connection.execute(
                "INSERT INTO source_sets (id,scope_id,client_id,"
                "security_label_set_version,source_count,version,purpose_kind,"
                "publication_state,created_by_principal_id,created_at,expires_at,"
                "lifecycle_state,updated_at,deleted_at,authority_role,"
                "origin_instance_id) VALUES (?,?,?,'member-local-v1',?,1,"
                "'meeting_support_materials','draft',?,?,NULL,'active',?,NULL,"
                "'local',?) ON CONFLICT(id) DO UPDATE SET source_count=?,"
                "version=source_sets.version+1,lifecycle_state='active',"
                "updated_at=excluded.updated_at,deleted_at=NULL",
                (
                    source_set_id,
                    scope_id,
                    project_id,
                    len(entries),
                    context.principal_id,
                    now,
                    now,
                    context.sandbox_id,
                    len(entries),
                ),
            )
            for ordinal, (source_asset_id, source_version) in enumerate(entries):
                member_id = "meeting_material_member_" + self._stable_segment(
                    f"{source_set_id}|{source_asset_id}"
                )
                connection.execute(
                    "INSERT INTO source_set_members (id,scope_id,source_set_id,"
                    "source_object_id,source_version,policy_version,"
                    "source_object_kind,ordinal,added_at,removed_at,version,"
                    "lifecycle_state,created_at,updated_at,deleted_at,"
                    "authority_role,origin_instance_id) VALUES (?,?,?,?,?,1,"
                    "'source_asset',?,?,NULL,1,'active',?,?,NULL,'local',?) "
                    "ON CONFLICT(id) DO UPDATE SET source_version=excluded.source_version,"
                    "ordinal=excluded.ordinal,removed_at=NULL,lifecycle_state='active',"
                    "updated_at=excluded.updated_at,deleted_at=NULL",
                    (
                        member_id,
                        scope_id,
                        source_set_id,
                        source_asset_id,
                        source_version,
                        ordinal,
                        now,
                        now,
                        now,
                        context.sandbox_id,
                    ),
                )
            connection.execute(
                "UPDATE source_sets SET source_count=(SELECT COUNT(*) FROM "
                "source_set_members WHERE scope_id=? AND source_set_id=? "
                "AND lifecycle_state='active') WHERE scope_id=? AND id=?",
                (scope_id, source_set_id, scope_id, source_set_id),
            )
            connection.commit()
        return {
            "meetingId": meeting_id,
            "sourceSetId": source_set_id,
            "linkedCount": len(entries),
        }

    def meeting_materials(
        self,
        *,
        project_id: str,
        meeting_id: str,
    ) -> list[dict[str, Any]]:
        """List member-local project files explicitly attached to a meeting."""
        context = self._context()
        source_set_id = "meeting_materials_" + self._stable_segment(
            f"{context.sandbox_id}|{meeting_id}"
        )
        with self.runtime._connection() as connection:
            scope_id = self.runtime._local_object_scope_id(
                connection,
                context.sandbox_id,
            )
            meeting = connection.execute(
                "SELECT id FROM meetings WHERE id=? AND scope_id=? AND client_id=? "
                "AND lifecycle_state='active'",
                (meeting_id, scope_id, project_id),
            ).fetchone()
            if meeting is None:
                raise LocalRuntimeError(
                    409,
                    "meeting_material_project_mismatch",
                    "会议与当前项目不一致，无法读取会议附件",
                )
            rows = connection.execute(
                "SELECT asset.id,asset.display_name,asset.media_type,asset.byte_size,"
                "manifest.local_original_path,member.ordinal "
                "FROM source_set_members AS member "
                "JOIN source_assets AS asset ON asset.scope_id=member.scope_id "
                "AND asset.id=member.source_object_id "
                "LEFT JOIN object_manifests AS manifest "
                "ON manifest.scope_id=asset.scope_id "
                "AND manifest.id=asset.object_manifest_id "
                "WHERE member.scope_id=? AND member.source_set_id=? "
                "AND member.source_object_kind='source_asset' "
                "AND member.lifecycle_state='active' "
                "AND asset.client_id=? AND asset.lifecycle_state='active' "
                "ORDER BY member.ordinal,asset.display_name,asset.id",
                (scope_id, source_set_id, project_id),
            ).fetchall()
        materials: list[dict[str, Any]] = []
        for row in rows:
            local_path = str(row["local_original_path"] or "").strip()
            available = bool(local_path and Path(local_path).is_file())
            materials.append(
                {
                    "id": str(row["id"]),
                    "fileName": str(row["display_name"] or "未命名附件"),
                    "mediaType": str(row["media_type"] or "application/octet-stream"),
                    "byteSize": int(row["byte_size"] or 0),
                    "localPath": local_path if available else None,
                    "availabilityState": "ready" if available else "missing",
                }
            )
        return materials

    def bind_task_attachment(
        self,
        *,
        project_id: str,
        document_id: str,
        task_id: str,
    ) -> dict[str, Any]:
        """Bind a member-local source to one strict task without a new table.

        The stable relation lives in ``source_assets.source_locator_nonlocal``;
        the actual path remains only in the local object manifest.
        """
        state = self._load_project_state(project_id)
        documents = dict(state.get("documents") or {})
        entry = documents.get(document_id)
        if not isinstance(entry, Mapping):
            raise LocalRuntimeError(404, "task_attachment_missing", "任务附件不存在")
        normalized = {**dict(entry), "taskId": task_id, "sourceKind": "task_attachment"}
        documents[document_id] = normalized
        state["documents"] = documents
        self._write_project_state(project_id, state)
        source_asset_id = self._ensure_local_source_asset(
            project_id=project_id,
            entry=normalized,
        )
        context = self._context()
        with self.runtime._connection() as connection:
            scope_id = self.runtime._local_object_scope_id(connection, context.sandbox_id)
            task = connection.execute(
                "SELECT client_id FROM tasks WHERE id=? AND scope_id=? "
                "AND lifecycle_state!='deleted'",
                (task_id, scope_id),
            ).fetchone()
            if task is None or str(task["client_id"] or "") != project_id:
                raise LocalRuntimeError(409, "task_attachment_project_mismatch", "任务附件与项目归属不一致")
            connection.execute(
                "UPDATE source_assets SET source_kind='task_attachment',"
                "source_locator_nonlocal=?,updated_at=? WHERE id=? AND scope_id=?",
                (f"task:{task_id}", utc_now(), source_asset_id, scope_id),
            )
            connection.commit()
        return {**normalized, "documentId": source_asset_id, "taskId": task_id}

    def task_attachments(self, task_id: str) -> list[dict[str, Any]]:
        context = self._context()
        with self.runtime._connection() as connection:
            scope_id = self.runtime._local_object_scope_id(connection, context.sandbox_id)
            task = connection.execute(
                "SELECT client_id FROM tasks WHERE id=? AND scope_id=? "
                "AND lifecycle_state!='deleted'",
                (task_id, scope_id),
            ).fetchone()
        if task is None or not str(task["client_id"] or ""):
            return []
        state = self._load_project_state(str(task["client_id"]))
        result: list[dict[str, Any]] = []
        for document_id, raw in dict(state.get("documents") or {}).items():
            if not isinstance(raw, Mapping) or str(raw.get("taskId") or "") != task_id:
                continue
            path = str(raw.get("managedPath") or "")
            media_type = str(raw.get("mediaType") or "application/octet-stream")
            is_audio = media_type.startswith("audio/") or Path(
                str(raw.get("fileName") or path)
            ).suffix.lower() in {
                ".m4a", ".mp3", ".wav", ".aac", ".flac", ".ogg", ".webm", ".mp4", ".mov"
            }
            transcript_path = ""
            transcript_object_id = str(raw.get("transcriptObjectId") or "")
            if transcript_object_id:
                manifest = self.runtime.local_storage_object_get(
                    sandbox_id=context.sandbox_id,
                    object_id=transcript_object_id,
                )
                if manifest is not None:
                    candidate = self._managed_path(str(manifest.get("storage_key") or ""))
                    if candidate.is_file():
                        transcript_path = str(candidate)
            result.append(
                {
                    "id": str(document_id),
                    "taskId": task_id,
                    "clientId": str(task["client_id"] or ""),
                    "title": str(raw.get("title") or raw.get("fileName") or "任务附件"),
                    "fileName": str(raw.get("fileName") or "任务附件"),
                    "mediaType": media_type,
                    "byteSize": int(raw.get("byteSize") or 0),
                    "contentHash": str(raw.get("contentHash") or ""),
                    "path": path,
                    "localAvailable": bool(path and Path(path).is_file()),
                    "sourceScope": "local_private",
                    "source": "task_attachment",
                    "isAudio": is_audio,
                    "processingStatus": (
                        str(raw.get("transcriptionStatus") or "not_requested")
                        if is_audio
                        else None
                    ),
                    "processingError": raw.get("transcriptionError"),
                    "processingProgress": int(raw.get("transcriptionProgress") or 0),
                    "processingStage": raw.get("transcriptionStage"),
                    "transcriptAttachmentId": raw.get("transcriptObjectId"),
                    "transcriptPath": transcript_path or None,
                    "version": int(raw.get("version") or 1),
                    "createdAt": raw.get("updatedAt"),
                    "updatedAt": raw.get("updatedAt"),
                }
            )
        return sorted(result, key=lambda item: (str(item.get("createdAt") or ""), item["id"]))

    def _task_attachment_entry(
        self,
        *,
        task_id: str,
        attachment_id: str,
    ) -> tuple[str, dict[str, Any], dict[str, Any]]:
        context = self._context()
        with self.runtime._connection() as connection:
            scope_id = self.runtime._local_object_scope_id(connection, context.sandbox_id)
            task = connection.execute(
                "SELECT client_id FROM tasks WHERE id=? AND scope_id=? "
                "AND lifecycle_state!='deleted'",
                (task_id, scope_id),
            ).fetchone()
        project_id = str(task["client_id"] or "") if task is not None else ""
        if not project_id:
            raise LocalRuntimeError(404, "task_missing", "任务不存在或没有项目归属")
        state = self._load_project_state(project_id)
        entry = (state.get("documents") or {}).get(attachment_id)
        if not isinstance(entry, dict) or str(entry.get("taskId") or "") != task_id:
            raise LocalRuntimeError(404, "task_attachment_missing", "任务附件不存在")
        return project_id, state, entry

    def set_task_transcription_state(
        self,
        *,
        task_id: str,
        attachment_id: str,
        status: str,
        error: str | None = None,
        progress: int | None = None,
        stage: str | None = None,
    ) -> dict[str, Any]:
        project_id, state, entry = self._task_attachment_entry(
            task_id=task_id, attachment_id=attachment_id
        )
        entry["transcriptionStatus"] = status
        entry["transcriptionError"] = error
        if progress is not None:
            entry["transcriptionProgress"] = max(0, min(100, int(progress)))
        if stage is not None:
            entry["transcriptionStage"] = str(stage)[:160]
        entry["updatedAt"] = utc_now()
        self._write_project_state(project_id, state)
        return dict(entry)

    def save_task_transcript(
        self,
        *,
        task_id: str,
        attachment_id: str,
        text: str,
        preserve_original: bool,
        expected_version: int | None = None,
    ) -> dict[str, Any]:
        project_id, state, entry = self._task_attachment_entry(
            task_id=task_id, attachment_id=attachment_id
        )
        normalized = str(text or "").strip()
        if not normalized:
            raise LocalRuntimeError(422, "task_transcript_empty", "转写文本不能为空")
        context = self._context()
        transcript_key = hashlib.sha256(
            f"{context.sandbox_id}\x1f{task_id}\x1f{attachment_id}".encode("utf-8")
        ).hexdigest()[:32]
        current_object_id = f"task-transcript-current-{transcript_key}"
        current = self.runtime.local_storage_object_get(
            sandbox_id=context.sandbox_id, object_id=current_object_id
        )
        current_version = int(current["version"]) if current is not None else 0
        source_stem = self._safe_name(
            Path(str(entry.get("fileName") or entry.get("title") or "录音")).stem
        )
        transcript_file_name = f"{source_stem}-录音转写.txt"
        storage_key = (
            f"task-transcripts/{self._stable_segment(context.sandbox_id)}/"
            f"{transcript_key}/v{current_version + 1}/{transcript_file_name}"
        )
        if expected_version is not None and current_version != expected_version:
            raise LocalRuntimeError(409, "task_transcript_version_conflict", "转写稿已变化，请刷新后重试")
        if preserve_original and not entry.get("originalTranscriptObjectId"):
            original_object_id = f"task-transcript-original-{transcript_key}"
            self._write_object(
                sandbox_id=context.sandbox_id,
                object_id=original_object_id,
                storage_key=(
                    f"task-transcripts/{self._stable_segment(context.sandbox_id)}/"
                    f"{transcript_key}/original/{transcript_file_name}"
                ),
                media_type="text/plain; charset=utf-8",
                data=normalized.encode("utf-8"),
                expected_version=0,
            )
            entry["originalTranscriptObjectId"] = original_object_id
        stored = self._write_object(
            sandbox_id=context.sandbox_id,
            object_id=current_object_id,
            storage_key=storage_key,
            media_type="text/plain; charset=utf-8",
            data=normalized.encode("utf-8"),
            expected_version=current_version,
        )
        entry["transcriptObjectId"] = current_object_id
        entry["transcriptionStatus"] = "ready"
        entry["transcriptionError"] = None
        entry["transcriptionProgress"] = 100
        entry["transcriptionStage"] = "转写完成"
        entry["updatedAt"] = stored["updatedAt"]
        self._write_project_state(project_id, state)
        return self.task_transcript(task_id=task_id, attachment_id=attachment_id)

    def task_transcript(self, *, task_id: str, attachment_id: str) -> dict[str, Any]:
        _project_id, _state, entry = self._task_attachment_entry(
            task_id=task_id, attachment_id=attachment_id
        )
        context = self._context()

        def read_object(object_id: str) -> tuple[str, int]:
            row = self.runtime.local_storage_object_get(
                sandbox_id=context.sandbox_id, object_id=object_id
            )
            if row is None or str(row.get("lifecycle_state") or "") != "active":
                raise LocalRuntimeError(404, "task_transcript_missing", "当前设备没有该转写稿")
            path = self._managed_path(str(row.get("storage_key") or ""))
            try:
                return path.read_text(encoding="utf-8"), int(row["version"])
            except OSError as exc:
                raise LocalRuntimeError(409, "task_transcript_unavailable", "转写稿文件暂时不可读取") from exc

        current_id = str(entry.get("transcriptObjectId") or "")
        if not current_id:
            raise LocalRuntimeError(404, "task_transcript_missing", "该录音尚未完成转写")
        current_text, version = read_object(current_id)
        original_id = str(entry.get("originalTranscriptObjectId") or current_id)
        original_text, _ = read_object(original_id)
        current_row = self.runtime.local_storage_object_get(
            sandbox_id=context.sandbox_id,
            object_id=current_id,
        )
        current_path = (
            self._managed_path(str(current_row.get("storage_key") or ""))
            if current_row is not None
            else None
        )
        return {
            "sourceAttachmentId": attachment_id,
            "transcriptAttachmentId": current_id,
            "transcriptDocumentId": current_id,
            "originalText": original_text,
            "currentText": current_text,
            "version": version,
            "path": str(current_path) if current_path is not None and current_path.is_file() else None,
        }

    def delete_task_attachment_local(
        self, *, task_id: str, attachment_id: str
    ) -> dict[str, Any]:
        context = self._context()
        with self.runtime._connection() as connection:
            scope_id = self.runtime._local_object_scope_id(connection, context.sandbox_id)
            row = connection.execute(
                "SELECT client_id FROM source_assets WHERE id=? AND scope_id=? "
                "AND source_kind='task_attachment' AND source_locator_nonlocal=? "
                "AND lifecycle_state='active'",
                (attachment_id, scope_id, f"task:{task_id}"),
            ).fetchone()
        if row is None:
            raise LocalRuntimeError(404, "task_attachment_missing", "任务附件不存在")
        return self.delete_document_local(str(row["client_id"]), attachment_id)

    def ensure_project_projection(
        self,
        project: Mapping[str, Any],
    ) -> None:
        """Keep the cloud-authoritative client identity as a local projection.

        Local ``source_assets`` must point at the real client row.  The owner is
        supplied by the cloud payload; this method must never guess it from the
        current user or from a project name.
        """
        project_id = str(project.get("projectId") or "").strip()
        owner_membership_id = str(
            project.get("ownerMembershipId") or ""
        ).strip()
        if not project_id or not owner_membership_id:
            raise LocalRuntimeError(
                409,
                "project_projection_identity_missing",
                "组织云项目缺少稳定归属信息，暂不能启动本机资料处理",
            )
        authorization = project.get("authorizationProjection")
        if not isinstance(authorization, Mapping):
            raise LocalRuntimeError(
                409,
                "project_authorization_projection_missing",
                "组织云项目缺少权限投影，暂不能在本机使用",
            )
        context = self._context()
        policy_version_id = str(
            authorization.get("policyVersionId") or ""
        ).strip()
        viewer_principal_id = str(
            authorization.get("viewerPrincipalId") or ""
        ).strip()
        viewer_membership_id = str(
            authorization.get("viewerMembershipId") or ""
        ).strip()
        if (
            not policy_version_id
            or viewer_principal_id != context.principal_id
            or viewer_membership_id != context.membership_id
        ):
            raise LocalRuntimeError(
                409,
                "project_authorization_scope_mismatch",
                "项目权限投影与当前登录身份不一致",
            )
        viewer_surfaces = authorization.get("viewerSurfaces")
        viewer_capabilities = authorization.get("viewerCapabilities")
        policy_spec = authorization.get("policySpec")
        if (
            not isinstance(viewer_surfaces, list)
            or not isinstance(viewer_capabilities, list)
            or not isinstance(policy_spec, Mapping)
            or "read" not in viewer_capabilities
        ):
            raise LocalRuntimeError(
                409,
                "project_authorization_projection_invalid",
                "项目权限投影内容无效",
            )
        now = utc_now()
        with self.runtime._connection() as connection:
            scope_id = self.runtime._local_object_scope_id(
                connection,
                context.sandbox_id,
            )
            owner = connection.execute(
                "SELECT id FROM organization_memberships "
                "WHERE id=? AND scope_id=? AND lifecycle_state='active'",
                (owner_membership_id, scope_id),
            ).fetchone()
            if owner is None:
                raise LocalRuntimeError(
                    409,
                    "project_owner_projection_missing",
                    "本机尚未同步项目负责人的组织身份",
                )
            viewer = connection.execute(
                "SELECT principal_id FROM organization_memberships "
                "WHERE id=? AND scope_id=? AND status='active' "
                "AND lifecycle_state='active'",
                (viewer_membership_id, scope_id),
            ).fetchone()
            if (
                viewer is None
                or str(viewer["principal_id"] or "") != viewer_principal_id
            ):
                raise LocalRuntimeError(
                    409,
                    "project_viewer_projection_missing",
                    "本机尚未同步当前项目成员身份",
                )
            policy_version = max(
                1,
                int(authorization.get("policyVersion") or 1),
            )
            source_version = max(
                1,
                int(authorization.get("sourceVersion") or policy_version),
            )
            generated_at = str(authorization.get("generatedAt") or now)
            lease_expires_at = _bounded_project_lease(
                authorization.get("leaseExpiresAt")
            )
            viewer_projection_id = self.runtime._stable_id(
                "viewer_client",
                context.sandbox_id,
                project_id,
                viewer_membership_id,
            )
            connection.execute("BEGIN IMMEDIATE")
            connection.execute(
                """
                INSERT INTO secured_resources (
                    id, scope_id, resource_kind, lifecycle_state, version,
                    resource_type_key, created_at, updated_at, deleted_at,
                    authority_role, origin_instance_id
                ) VALUES (?, ?, 'client', 'active', ?, 'client', ?, ?, NULL,
                          'cloud_projection', ?)
                ON CONFLICT(id) DO UPDATE SET
                    scope_id=excluded.scope_id,
                    resource_kind='client',
                    lifecycle_state=excluded.lifecycle_state,
                    version=excluded.version,
                    resource_type_key='client',
                    updated_at=excluded.updated_at,
                    deleted_at=NULL,
                    authority_role='cloud_projection',
                    origin_instance_id=excluded.origin_instance_id
                """,
                (
                    project_id,
                    scope_id,
                    max(1, int(project.get("version") or 1)),
                    str(project.get("createdAt") or now),
                    now,
                    context.cloud_instance_id,
                ),
            )
            connection.execute(
                """
                INSERT INTO policy_versions (
                    id, scope_id, secured_resource_id, policy_scope_kind,
                    version, policy_spec_schema_version, policy_spec,
                    effective_at, created_at, lifecycle_state, updated_at,
                    deleted_at, sandbox_id, source_version,
                    projection_state, projected_at, stale_at, lease_expires_at
                ) VALUES (?, ?, ?, 'secured_resource', ?, ?, ?, ?, ?, 'active',
                          ?, NULL, ?, ?, 'fresh', ?, NULL, ?)
                ON CONFLICT(id) DO UPDATE SET
                    scope_id=excluded.scope_id,
                    secured_resource_id=excluded.secured_resource_id,
                    policy_scope_kind='secured_resource',
                    version=excluded.version,
                    policy_spec_schema_version=excluded.policy_spec_schema_version,
                    policy_spec=excluded.policy_spec,
                    effective_at=excluded.effective_at,
                    lifecycle_state='active',
                    updated_at=excluded.updated_at,
                    deleted_at=NULL,
                    sandbox_id=excluded.sandbox_id,
                    source_version=excluded.source_version,
                    projection_state='fresh',
                    projected_at=excluded.projected_at,
                    stale_at=NULL,
                    lease_expires_at=excluded.lease_expires_at
                """,
                (
                    policy_version_id,
                    scope_id,
                    project_id,
                    policy_version,
                    str(
                        authorization.get("policySpecSchemaVersion")
                        or "gc02.client-access.v1"
                    ),
                    canonical_json(dict(policy_spec)),
                    generated_at,
                    generated_at,
                    now,
                    context.sandbox_id,
                    source_version,
                    now,
                    lease_expires_at,
                ),
            )
            connection.execute(
                """
                INSERT INTO clients (
                    id, scope_id, owner_principal_id, owner_membership_id,
                    lifecycle_state, version, name, alias, summary, domain,
                    color, visibility_scope, is_default_internal, archived_at,
                    created_at, updated_at, deleted_at, sandbox_id,
                    source_version, projection_state, projected_at, stale_at,
                    lease_expires_at
                ) VALUES (?, ?, NULL, ?, ?, ?, ?, ?, ?, ?, ?, 'organization',
                          ?, NULL, ?, ?, NULL, ?, ?, 'current', ?, NULL, ?)
                ON CONFLICT(id) DO UPDATE SET
                    scope_id=excluded.scope_id,
                    owner_principal_id=NULL,
                    owner_membership_id=excluded.owner_membership_id,
                    lifecycle_state=excluded.lifecycle_state,
                    version=excluded.version,
                    name=excluded.name,
                    alias=excluded.alias,
                    summary=excluded.summary,
                    domain=excluded.domain,
                    color=excluded.color,
                    visibility_scope='organization',
                    is_default_internal=excluded.is_default_internal,
                    archived_at=NULL,
                    updated_at=excluded.updated_at,
                    deleted_at=NULL,
                    sandbox_id=excluded.sandbox_id,
                    source_version=excluded.source_version,
                    projection_state='current',
                    projected_at=excluded.projected_at,
                    stale_at=NULL,
                    lease_expires_at=excluded.lease_expires_at
                """,
                (
                    project_id,
                    scope_id,
                    owner_membership_id,
                    str(project.get("lifecycleState") or "active"),
                    max(1, int(project.get("version") or 1)),
                    str(project.get("name") or ""),
                    str(project.get("alias") or ""),
                    str(project.get("summary") or ""),
                    str(project.get("domain") or "项目"),
                    str(project.get("color") or "#5B7BFE"),
                    1 if project.get("isDefaultInternalProject") else 0,
                    str(project.get("createdAt") or now),
                    now,
                    context.sandbox_id,
                    max(1, int(project.get("version") or 1)),
                    now,
                    lease_expires_at,
                ),
            )
            connection.execute(
                """
                UPDATE viewer_projections
                SET invalidated_at=?, projection_state='stale', stale_at=?
                WHERE scope_id=? AND secured_resource_id=?
                  AND viewer_membership_id=? AND sandbox_id=? AND id!=?
                  AND invalidated_at IS NULL
                """,
                (
                    now,
                    now,
                    scope_id,
                    project_id,
                    viewer_membership_id,
                    context.sandbox_id,
                    viewer_projection_id,
                ),
            )
            connection.execute(
                """
                INSERT INTO viewer_projections (
                    id, scope_id, secured_resource_id, viewer_principal_id,
                    viewer_membership_id, policy_version_id,
                    viewer_surfaces, viewer_capabilities,
                    viewer_surfaces_schema_version,
                    viewer_capabilities_schema_version, lease_expires_at,
                    generated_at, source_version, invalidated_at, sandbox_id,
                    projection_state, projected_at, stale_at
                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, '1', '1', ?, ?, ?, NULL,
                          ?, 'fresh', ?, NULL)
                ON CONFLICT(id) DO UPDATE SET
                    scope_id=excluded.scope_id,
                    secured_resource_id=excluded.secured_resource_id,
                    viewer_principal_id=excluded.viewer_principal_id,
                    viewer_membership_id=excluded.viewer_membership_id,
                    policy_version_id=excluded.policy_version_id,
                    viewer_surfaces=excluded.viewer_surfaces,
                    viewer_capabilities=excluded.viewer_capabilities,
                    lease_expires_at=excluded.lease_expires_at,
                    generated_at=excluded.generated_at,
                    source_version=excluded.source_version,
                    invalidated_at=NULL,
                    sandbox_id=excluded.sandbox_id,
                    projection_state='fresh',
                    projected_at=excluded.projected_at,
                    stale_at=NULL
                """,
                (
                    viewer_projection_id,
                    scope_id,
                    project_id,
                    viewer_principal_id,
                    viewer_membership_id,
                    policy_version_id,
                    canonical_json(sorted({str(item) for item in viewer_surfaces})),
                    canonical_json(
                        sorted({str(item) for item in viewer_capabilities})
                    ),
                    lease_expires_at,
                    generated_at,
                    source_version,
                    context.sandbox_id,
                    now,
                ),
            )
            connection.commit()

    def _ensure_local_source_asset(
        self,
        *,
        project_id: str,
        entry: Mapping[str, Any],
    ) -> str:
        context = self._context()
        local_object_id = str(entry.get("localSourceId") or "").strip()
        source_asset_id = str(
            entry.get("cloudDocumentId")
            or entry.get("documentId")
            or ""
        ).strip()
        if source_asset_id.startswith("local-pending:"):
            source_asset_id = local_object_id
        if not local_object_id or not source_asset_id:
            raise LocalRuntimeError(
                409,
                "local_source_asset_identity_missing",
                "本机资料缺少稳定来源标识",
            )
        source = self.runtime.local_storage_object_get(
            sandbox_id=context.sandbox_id,
            object_id=local_object_id,
        )
        if source is None:
            raise LocalRuntimeError(
                404,
                "local_document_source_missing",
                "当前设备没有该资料的源文件",
            )
        now = utc_now()
        with self.runtime._connection() as connection:
            scope_id = self.runtime._local_object_scope_id(
                connection,
                context.sandbox_id,
            )
            project = connection.execute(
                "SELECT id FROM clients WHERE id=? AND scope_id=? "
                "AND lifecycle_state='active'",
                (project_id, scope_id),
            ).fetchone()
            if project is None:
                raise LocalRuntimeError(
                    409,
                    "project_projection_missing",
                    "当前项目尚未形成可验证的本机投影",
                )
            connection.execute("BEGIN IMMEDIATE")
            connection.execute(
                """
                INSERT INTO secured_resources (
                    id, scope_id, resource_kind, lifecycle_state, version,
                    resource_type_key, created_at, updated_at, deleted_at,
                    authority_role, origin_instance_id
                ) VALUES (?, ?, 'source_asset', 'active', 1,
                          'local_original', ?, ?, NULL, 'cloud_projection', ?)
                ON CONFLICT(id) DO UPDATE SET
                    lifecycle_state='active', updated_at=excluded.updated_at,
                    deleted_at=NULL, authority_role='cloud_projection',
                    origin_instance_id=excluded.origin_instance_id
                """,
                (
                    source_asset_id,
                    scope_id,
                    now,
                    now,
                    context.cloud_instance_id,
                ),
            )
            connection.execute(
                """
                INSERT INTO source_assets (
                    id, scope_id, client_id, object_manifest_id, content_hash,
                    record_kind, source_kind, display_name, media_type,
                    byte_size, source_locator_nonlocal, parent_folder_id,
                    asset_id, folder_id, created_by_membership_id,
                    availability_state, archived_at, version, lifecycle_state,
                    created_at, updated_at, deleted_at, authority_role,
                    origin_instance_id
                ) VALUES (?, ?, ?, ?, ?, 'asset', 'local_original', ?, ?, ?,
                          NULL, NULL, NULL, NULL, ?, 'ready', NULL, 1, 'active',
                          ?, ?, NULL, 'cloud_projection', ?)
                ON CONFLICT(id) DO UPDATE SET
                    scope_id=excluded.scope_id,
                    client_id=excluded.client_id,
                    object_manifest_id=excluded.object_manifest_id,
                    version=CASE
                        WHEN source_assets.content_hash != excluded.content_hash
                        THEN source_assets.version + 1
                        ELSE source_assets.version
                    END,
                    content_hash=excluded.content_hash,
                    source_kind='local_original',
                    display_name=excluded.display_name,
                    media_type=excluded.media_type,
                    byte_size=excluded.byte_size,
                    created_by_membership_id=excluded.created_by_membership_id,
                    availability_state='ready',
                    lifecycle_state='active',
                    updated_at=excluded.updated_at,
                    deleted_at=NULL,
                    authority_role='cloud_projection',
                    origin_instance_id=excluded.origin_instance_id
                """,
                (
                    source_asset_id,
                    scope_id,
                    project_id,
                    str(source["manifest_id"]),
                    str(source["content_hash"]),
                    str(entry.get("fileName") or ""),
                    str(entry.get("mediaType") or source["media_type"] or ""),
                    int(entry.get("byteSize") or source["byte_size"] or 0),
                    context.membership_id,
                    str(entry.get("updatedAt") or now),
                    now,
                    context.cloud_instance_id,
                ),
            )
            connection.commit()
        return source_asset_id

    @staticmethod
    def _processing_error_state(exc: LocalRuntimeError) -> tuple[str, bool]:
        if exc.code in {
            "local_document_ocr_required",
            "local_document_preview_unsupported",
            "local_document_pdf_encrypted",
            "local_document_encoding_unsupported",
            "local_document_source_missing",
            "local_storage_object_missing",
        }:
            return "blocked", False
        if exc.status_code >= 500:
            return "failed_retryable", True
        return "blocked", False

    def _latest_processing_attempt(
        self,
        source_asset_id: str,
        *,
        processor_kind: str,
    ) -> dict[str, Any] | None:
        context = self._context()
        with self.runtime._connection() as connection:
            scope_id = self.runtime._local_object_scope_id(
                connection,
                context.sandbox_id,
            )
            row = connection.execute(
                """
                SELECT * FROM processing_attempts
                WHERE scope_id=? AND source_asset_id=? AND processor_kind=?
                ORDER BY attempt_no DESC, started_at DESC, id DESC
                LIMIT 1
                """,
                (scope_id, source_asset_id, processor_kind),
            ).fetchone()
        return dict(row) if row is not None else None

    def processing_state(self, entry: Mapping[str, Any]) -> dict[str, Any]:
        source_asset_id = str(
            entry.get("cloudDocumentId")
            or entry.get("documentId")
            or ""
        ).strip()
        if source_asset_id.startswith("local-pending:"):
            source_asset_id = str(entry.get("localSourceId") or "").strip()
        if not source_asset_id:
            return {
                "parseStatus": "blocked",
                "wikiStatus": "not_requested",
                "processingErrorCode": "local_source_asset_identity_missing",
                "processingMessage": "本机资料缺少稳定来源标识",
                "processingRetryable": False,
            }
        parsed = self._latest_processing_attempt(
            source_asset_id,
            processor_kind="local_text_extraction",
        )
        wiki = self._latest_processing_attempt(
            source_asset_id,
            processor_kind="local_wiki_projection",
        )
        if parsed is None:
            return {
                "parseStatus": "not_requested",
                "wikiStatus": "not_requested",
                "processingErrorCode": None,
                "processingMessage": "等待本机解析",
                "processingRetryable": True,
            }
        parsed_status = str(parsed.get("status") or "not_requested")
        wiki_status = str((wiki or {}).get("status") or "not_requested")
        active_error = parsed
        if parsed_status == "ready" and wiki_status in {
            "blocked",
            "failed_retryable",
        }:
            active_error = wiki or parsed
        return {
            "parseStatus": parsed_status,
            "wikiStatus": wiki_status,
            "processingAttemptId": str(parsed.get("id") or ""),
            "processingAttemptNo": int(parsed.get("attempt_no") or 0),
            "processingStage": (
                "wiki" if active_error is wiki else "text_extraction"
            ),
            "processingErrorCode": active_error.get("error_code") or None,
            "processingMessage": (
                active_error.get("error_message_safe") or None
            ),
            "processingRetryable": parsed_status
            in {"not_requested", "failed_retryable"}
            or wiki_status in {"not_requested", "failed_retryable"},
            "processedAt": (
                (wiki or {}).get("finished_at")
                or (wiki or {}).get("started_at")
                or parsed.get("finished_at")
                or parsed.get("started_at")
            ),
        }

    def process_document(
        self,
        *,
        project_id: str,
        document_id: str,
        force: bool = False,
    ) -> dict[str, Any]:
        context = self._context()
        _, state, entry = self._document_entry(document_id)
        if str(state.get("projectId") or "") != project_id:
            raise LocalRuntimeError(
                409,
                "local_document_project_mismatch",
                "本机资料不属于当前项目",
            )
        source_id = self._ensure_local_source_asset(
            project_id=project_id,
            entry=entry,
        )
        current = self._latest_processing_attempt(
            source_id,
            processor_kind="local_text_extraction",
        )
        if current is not None and not force and str(current.get("status")) in {
            "ready",
            "blocked",
        }:
            return {"documentId": document_id, **self.processing_state(entry)}
        attempt_no = int((current or {}).get("attempt_no") or 0) + 1
        attempt_id = new_id()
        now = utc_now()
        with self.runtime._connection() as connection:
            scope_id = self.runtime._local_object_scope_id(
                connection,
                context.sandbox_id,
            )
            connection.execute(
                """
                INSERT INTO processing_attempts (
                    id, scope_id, operation_id, source_asset_id, recording_id,
                    attempt_no, status, error_code, processor_kind,
                    provider_resource_id, error_message_safe, next_retry_at,
                    started_at, finished_at, authority_role, origin_instance_id
                ) VALUES (?, ?, NULL, ?, NULL, ?, 'processing', NULL,
                          'local_text_extraction', NULL, NULL, NULL, ?, NULL,
                          'local', ?)
                """,
                (
                    attempt_id,
                    scope_id,
                    source_id,
                    attempt_no,
                    now,
                    context.sandbox_id,
                ),
            )
            connection.commit()
        try:
            local = self.document_text(document_id)
            content = str(local.get("content") or "").strip()
            if not content:
                raise LocalRuntimeError(
                    422,
                    "local_document_empty",
                    "本机资料没有可读取正文",
                )
        except LocalRuntimeError as exc:
            status, retryable = self._processing_error_state(exc)
            finished = utc_now()
            with self.runtime._connection() as connection:
                connection.execute(
                    """
                    UPDATE processing_attempts
                    SET status=?, error_code=?, error_message_safe=?,
                        next_retry_at=?, finished_at=?
                    WHERE id=? AND authority_role='local'
                    """,
                    (
                        status,
                        exc.code,
                        exc.message,
                        finished if retryable else None,
                        finished,
                        attempt_id,
                    ),
                )
                connection.commit()
            return {"documentId": document_id, **self.processing_state(entry)}
        except Exception:
            finished = utc_now()
            with self.runtime._connection() as connection:
                connection.execute(
                    """
                    UPDATE processing_attempts
                    SET status='failed_retryable',
                        error_code='local_parser_exception',
                        error_message_safe='本机解析暂时失败，可以重试',
                        next_retry_at=?, finished_at=?
                    WHERE id=? AND authority_role='local'
                    """,
                    (finished, finished, attempt_id),
                )
                connection.commit()
            return {"documentId": document_id, **self.processing_state(entry)}
        finished = utc_now()
        with self.runtime._connection() as connection:
            connection.execute("BEGIN IMMEDIATE")
            connection.execute(
                """
                UPDATE processing_attempts
                SET status='ready', error_code=NULL, error_message_safe=NULL,
                    next_retry_at=NULL, finished_at=?
                WHERE id=? AND authority_role='local'
                """,
                (finished, attempt_id),
            )
            wiki = connection.execute(
                """
                SELECT id FROM processing_attempts
                WHERE scope_id=? AND source_asset_id=?
                  AND processor_kind='local_wiki_projection'
                  AND status IN ('queued','processing','ready')
                ORDER BY attempt_no DESC LIMIT 1
                """,
                (scope_id, source_id),
            ).fetchone()
            if wiki is None:
                connection.execute(
                    """
                    INSERT INTO processing_attempts (
                        id, scope_id, operation_id, source_asset_id,
                        recording_id, attempt_no, status, error_code,
                        processor_kind, provider_resource_id,
                        error_message_safe, next_retry_at, started_at,
                        finished_at, authority_role, origin_instance_id
                    ) VALUES (?, ?, NULL, ?, NULL, 1, 'queued', NULL,
                              'local_wiki_projection', NULL, NULL, NULL, ?,
                              NULL, 'local', ?)
                    """,
                    (new_id(), scope_id, source_id, finished, context.sandbox_id),
                )
            connection.commit()
        return {
            "documentId": document_id,
            "contentCharacterCount": len(content),
            **self.processing_state(entry),
        }

    def build_local_wiki_document(
        self,
        *,
        project_id: str,
        document_id: str,
        force: bool = False,
    ) -> dict[str, Any]:
        """Build the private local Wiki projection without sharing body text."""
        context = self._context()
        _, state, entry = self._document_entry(document_id)
        if str(state.get("projectId") or "") != project_id:
            raise LocalRuntimeError(
                409,
                "local_document_project_mismatch",
                "本机资料不属于当前项目",
            )
        source_asset_id = self._ensure_local_source_asset(
            project_id=project_id,
            entry=entry,
        )
        parsed = self._latest_processing_attempt(
            source_asset_id,
            processor_kind="local_text_extraction",
        )
        if parsed is None or str(parsed.get("status") or "") != "ready":
            return {"documentId": document_id, **self.processing_state(entry)}

        current_attempt = self._latest_processing_attempt(
            source_asset_id,
            processor_kind="local_wiki_projection",
        )
        if (
            current_attempt is not None
            and not force
            and str(current_attempt.get("status") or "") == "ready"
        ):
            return {"documentId": document_id, **self.processing_state(entry)}

        attempt_id = str((current_attempt or {}).get("id") or "")
        attempt_no = int((current_attempt or {}).get("attempt_no") or 0)
        now = utc_now()
        with self.runtime._connection() as connection:
            scope_id = self.runtime._local_object_scope_id(
                connection,
                context.sandbox_id,
            )
            if (
                current_attempt is not None
                and str(current_attempt.get("status") or "") == "queued"
                and not force
            ):
                connection.execute(
                    """
                    UPDATE processing_attempts
                    SET status='processing', error_code=NULL,
                        error_message_safe=NULL, next_retry_at=NULL,
                        started_at=?, finished_at=NULL
                    WHERE id=? AND authority_role='local'
                    """,
                    (now, attempt_id),
                )
            else:
                attempt_id = new_id()
                attempt_no += 1
                connection.execute(
                    """
                    INSERT INTO processing_attempts (
                        id, scope_id, operation_id, source_asset_id,
                        recording_id, attempt_no, status, error_code,
                        processor_kind, provider_resource_id,
                        error_message_safe, next_retry_at, started_at,
                        finished_at, authority_role, origin_instance_id
                    ) VALUES (?, ?, NULL, ?, NULL, ?, 'processing', NULL,
                              'local_wiki_projection', NULL, NULL, NULL, ?,
                              NULL, 'local', ?)
                    """,
                    (
                        attempt_id,
                        scope_id,
                        source_asset_id,
                        attempt_no,
                        now,
                        context.sandbox_id,
                    ),
                )
            connection.commit()

        try:
            local = self.document_text(document_id)
            content = str(local.get("content") or "").replace(
                "\r\n", "\n"
            ).replace("\r", "\n").strip()
            if not content:
                raise LocalRuntimeError(
                    422,
                    "local_document_empty",
                    "本机资料没有可加工正文",
                )
            chunks = self._chunk_wiki_text(content)
            if not chunks:
                raise LocalRuntimeError(
                    422,
                    "local_wiki_chunks_empty",
                    "本机资料没有形成可用知识分块",
                )
            content_hash = hashlib.sha256(content.encode("utf-8")).hexdigest()
            knowledge_document_id = (
                "knowledge_local_"
                + self._stable_segment(
                    f"{context.sandbox_id}|{source_asset_id}"
                )
            )

            with self.runtime._connection() as connection:
                scope_id = self.runtime._local_object_scope_id(
                    connection,
                    context.sandbox_id,
                )
                source_row = connection.execute(
                    """
                    SELECT version, content_hash
                    FROM source_assets
                    WHERE id=? AND scope_id=? AND client_id=?
                      AND lifecycle_state='active'
                    """,
                    (source_asset_id, scope_id, project_id),
                ).fetchone()
                if source_row is None:
                    raise LocalRuntimeError(
                        409,
                        "local_source_asset_projection_missing",
                        "本机来源投影缺失，暂不能构建知识",
                    )
                source_version = int(source_row["version"] or 1)
                current = connection.execute(
                    """
                    SELECT d.current_version, d.version AS aggregate_version,
                           v.id AS version_id, v.content_hash
                    FROM knowledge_documents d
                    LEFT JOIN document_versions v
                      ON v.scope_id=d.scope_id AND v.document_id=d.id
                     AND v.version=d.current_version
                    WHERE d.id=? AND d.scope_id=?
                    """,
                    (knowledge_document_id, scope_id),
                ).fetchone()
            same_content = (
                current is not None
                and str(current["content_hash"] or "") == content_hash
            )
            content_version = (
                int(current["current_version"] or 0)
                if same_content
                else int((current or {})["current_version"] or 0) + 1
                if current is not None
                else 1
            )
            aggregate_version = (
                int(current["aggregate_version"] or 1)
                if same_content
                else int(current["aggregate_version"] or 1) + 1
                if current is not None
                else 1
            )
            document_version_id = (
                "document_version_local_"
                + self._stable_segment(
                    f"{knowledge_document_id}|{content_version}|{content_hash}"
                )
            )
            source_set_id = (
                "source_set_local_"
                + self._stable_segment(document_version_id)
            )
            source_set_member_id = (
                "source_member_local_"
                + self._stable_segment(f"{source_set_id}|{source_asset_id}")
            )
            lineage_id = (
                "lineage_local_"
                + self._stable_segment(document_version_id)
            )
            search_manifest_id = (
                "search_manifest_local_"
                + self._stable_segment(document_version_id)
            )
            vector_manifest_id = (
                "vector_manifest_local_"
                + self._stable_segment(document_version_id)
            )
            base = (
                "local-wiki/"
                f"{self._stable_segment(context.sandbox_id)}/"
                f"{self._stable_segment(project_id)}"
            )
            document_object = self._ensure_wiki_object(
                object_id=f"wiki-document:{document_version_id}",
                storage_key=f"{base}/documents/{document_version_id}.txt",
                media_type=self.WIKI_DOCUMENT_MEDIA_TYPE,
                content=content,
            )
            chunk_objects: list[dict[str, Any]] = []
            for chunk in chunks:
                chunk_hash = hashlib.sha256(
                    str(chunk["content"]).encode("utf-8")
                ).hexdigest()
                chunk_id = (
                    "content_chunk_local_"
                    + self._stable_segment(
                        f"{document_version_id}|{chunk['ordinal']}|{chunk_hash}"
                    )
                )
                stored = self._ensure_wiki_object(
                    object_id=f"wiki-chunk:{chunk_id}",
                    storage_key=f"{base}/chunks/{chunk_id}.txt",
                    media_type=self.WIKI_CHUNK_MEDIA_TYPE,
                    content=str(chunk["content"]),
                )
                chunk_objects.append(
                    {
                        **chunk,
                        "id": chunk_id,
                        "hash": chunk_hash,
                        "manifestId": stored["manifestId"],
                    }
                )

            finished = utc_now()
            title = str(
                entry.get("title")
                or entry.get("fileName")
                or local.get("title")
                or "本机资料"
            )
            with self.runtime._connection() as connection:
                connection.execute("BEGIN IMMEDIATE")
                scope_id = self.runtime._local_object_scope_id(
                    connection,
                    context.sandbox_id,
                )
                previous_version_id = (
                    str(current["version_id"] or "")
                    if current is not None
                    else ""
                )
                if previous_version_id and previous_version_id != document_version_id:
                    connection.execute(
                        """
                        UPDATE content_chunks
                        SET lifecycle_state='deleted', deleted_at=?,
                            updated_at=?, version=version+1
                        WHERE scope_id=? AND document_version_id=?
                          AND lifecycle_state='active'
                        """,
                        (finished, finished, scope_id, previous_version_id),
                    )
                    previous_lineages = connection.execute(
                        """
                        SELECT id FROM derivation_lineage
                        WHERE scope_id=? AND derivative_kind='local_wiki_version'
                          AND derivative_object_id=? AND invalidated_at IS NULL
                        """,
                        (scope_id, previous_version_id),
                    ).fetchall()
                    for row in previous_lineages:
                        previous_lineage_id = str(row["id"])
                        connection.execute(
                            """
                            UPDATE search_index_manifests
                            SET status='invalidated', invalidated_at=?
                            WHERE scope_id=? AND lineage_id=?
                              AND invalidated_at IS NULL
                            """,
                            (finished, scope_id, previous_lineage_id),
                        )
                        connection.execute(
                            """
                            UPDATE vector_index_manifests
                            SET status='invalidated', invalidated_at=?
                            WHERE scope_id=? AND lineage_id=?
                              AND invalidated_at IS NULL
                            """,
                            (finished, scope_id, previous_lineage_id),
                        )
                    connection.execute(
                        """
                        UPDATE derivation_lineage SET invalidated_at=?
                        WHERE scope_id=? AND derivative_kind='local_wiki_version'
                          AND derivative_object_id=? AND invalidated_at IS NULL
                        """,
                        (finished, scope_id, previous_version_id),
                    )

                connection.execute(
                    """
                    INSERT INTO secured_resources (
                        id, scope_id, resource_kind, lifecycle_state, version,
                        resource_type_key, created_at, updated_at, deleted_at,
                        authority_role, origin_instance_id
                    ) VALUES (?, ?, 'knowledge_document', 'active', ?,
                              'local_private_wiki', ?, ?, NULL, 'local', ?)
                    ON CONFLICT(id) DO UPDATE SET
                        lifecycle_state='active', version=excluded.version,
                        updated_at=excluded.updated_at, deleted_at=NULL,
                        authority_role='local',
                        origin_instance_id=excluded.origin_instance_id
                    """,
                    (
                        knowledge_document_id,
                        scope_id,
                        aggregate_version,
                        finished,
                        finished,
                        context.sandbox_id,
                    ),
                )
                connection.execute(
                    """
                    INSERT INTO knowledge_documents (
                        id, scope_id, source_asset_id, client_id,
                        current_version, owner_membership_id, title,
                        document_kind, visibility_scope, parse_state,
                        publication_state, published_at, version,
                        lifecycle_state, created_at, updated_at, deleted_at,
                        sandbox_id, source_version, projection_state,
                        projected_at, stale_at, lease_expires_at
                    ) VALUES (?, ?, ?, ?, ?, ?, ?, 'local_private_wiki',
                              'self', 'ready', 'draft', NULL, ?, 'active', ?,
                              ?, NULL, ?, ?, 'current', ?, NULL, NULL)
                    ON CONFLICT(id) DO UPDATE SET
                        source_asset_id=excluded.source_asset_id,
                        client_id=excluded.client_id,
                        current_version=excluded.current_version,
                        owner_membership_id=excluded.owner_membership_id,
                        title=excluded.title,
                        document_kind='local_private_wiki',
                        visibility_scope='self', parse_state='ready',
                        publication_state='draft', version=excluded.version,
                        lifecycle_state='active', updated_at=excluded.updated_at,
                        deleted_at=NULL, sandbox_id=excluded.sandbox_id,
                        source_version=excluded.source_version,
                        projection_state='current',
                        projected_at=excluded.projected_at,
                        stale_at=NULL, lease_expires_at=NULL
                    """,
                    (
                        knowledge_document_id,
                        scope_id,
                        source_asset_id,
                        project_id,
                        content_version,
                        context.membership_id,
                        title,
                        aggregate_version,
                        finished,
                        finished,
                        context.sandbox_id,
                        source_version,
                        finished,
                    ),
                )
                connection.execute(
                    """
                    INSERT INTO document_versions (
                        id, scope_id, document_id, version, content_hash,
                        created_at, object_manifest_id, source_asset_version,
                        publication_state, created_by_membership_id,
                        origin_instance_id, integrity_hash, sandbox_id,
                        source_version, projection_state, projected_at,
                        stale_at, lease_expires_at
                    ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, 'draft', ?, ?, ?, ?, ?,
                              'current', ?, NULL, NULL)
                    ON CONFLICT(id) DO UPDATE SET
                        content_hash=excluded.content_hash,
                        object_manifest_id=excluded.object_manifest_id,
                        source_asset_version=excluded.source_asset_version,
                        publication_state='draft',
                        created_by_membership_id=excluded.created_by_membership_id,
                        integrity_hash=excluded.integrity_hash,
                        sandbox_id=excluded.sandbox_id,
                        source_version=excluded.source_version,
                        projection_state='current',
                        projected_at=excluded.projected_at,
                        stale_at=NULL, lease_expires_at=NULL
                    """,
                    (
                        document_version_id,
                        scope_id,
                        knowledge_document_id,
                        content_version,
                        content_hash,
                        finished,
                        document_object["manifestId"],
                        source_version,
                        context.membership_id,
                        context.sandbox_id,
                        hashlib.sha256(
                            f"{content_hash}|{source_asset_id}|{source_version}".encode(
                                "utf-8"
                            )
                        ).hexdigest(),
                        context.sandbox_id,
                        source_version,
                        finished,
                    ),
                )
                connection.execute(
                    """
                    INSERT INTO source_sets (
                        id, scope_id, client_id, security_label_set_version,
                        source_count, version, purpose_kind,
                        publication_state, created_by_principal_id, created_at,
                        expires_at, lifecycle_state, updated_at, deleted_at,
                        authority_role, origin_instance_id
                    ) VALUES (?, ?, ?, 'local-self-v1', 1, 1,
                              'local_wiki_build', 'draft', ?, ?, NULL, 'active',
                              ?, NULL, 'local', ?)
                    ON CONFLICT(id) DO UPDATE SET
                        source_count=1, lifecycle_state='active',
                        updated_at=excluded.updated_at, deleted_at=NULL
                    """,
                    (
                        source_set_id,
                        scope_id,
                        project_id,
                        context.principal_id,
                        finished,
                        finished,
                        context.sandbox_id,
                    ),
                )
                connection.execute(
                    """
                    INSERT INTO source_set_members (
                        id, scope_id, source_set_id, source_object_id,
                        source_version, policy_version, source_object_kind,
                        ordinal, added_at, removed_at, version,
                        lifecycle_state, created_at, updated_at, deleted_at,
                        authority_role, origin_instance_id
                    ) VALUES (?, ?, ?, ?, ?, 1, 'source_asset', 0, ?, NULL, 1,
                              'active', ?, ?, NULL, 'local', ?)
                    ON CONFLICT(id) DO UPDATE SET
                        source_version=excluded.source_version,
                        removed_at=NULL, lifecycle_state='active',
                        updated_at=excluded.updated_at, deleted_at=NULL
                    """,
                    (
                        source_set_member_id,
                        scope_id,
                        source_set_id,
                        source_asset_id,
                        source_version,
                        finished,
                        finished,
                        finished,
                        context.sandbox_id,
                    ),
                )
                active_chunk_ids = [str(item["id"]) for item in chunk_objects]
                for item in chunk_objects:
                    connection.execute(
                        """
                        INSERT INTO content_chunks (
                            id, scope_id, document_version_id, ordinal,
                            policy_version, chunk_hash, object_manifest_id,
                            start_locator, end_locator, embedding_eligibility,
                            created_at, version, lifecycle_state, updated_at,
                            deleted_at, authority_role, origin_instance_id
                        ) VALUES (?, ?, ?, ?, 1, ?, ?, ?, ?, 'eligible', ?, 1,
                                  'active', ?, NULL, 'local', ?)
                        ON CONFLICT(id) DO UPDATE SET
                            chunk_hash=excluded.chunk_hash,
                            object_manifest_id=excluded.object_manifest_id,
                            start_locator=excluded.start_locator,
                            end_locator=excluded.end_locator,
                            embedding_eligibility='eligible',
                            lifecycle_state='active', updated_at=excluded.updated_at,
                            deleted_at=NULL
                        """,
                        (
                            item["id"],
                            scope_id,
                            document_version_id,
                            int(item["ordinal"]),
                            item["hash"],
                            item["manifestId"],
                            f"char:{item['start']}",
                            f"char:{item['end']}",
                            finished,
                            finished,
                            context.sandbox_id,
                        ),
                    )
                if active_chunk_ids:
                    placeholders = ",".join("?" for _ in active_chunk_ids)
                    connection.execute(
                        f"""
                        UPDATE content_chunks
                        SET lifecycle_state='deleted', deleted_at=?,
                            updated_at=?, version=version+1
                        WHERE scope_id=? AND document_version_id=?
                          AND lifecycle_state='active'
                          AND id NOT IN ({placeholders})
                        """,
                        (
                            finished,
                            finished,
                            scope_id,
                            document_version_id,
                            *active_chunk_ids,
                        ),
                    )
                connection.execute(
                    """
                    INSERT INTO derivation_lineage (
                        id, scope_id, source_set_id, policy_version_id,
                        grant_generation, derivative_kind,
                        derivative_object_id, generator_version, generated_at,
                        invalidated_at, source_version, authority_role,
                        origin_instance_id
                    ) VALUES (?, ?, ?, NULL, 1, 'local_wiki_version', ?, ?, ?,
                              NULL, ?, 'local', ?)
                    ON CONFLICT(id) DO UPDATE SET
                        source_set_id=excluded.source_set_id,
                        derivative_object_id=excluded.derivative_object_id,
                        generator_version=excluded.generator_version,
                        generated_at=excluded.generated_at,
                        invalidated_at=NULL,
                        source_version=excluded.source_version
                    """,
                    (
                        lineage_id,
                        scope_id,
                        source_set_id,
                        document_version_id,
                        self.WIKI_GENERATOR_VERSION,
                        finished,
                        source_version,
                        context.sandbox_id,
                    ),
                )
                connection.execute(
                    """
                    INSERT INTO search_index_manifests (
                        id, scope_id, lineage_id, index_version, status,
                        reconciled_at, index_kind, index_artifact_ref,
                        generator_version, invalidated_at, source_version,
                        generated_at, authority_role, origin_instance_id
                    ) VALUES (?, ?, ?, 1, 'queued', NULL,
                              'local_keyword_chunks', NULL, ?, NULL, ?, ?,
                              'local', ?)
                    ON CONFLICT(id) DO UPDATE SET
                        lineage_id=excluded.lineage_id, status='queued',
                        reconciled_at=NULL, index_artifact_ref=NULL,
                        generator_version=excluded.generator_version,
                        invalidated_at=NULL,
                        source_version=excluded.source_version,
                        generated_at=excluded.generated_at
                    """,
                    (
                        search_manifest_id,
                        scope_id,
                        lineage_id,
                        self.WIKI_GENERATOR_VERSION,
                        source_version,
                        finished,
                        context.sandbox_id,
                    ),
                )
                connection.execute(
                    """
                    INSERT INTO vector_index_manifests (
                        id, scope_id, lineage_id, provider_resource_id,
                        policy_version, status, embedding_model,
                        embedding_dimensions, index_artifact_ref,
                        generator_version, reconciled_at, invalidated_at,
                        source_version, generated_at, authority_role,
                        origin_instance_id
                    ) VALUES (?, ?, ?, NULL, 1, 'not_connected', NULL, NULL,
                              NULL, ?, NULL, NULL, ?, ?, 'local', ?)
                    ON CONFLICT(id) DO UPDATE SET
                        lineage_id=excluded.lineage_id,
                        status='not_connected', embedding_model=NULL,
                        embedding_dimensions=NULL, index_artifact_ref=NULL,
                        generator_version=excluded.generator_version,
                        reconciled_at=NULL, invalidated_at=NULL,
                        source_version=excluded.source_version,
                        generated_at=excluded.generated_at
                    """,
                    (
                        vector_manifest_id,
                        scope_id,
                        lineage_id,
                        self.WIKI_GENERATOR_VERSION,
                        source_version,
                        finished,
                        context.sandbox_id,
                    ),
                )
                connection.execute(
                    """
                    UPDATE processing_attempts
                    SET status='ready', error_code=NULL,
                        error_message_safe=NULL, next_retry_at=NULL,
                        finished_at=?
                    WHERE id=? AND authority_role='local'
                    """,
                    (finished, attempt_id),
                )
                connection.commit()
            return {
                "documentId": document_id,
                "knowledgeDocumentId": knowledge_document_id,
                "documentVersionId": document_version_id,
                "contentVersion": content_version,
                "chunkCount": len(chunk_objects),
                "maxChunkCharacters": max(
                    len(str(item["content"])) for item in chunk_objects
                ),
                "searchIndexStatus": "queued",
                "vectorIndexStatus": "not_connected",
                **self.processing_state(entry),
            }
        except Exception as exc:
            finished = utc_now()
            error_code = (
                exc.code
                if isinstance(exc, LocalRuntimeError)
                else "local_wiki_projection_exception"
            )
            safe_message = (
                exc.message
                if isinstance(exc, LocalRuntimeError)
                else "本机知识构建暂时失败，可以重试"
            )
            with self.runtime._connection() as connection:
                connection.execute(
                    """
                    UPDATE processing_attempts
                    SET status='failed_retryable', error_code=?,
                        error_message_safe=?, next_retry_at=?, finished_at=?
                    WHERE id=? AND authority_role='local'
                    """,
                    (
                        error_code,
                        safe_message,
                        finished,
                        finished,
                        attempt_id,
                    ),
                )
                connection.commit()
            return {"documentId": document_id, **self.processing_state(entry)}

    def build_local_wiki_retrieval(
        self,
        *,
        project_id: str,
        document_id: str,
        force: bool = False,
    ) -> dict[str, Any]:
        """Build candidate facts, exact evidence and local hybrid index."""
        context = self._context()
        _, state, entry = self._document_entry(document_id)
        if str(state.get("projectId") or "") != project_id:
            raise LocalRuntimeError(
                409,
                "local_document_project_mismatch",
                "本机资料不属于当前项目",
            )
        source_asset_id = self._ensure_local_source_asset(
            project_id=project_id,
            entry=entry,
        )
        with self.runtime._connection() as connection:
            scope_id = self.runtime._local_object_scope_id(
                connection,
                context.sandbox_id,
            )
            current = connection.execute(
                """
                SELECT d.id AS knowledge_document_id, d.title,
                       d.source_asset_id, d.source_version,
                       dv.id AS document_version_id,
                       dv.version AS document_version,
                       l.id AS lineage_id, l.source_set_id,
                       s.id AS search_manifest_id, s.status AS search_status,
                       s.generator_version AS search_generator_version,
                       v.id AS vector_manifest_id, v.status AS vector_status,
                       v.generator_version AS vector_generator_version,
                       v.embedding_model, v.embedding_dimensions
                FROM knowledge_documents d
                JOIN document_versions dv
                  ON dv.scope_id=d.scope_id AND dv.document_id=d.id
                 AND dv.version=d.current_version
                JOIN derivation_lineage l
                  ON l.scope_id=d.scope_id
                 AND l.derivative_kind='local_wiki_version'
                 AND l.derivative_object_id=dv.id
                 AND l.invalidated_at IS NULL
                JOIN search_index_manifests s
                  ON s.scope_id=d.scope_id AND s.lineage_id=l.id
                 AND s.invalidated_at IS NULL
                JOIN vector_index_manifests v
                  ON v.scope_id=d.scope_id AND v.lineage_id=l.id
                 AND v.invalidated_at IS NULL
                WHERE d.scope_id=? AND d.client_id=? AND d.source_asset_id=?
                  AND d.sandbox_id=? AND d.document_kind='local_private_wiki'
                  AND d.lifecycle_state='active' AND d.parse_state='ready'
                """,
                (
                    scope_id,
                    project_id,
                    source_asset_id,
                    context.sandbox_id,
                ),
            ).fetchone()
            if current is None:
                raise LocalRuntimeError(
                    409,
                    "local_wiki_lineage_missing",
                    "本机知识版本或索引血统缺失，请重新构建知识",
                )
            existing_fact_count = int(
                connection.execute(
                    """
                    SELECT COUNT(*)
                    FROM atomic_facts f
                    JOIN content_chunks c
                      ON c.scope_id=f.scope_id AND c.id=f.chunk_id
                    WHERE f.scope_id=? AND c.document_version_id=?
                      AND f.lifecycle_state='active'
                      AND c.lifecycle_state='active'
                    """,
                    (scope_id, current["document_version_id"]),
                ).fetchone()[0]
            )
            chunks = connection.execute(
                """
                SELECT id, ordinal, object_manifest_id, start_locator,
                       end_locator, version
                FROM content_chunks
                WHERE scope_id=? AND document_version_id=?
                  AND lifecycle_state='active'
                ORDER BY ordinal, id
                """,
                (scope_id, current["document_version_id"]),
            ).fetchall()
        if (
            not force
            and existing_fact_count > 0
            and str(current["search_status"] or "") == "ready"
            and str(current["vector_status"] or "") == "ready"
            and str(current["search_generator_version"] or "")
            == self.WIKI_RETRIEVAL_GENERATOR_VERSION
            and str(current["vector_generator_version"] or "")
            == self.WIKI_RETRIEVAL_GENERATOR_VERSION
            and str(current["embedding_model"] or "")
            == self.WIKI_SPARSE_VECTOR_MODEL
            and int(current["embedding_dimensions"] or 0)
            == self.WIKI_SPARSE_VECTOR_DIMENSIONS
        ):
            return {
                "documentId": document_id,
                "knowledgeDocumentId": str(current["knowledge_document_id"]),
                "documentVersionId": str(current["document_version_id"]),
                "factCount": existing_fact_count,
                "searchIndexStatus": "ready",
                "vectorIndexStatus": "ready",
                **self.processing_state(entry),
            }
        if not chunks:
            raise LocalRuntimeError(
                409,
                "local_wiki_chunks_missing",
                "本机知识没有可检索分块，请重新构建知识",
            )

        fact_rows: list[dict[str, Any]] = []
        for chunk in chunks:
            manifest_id = str(chunk["object_manifest_id"] or "")
            content = self._read_local_manifest_text(manifest_id)
            start_locator = str(chunk["start_locator"] or "char:0")
            document_offset = (
                int(start_locator.split(":", 1)[1])
                if start_locator.startswith("char:")
                else 0
            )
            for candidate in self._fact_candidates(
                content,
                document_offset=document_offset,
            ):
                fact_hash = hashlib.sha256(
                    str(candidate["content"]).encode("utf-8")
                ).hexdigest()
                fact_id = (
                    "atomic_fact_local_"
                    + self._stable_segment(
                        f"{chunk['id']}|{candidate['chunkStart']}|"
                        f"{candidate['chunkEnd']}|{fact_hash}"
                    )
                )
                evidence_id = (
                    "evidence_local_"
                    + self._stable_segment(f"{fact_id}|{chunk['id']}")
                )
                locator = canonical_json(
                    {
                        "schema": "yiyu.local-evidence-locator.v1",
                        "chunkStart": int(candidate["chunkStart"]),
                        "chunkEnd": int(candidate["chunkEnd"]),
                        "documentStart": int(candidate["documentStart"]),
                        "documentEnd": int(candidate["documentEnd"]),
                    }
                )
                vector = self._sparse_vector(
                    self._retrieval_tokens(str(candidate["content"]))
                )
                strongest = sorted(vector.items(), key=lambda item: item[0])
                fact_rows.append(
                    {
                        **candidate,
                        "id": fact_id,
                        "evidenceId": evidence_id,
                        "hash": fact_hash,
                        "chunkId": str(chunk["id"]),
                        "chunkManifestId": manifest_id,
                        "chunkVersion": int(chunk["version"] or 1),
                        "locator": locator,
                        "locatorHash": hashlib.sha256(
                            locator.encode("utf-8")
                        ).hexdigest(),
                        "vector": [
                            [bucket, round(weight, 8)]
                            for bucket, weight in strongest
                        ],
                    }
                )

        if not fact_rows:
            raise LocalRuntimeError(
                422,
                "local_wiki_facts_empty",
                "本机知识分块没有形成可检索事实",
            )
        index_payload = {
            "schema": "yiyu.local-wiki-hybrid-index.v1",
            "clientId": project_id,
            "knowledgeDocumentId": str(current["knowledge_document_id"]),
            "documentVersionId": str(current["document_version_id"]),
            "sourceAssetId": source_asset_id,
            "sourceVersion": int(current["source_version"] or 1),
            "generatorVersion": self.WIKI_RETRIEVAL_GENERATOR_VERSION,
            "sparseVectorModel": self.WIKI_SPARSE_VECTOR_MODEL,
            "sparseVectorDimensions": self.WIKI_SPARSE_VECTOR_DIMENSIONS,
            "entries": [
                {
                    "factId": row["id"],
                    "chunkId": row["chunkId"],
                    "chunkStart": row["chunkStart"],
                    "chunkEnd": row["chunkEnd"],
                    "vector": row["vector"],
                }
                for row in fact_rows
            ],
        }
        base = (
            "local-wiki/"
            f"{self._stable_segment(context.sandbox_id)}/"
            f"{self._stable_segment(project_id)}"
        )
        index_object = self._ensure_wiki_object(
            object_id=f"wiki-hybrid-index:{current['document_version_id']}",
            storage_key=(
                f"{base}/indexes/{current['document_version_id']}.json"
            ),
            media_type=self.WIKI_SEARCH_INDEX_MEDIA_TYPE,
            content=canonical_json(index_payload),
        )
        finished = utc_now()
        fact_ids = [str(row["id"]) for row in fact_rows]
        with self.runtime._connection() as connection:
            connection.execute("BEGIN IMMEDIATE")
            scope_id = self.runtime._local_object_scope_id(
                connection,
                context.sandbox_id,
            )
            fact_placeholders = ",".join("?" for _ in fact_ids)
            connection.execute(
                f"""
                UPDATE atomic_facts
                SET lifecycle_state='deleted', deleted_at=?, updated_at=?,
                    version=version+1
                WHERE scope_id=? AND lifecycle_state='active'
                  AND chunk_id IN (
                      SELECT id FROM content_chunks
                      WHERE scope_id=? AND document_version_id=?
                  )
                  AND id NOT IN ({fact_placeholders})
                """,
                (
                    finished,
                    finished,
                    scope_id,
                    scope_id,
                    current["document_version_id"],
                    *fact_ids,
                ),
            )
            for row in fact_rows:
                connection.execute(
                    """
                    INSERT INTO secured_resources (
                        id, scope_id, resource_kind, lifecycle_state, version,
                        resource_type_key, created_at, updated_at, deleted_at,
                        authority_role, origin_instance_id
                    ) VALUES (?, ?, 'atomic_fact', 'active', 1,
                              'local_candidate_fact', ?, ?, NULL, 'local', ?)
                    ON CONFLICT(id) DO UPDATE SET
                        lifecycle_state='active', updated_at=excluded.updated_at,
                        deleted_at=NULL, authority_role='local',
                        origin_instance_id=excluded.origin_instance_id
                    """,
                    (
                        row["id"],
                        scope_id,
                        finished,
                        finished,
                        context.sandbox_id,
                    ),
                )
                connection.execute(
                    """
                    INSERT INTO atomic_facts (
                        id, scope_id, chunk_id, fact_hash, confidence, version,
                        source_set_id, fact_object_manifest_id,
                        verification_state, confirmed_by_membership_id,
                        confirmed_at, lifecycle_state, created_at, updated_at,
                        deleted_at, authority_role, origin_instance_id
                    ) VALUES (?, ?, ?, ?, NULL, 1, ?, ?, 'candidate', NULL,
                              NULL, 'active', ?, ?, NULL, 'local', ?)
                    ON CONFLICT(id) DO UPDATE SET
                        chunk_id=excluded.chunk_id,
                        fact_hash=excluded.fact_hash,
                        source_set_id=excluded.source_set_id,
                        fact_object_manifest_id=excluded.fact_object_manifest_id,
                        verification_state='candidate',
                        lifecycle_state='active', updated_at=excluded.updated_at,
                        deleted_at=NULL, authority_role='local',
                        origin_instance_id=excluded.origin_instance_id
                    """,
                    (
                        row["id"],
                        scope_id,
                        row["chunkId"],
                        row["hash"],
                        current["source_set_id"],
                        row["chunkManifestId"],
                        finished,
                        finished,
                        context.sandbox_id,
                    ),
                )
                connection.execute(
                    """
                    INSERT INTO evidence_links (
                        id, scope_id, fact_id, source_object_id,
                        source_version, locator, source_object_kind,
                        locator_kind, page_no, paragraph_no, locator_hash,
                        created_at
                    ) VALUES (?, ?, ?, ?, ?, ?, 'content_chunk',
                              'local_character_range', NULL, NULL, ?, ?)
                    ON CONFLICT(id) DO UPDATE SET
                        fact_id=excluded.fact_id,
                        source_object_id=excluded.source_object_id,
                        source_version=excluded.source_version,
                        locator=excluded.locator,
                        source_object_kind='content_chunk',
                        locator_kind='local_character_range',
                        locator_hash=excluded.locator_hash
                    """,
                    (
                        row["evidenceId"],
                        scope_id,
                        row["id"],
                        row["chunkId"],
                        row["chunkVersion"],
                        row["locator"],
                        row["locatorHash"],
                        finished,
                    ),
                )
            connection.execute(
                """
                UPDATE search_index_manifests
                SET status='ready', index_version=1,
                    index_kind='local_keyword_chunks',
                    index_artifact_ref=?, generator_version=?,
                    reconciled_at=?, generated_at=?, invalidated_at=NULL
                WHERE id=? AND scope_id=?
                """,
                (
                    index_object["manifestId"],
                    self.WIKI_RETRIEVAL_GENERATOR_VERSION,
                    finished,
                    finished,
                    current["search_manifest_id"],
                    scope_id,
                ),
            )
            connection.execute(
                """
                UPDATE vector_index_manifests
                SET status='ready', provider_resource_id=NULL,
                    embedding_model=?, embedding_dimensions=?,
                    index_artifact_ref=?, generator_version=?,
                    reconciled_at=?, generated_at=?, invalidated_at=NULL
                WHERE id=? AND scope_id=?
                """,
                (
                    self.WIKI_SPARSE_VECTOR_MODEL,
                    self.WIKI_SPARSE_VECTOR_DIMENSIONS,
                    index_object["manifestId"],
                    self.WIKI_RETRIEVAL_GENERATOR_VERSION,
                    finished,
                    finished,
                    current["vector_manifest_id"],
                    scope_id,
                ),
            )
            connection.commit()
        return {
            "documentId": document_id,
            "knowledgeDocumentId": str(current["knowledge_document_id"]),
            "documentVersionId": str(current["document_version_id"]),
            "factCount": len(fact_ids),
            "evidenceCount": len(fact_ids),
            "searchIndexStatus": "ready",
            "vectorIndexStatus": "ready",
            "semanticMode": "local_sparse_vector",
            **self.processing_state(entry),
        }

    def process_pending_documents(
        self,
        *,
        project_id: str,
        document_ids: Iterable[Any] = (),
        force: bool = False,
    ) -> dict[str, Any]:
        state = self._load_project_state(project_id)
        requested = {str(value) for value in document_ids if str(value)}
        selected = [
            str(document_id)
            for document_id in dict(state.get("documents") or {})
            if not requested or str(document_id) in requested
        ]
        parsed_items = [
            self.process_document(
                project_id=project_id,
                document_id=document_id,
                force=force,
            )
            for document_id in selected
        ]
        wiki_items = [
            self.build_local_wiki_document(
                project_id=project_id,
                document_id=document_id,
                force=force,
            )
            if str(parsed.get("parseStatus") or "") == "ready"
            else parsed
            for document_id, parsed in zip(selected, parsed_items, strict=True)
        ]
        items = [
            self.build_local_wiki_retrieval(
                project_id=project_id,
                document_id=document_id,
                force=force,
            )
            if str(wiki.get("wikiStatus") or "") == "ready"
            else wiki
            for document_id, wiki in zip(selected, wiki_items, strict=True)
        ]
        counts = {
            status: sum(item.get("parseStatus") == status for item in items)
            for status in (
                "ready",
                "processing",
                "queued",
                "blocked",
                "failed_retryable",
            )
        }
        return {
            "clientId": project_id,
            "attempted": len(items),
            "items": items,
            "counts": counts,
            "wikiCounts": {
                status: sum(item.get("wikiStatus") == status for item in items)
                for status in (
                    "ready",
                    "processing",
                    "queued",
                    "blocked",
                    "failed_retryable",
                )
            },
            "state": (
                "failed_retryable"
                if counts["failed_retryable"]
                else "blocked"
                if counts["blocked"] and not counts["ready"]
                else "ready"
            ),
            "retryable": bool(counts["failed_retryable"]),
            "updatedAt": utc_now(),
        }

    def retry_document_processing(
        self,
        *,
        project_id: str,
        document_id: str,
    ) -> dict[str, Any]:
        parsed = self.process_document(
            project_id=project_id,
            document_id=document_id,
            force=True,
        )
        if str(parsed.get("parseStatus") or "") != "ready":
            return parsed
        wiki = self.build_local_wiki_document(
            project_id=project_id,
            document_id=document_id,
            force=True,
        )
        if str(wiki.get("wikiStatus") or "") != "ready":
            return wiki
        return self.build_local_wiki_retrieval(
            project_id=project_id,
            document_id=document_id,
            force=True,
        )

    def local_wiki_status(self, project_id: str) -> dict[str, Any]:
        context = self._context()
        with self.runtime._connection() as connection:
            scope_id = self.runtime._local_object_scope_id(
                connection,
                context.sandbox_id,
            )
            row = connection.execute(
                """
                SELECT
                    COUNT(DISTINCT d.id) AS document_count,
                    COUNT(c.id) AS chunk_count,
                    COUNT(DISTINCT CASE WHEN s.status='ready' THEN s.id END)
                        AS search_ready_count,
                    COUNT(DISTINCT CASE WHEN s.status='queued' THEN s.id END)
                        AS search_queued_count,
                    COUNT(DISTINCT CASE WHEN v.status='ready' THEN v.id END)
                        AS vector_ready_count,
                    MAX(d.updated_at) AS updated_at
                FROM knowledge_documents d
                LEFT JOIN document_versions dv
                  ON dv.scope_id=d.scope_id AND dv.document_id=d.id
                 AND dv.version=d.current_version
                LEFT JOIN content_chunks c
                  ON c.scope_id=d.scope_id AND c.document_version_id=dv.id
                 AND c.lifecycle_state='active'
                LEFT JOIN derivation_lineage l
                  ON l.scope_id=d.scope_id
                 AND l.derivative_kind='local_wiki_version'
                 AND l.derivative_object_id=dv.id
                 AND l.invalidated_at IS NULL
                LEFT JOIN search_index_manifests s
                  ON s.scope_id=d.scope_id AND s.lineage_id=l.id
                 AND s.invalidated_at IS NULL
                LEFT JOIN vector_index_manifests v
                  ON v.scope_id=d.scope_id AND v.lineage_id=l.id
                 AND v.invalidated_at IS NULL
                WHERE d.scope_id=? AND d.client_id=? AND d.sandbox_id=?
                  AND d.document_kind='local_private_wiki'
                  AND d.lifecycle_state='active' AND d.parse_state='ready'
                """,
                (scope_id, project_id, context.sandbox_id),
            ).fetchone()
            evidence_row = connection.execute(
                """
                SELECT COUNT(DISTINCT f.id) AS fact_count,
                       COUNT(DISTINCT e.id) AS evidence_count
                FROM knowledge_documents d
                JOIN document_versions dv
                  ON dv.scope_id=d.scope_id AND dv.document_id=d.id
                 AND dv.version=d.current_version
                JOIN content_chunks c
                  ON c.scope_id=d.scope_id AND c.document_version_id=dv.id
                 AND c.lifecycle_state='active'
                JOIN atomic_facts f
                  ON f.scope_id=d.scope_id AND f.chunk_id=c.id
                 AND f.lifecycle_state='active'
                LEFT JOIN evidence_links e
                  ON e.scope_id=d.scope_id AND e.fact_id=f.id
                 AND e.source_object_id=c.id
                WHERE d.scope_id=? AND d.client_id=? AND d.sandbox_id=?
                  AND d.document_kind='local_private_wiki'
                  AND d.lifecycle_state='active' AND d.parse_state='ready'
                """,
                (scope_id, project_id, context.sandbox_id),
            ).fetchone()
        return {
            "documentCount": int(row["document_count"] or 0),
            "chunkCount": int(row["chunk_count"] or 0),
            "factCount": int(evidence_row["fact_count"] or 0),
            "evidenceCount": int(evidence_row["evidence_count"] or 0),
            "searchReadyCount": int(row["search_ready_count"] or 0),
            "searchQueuedCount": int(row["search_queued_count"] or 0),
            "vectorReadyCount": int(row["vector_ready_count"] or 0),
            "updatedAt": row["updated_at"],
        }

    def ensure_local_wiki_retrieval(self, project_id: str) -> dict[str, Any]:
        context = self._context()
        with self.runtime._connection() as connection:
            scope_id = self.runtime._local_object_scope_id(
                connection,
                context.sandbox_id,
            )
            pending_source_ids = {
                str(row["source_asset_id"])
                for row in connection.execute(
                    """
                    SELECT d.source_asset_id
                    FROM knowledge_documents d
                    JOIN document_versions dv
                      ON dv.scope_id=d.scope_id AND dv.document_id=d.id
                     AND dv.version=d.current_version
                    JOIN derivation_lineage l
                      ON l.scope_id=d.scope_id
                     AND l.derivative_kind='local_wiki_version'
                     AND l.derivative_object_id=dv.id
                     AND l.invalidated_at IS NULL
                    JOIN search_index_manifests s
                      ON s.scope_id=d.scope_id AND s.lineage_id=l.id
                     AND s.invalidated_at IS NULL
                    JOIN vector_index_manifests v
                      ON v.scope_id=d.scope_id AND v.lineage_id=l.id
                     AND v.invalidated_at IS NULL
                    WHERE d.scope_id=? AND d.client_id=? AND d.sandbox_id=?
                      AND d.lifecycle_state='active'
                      AND (
                          s.status!='ready' OR v.status!='ready'
                          OR COALESCE(s.generator_version, '')!=?
                          OR COALESCE(v.generator_version, '')!=?
                          OR COALESCE(v.embedding_model, '')!=?
                          OR COALESCE(v.embedding_dimensions, 0)!=?
                      )
                    """,
                    (
                        scope_id,
                        project_id,
                        context.sandbox_id,
                        self.WIKI_RETRIEVAL_GENERATOR_VERSION,
                        self.WIKI_RETRIEVAL_GENERATOR_VERSION,
                        self.WIKI_SPARSE_VECTOR_MODEL,
                        self.WIKI_SPARSE_VECTOR_DIMENSIONS,
                    ),
                ).fetchall()
            }
        if not pending_source_ids:
            status = self.local_wiki_status(project_id)
            return {
                "clientId": project_id,
                "attempted": 0,
                "ready": int(status.get("searchReadyCount") or 0),
                "items": [],
            }
        state = self._load_project_state(project_id)
        items: list[dict[str, Any]] = []
        for document_id, raw in dict(state.get("documents") or {}).items():
            if not isinstance(raw, Mapping):
                continue
            source_asset_id = str(
                raw.get("cloudDocumentId")
                or raw.get("documentId")
                or document_id
            )
            if source_asset_id.startswith("local-pending:"):
                source_asset_id = str(raw.get("localSourceId") or "")
            if source_asset_id not in pending_source_ids:
                continue
            status = self.processing_state(raw)
            if (
                str(status.get("parseStatus") or "") == "ready"
                and str(status.get("wikiStatus") or "") == "ready"
            ):
                items.append(
                    self.build_local_wiki_retrieval(
                        project_id=project_id,
                        document_id=str(document_id),
                        force=False,
                    )
                )
        return {
            "clientId": project_id,
            "attempted": len(items),
            "ready": sum(
                item.get("searchIndexStatus") == "ready" for item in items
            ),
            "items": items,
        }

    def search_local_wiki(
        self,
        *,
        project_id: str,
        query: str,
        limit: int | None = 20,
    ) -> dict[str, Any]:
        normalized_query = str(query or "").strip()
        if not normalized_query:
            raise LocalRuntimeError(422, "knowledge_search_query_required", "请输入搜索内容")
        self.ensure_local_wiki_retrieval(project_id)
        context = self._context()
        project_state = self._load_project_state(project_id)
        documents = {
            str(document_id): dict(item)
            for document_id, item in dict(
                project_state.get("documents") or {}
            ).items()
            if isinstance(item, Mapping)
        }
        with self.runtime._connection() as connection:
            scope_id = self.runtime._local_object_scope_id(
                connection,
                context.sandbox_id,
            )
            rows = connection.execute(
                """
                SELECT f.id AS fact_id, f.fact_hash,
                       e.id AS evidence_id, e.locator, e.locator_hash,
                       c.id AS chunk_id, c.ordinal, c.object_manifest_id,
                       d.id AS knowledge_document_id, d.title,
                       d.source_asset_id, d.source_version,
                       dv.id AS document_version_id,
                       s.index_artifact_ref, s.status AS search_status,
                       v.status AS vector_status, v.embedding_model
                FROM knowledge_documents d
                JOIN document_versions dv
                  ON dv.scope_id=d.scope_id AND dv.document_id=d.id
                 AND dv.version=d.current_version
                JOIN content_chunks c
                  ON c.scope_id=d.scope_id AND c.document_version_id=dv.id
                 AND c.lifecycle_state='active'
                JOIN atomic_facts f
                  ON f.scope_id=d.scope_id AND f.chunk_id=c.id
                 AND f.lifecycle_state='active'
                JOIN evidence_links e
                  ON e.scope_id=d.scope_id AND e.fact_id=f.id
                 AND e.source_object_id=c.id
                JOIN derivation_lineage l
                  ON l.scope_id=d.scope_id
                 AND l.derivative_kind='local_wiki_version'
                 AND l.derivative_object_id=dv.id
                 AND l.invalidated_at IS NULL
                JOIN search_index_manifests s
                  ON s.scope_id=d.scope_id AND s.lineage_id=l.id
                 AND s.invalidated_at IS NULL
                JOIN vector_index_manifests v
                  ON v.scope_id=d.scope_id AND v.lineage_id=l.id
                 AND v.invalidated_at IS NULL
                WHERE d.scope_id=? AND d.client_id=? AND d.sandbox_id=?
                  AND d.document_kind='local_private_wiki'
                  AND d.lifecycle_state='active' AND d.parse_state='ready'
                  AND s.status='ready' AND v.status='ready'
                ORDER BY d.title, c.ordinal, f.id
                """,
                (scope_id, project_id, context.sandbox_id),
            ).fetchall()

        query_tokens = self._retrieval_tokens(normalized_query)
        query_vector = self._sparse_vector(query_tokens)
        visible_terms = [
            match.group(0).casefold()
            for match in re.finditer(
                r"[A-Za-z0-9_]{2,}|[\u4e00-\u9fff]{2,24}",
                normalized_query,
            )
        ]
        chunk_cache: dict[str, str] = {}
        prepared: list[dict[str, Any]] = []
        keyword_candidates: list[dict[str, Any]] = []
        for row in rows:
            manifest_id = str(row["object_manifest_id"] or "")
            if manifest_id not in chunk_cache:
                chunk_cache[manifest_id] = self._read_local_manifest_text(
                    manifest_id
                )
            chunk_text = chunk_cache[manifest_id]
            try:
                locator = json.loads(str(row["locator"] or "{}"))
            except json.JSONDecodeError:
                continue
            start = int(locator.get("chunkStart") or 0)
            end = int(locator.get("chunkEnd") or 0)
            fact_text = chunk_text[start:end].strip()
            if not fact_text:
                continue
            lowered = fact_text.casefold()
            matched_terms = [term for term in visible_terms if term in lowered]
            exact_phrase = normalized_query.casefold() in lowered
            keyword_score = (
                (3.0 if exact_phrase else 0.0)
                + len(matched_terms) / max(1, len(visible_terms))
            )
            item = {
                "row": row,
                "text": fact_text,
                "matchedTerms": matched_terms,
                "keywordScore": keyword_score,
                "semanticScore": 0.0,
                "exactPhrase": exact_phrase,
            }
            prepared.append(item)
            if keyword_score > 0:
                keyword_candidates.append(item)

        candidates = keyword_candidates
        if not candidates:
            artifact_vectors: dict[str, dict[int, float]] = {}
            loaded_artifacts: set[str] = set()
            for row in rows:
                artifact_id = str(row["index_artifact_ref"] or "")
                if not artifact_id or artifact_id in loaded_artifacts:
                    continue
                loaded_artifacts.add(artifact_id)
                try:
                    payload = json.loads(
                        self._read_local_manifest_text(artifact_id)
                    )
                except (LocalRuntimeError, json.JSONDecodeError):
                    continue
                if payload.get("schema") != "yiyu.local-wiki-hybrid-index.v1":
                    continue
                for entry in payload.get("entries") or []:
                    if not isinstance(entry, Mapping):
                        continue
                    vector: dict[int, float] = {}
                    for pair in entry.get("vector") or []:
                        if isinstance(pair, list) and len(pair) == 2:
                            vector[int(pair[0])] = float(pair[1])
                    artifact_vectors[str(entry.get("factId") or "")] = vector
            semantic_candidates: list[dict[str, Any]] = []
            for item in prepared:
                semantic_score = self._sparse_cosine(
                    query_vector,
                    artifact_vectors.get(str(item["row"]["fact_id"]), {}),
                )
                if semantic_score < 0.08:
                    continue
                item["semanticScore"] = semantic_score
                semantic_candidates.append(item)
            candidates = semantic_candidates

        keyword_rank = {
            str(item["row"]["fact_id"]): rank
            for rank, item in enumerate(
                sorted(
                    candidates,
                    key=lambda item: (-item["keywordScore"], item["row"]["fact_id"]),
                ),
                start=1,
            )
            if item["keywordScore"] > 0
        }
        semantic_rank = {
            str(item["row"]["fact_id"]): rank
            for rank, item in enumerate(
                sorted(
                    candidates,
                    key=lambda item: (-item["semanticScore"], item["row"]["fact_id"]),
                ),
                start=1,
            )
            if item["semanticScore"] >= 0.08
        }
        hits: list[dict[str, Any]] = []
        for item in candidates:
            row = item["row"]
            fact_id = str(row["fact_id"])
            fused = 0.0
            if fact_id in keyword_rank:
                fused += 1.0 / (60 + keyword_rank[fact_id])
            if fact_id in semantic_rank:
                fused += 1.0 / (60 + semantic_rank[fact_id])
            score = max(
                fused * 130,
                4.5 if item["exactPhrase"] else 0.0,
                3.2 if item["keywordScore"] >= 1.5 else 0.0,
                item["semanticScore"] * 3.0,
            )
            source_asset_id = str(row["source_asset_id"] or "")
            local_document = documents.get(source_asset_id) or {}
            original_path = str(
                local_document.get("originalSourcePath") or ""
            )
            managed_path = str(local_document.get("managedPath") or "")
            path = (
                original_path
                if original_path and Path(original_path).is_file()
                else managed_path
            )
            original_available = bool(path and Path(path).is_file())
            hits.append(
                {
                    "title": str(row["title"] or local_document.get("title") or "本机资料"),
                    "excerpt": item["text"],
                    "score": round(score, 6),
                    "keywordScore": round(float(item["keywordScore"]), 6),
                    "semanticScore": round(float(item["semanticScore"]), 6),
                    "stage": "raw_chunk",
                    "sourceType": "local_document",
                    "documentId": source_asset_id,
                    "knowledgeDocumentId": str(row["knowledge_document_id"]),
                    "documentVersionId": str(row["document_version_id"]),
                    "chunkId": str(row["chunk_id"]),
                    "factId": fact_id,
                    "evidenceId": str(row["evidence_id"]),
                    "evidenceLocator": str(row["locator"] or ""),
                    "path": path or None,
                    "originalPath": path or None,
                    "sourceAvailability": (
                        "original_available"
                        if original_available
                        else "machine_readable_only"
                    ),
                    "originalAvailable": original_available,
                    "machineReadableAvailable": True,
                    "openableKind": (
                        "original_file" if original_available else "machine_markdown"
                    ),
                    "sectionLabel": f"正文第 {int(row['ordinal']) + 1} 段",
                    "matchedTerms": item["matchedTerms"][:8],
                    "retrievalMode": (
                        "keyword_sparse_hybrid"
                        if fact_id in keyword_rank and fact_id in semantic_rank
                        else "keyword"
                        if fact_id in keyword_rank
                        else "local_sparse_vector"
                    ),
                    "citationRole": (
                        "direct_quote" if item["exactPhrase"] else "direct_support"
                    ),
                    "citationPriority": 100 if item["exactPhrase"] else 80,
                    "citationReason": "本机知识分块中的可追溯原文",
                }
            )
        hits.sort(
            key=lambda item: (
                -float(item.get("score") or 0),
                str(item.get("title") or ""),
                str(item.get("factId") or ""),
            )
        )
        limited_hits = (
            hits
            if limit is None
            else hits[: max(1, min(int(limit), 50))]
        )
        matched_documents = {str(item.get("documentId") or "") for item in hits}
        return {
            "searchId": new_id(),
            "clientId": project_id,
            "query": normalized_query,
            "coverage": round(len(matched_documents) / max(1, len(documents)), 3),
            "matchedTerms": sorted(
                {term for item in limited_hits for term in item["matchedTerms"]}
            ),
            "masterHitCount": 0,
            "surrogateHitCount": 0,
            "rawChunkHitCount": len(limited_hits),
            "drillthroughUsed": bool(limited_hits),
            "phase": "completed",
            "progress": 100,
            "progressFloor": 100,
            "progressCeiling": 100,
            "lastUpdatedAt": utc_now(),
            "hits": limited_hits,
            "previewSummary": "；".join(
                str(item["excerpt"])[:120] for item in limited_hits[:3]
            ),
            "strictState": "ready",
            "retrievalMode": "local_hybrid",
            "keywordIndexState": "ready",
            "semanticIndexState": "ready",
            "semanticModel": self.WIKI_SPARSE_VECTOR_MODEL,
        }

    def strategic_profile_corpus(self, project_id: str) -> dict[str, Any]:
        """Return a bounded, per-document evidence inventory for profile synthesis.

        Search answers should rank only relevant chunks.  A client profile has a
        different duty: it must first inspect every parse-ready project document,
        then cite only the documents that support each rendered section.
        """

        self.ensure_local_wiki_retrieval(project_id)
        context = self._context()
        project_state = self._load_project_state(project_id)
        state_documents = {
            str(document_id): dict(item)
            for document_id, item in dict(project_state.get("documents") or {}).items()
            if isinstance(item, Mapping)
        }
        with self.runtime._connection() as connection:
            scope_id = self.runtime._local_object_scope_id(connection, context.sandbox_id)
            rows = connection.execute(
                """
                SELECT d.id AS knowledge_document_id, d.title, d.source_asset_id,
                       d.source_version, dv.id AS document_version_id,
                       c.id AS chunk_id, c.ordinal, c.object_manifest_id,
                       e.locator
                FROM knowledge_documents AS d
                JOIN document_versions AS dv
                  ON dv.scope_id=d.scope_id AND dv.document_id=d.id
                 AND dv.version=d.current_version
                JOIN content_chunks AS c
                  ON c.scope_id=d.scope_id AND c.document_version_id=dv.id
                 AND c.lifecycle_state='active'
                JOIN atomic_facts AS f
                  ON f.scope_id=d.scope_id AND f.chunk_id=c.id
                 AND f.lifecycle_state='active'
                JOIN evidence_links AS e
                  ON e.scope_id=d.scope_id AND e.fact_id=f.id
                 AND e.source_object_id=c.id
                WHERE d.scope_id=? AND d.client_id=? AND d.sandbox_id=?
                  AND d.document_kind='local_private_wiki'
                  AND d.lifecycle_state='active' AND d.parse_state='ready'
                ORDER BY d.title, d.id, c.ordinal, f.id
                """,
                (scope_id, project_id, context.sandbox_id),
            ).fetchall()

        grouped: dict[str, dict[str, Any]] = {}
        manifest_text: dict[str, str] = {}
        for row in rows:
            source_asset_id = str(row["source_asset_id"] or "").strip()
            if not source_asset_id:
                continue
            manifest_id = str(row["object_manifest_id"] or "")
            if manifest_id not in manifest_text:
                manifest_text[manifest_id] = self._read_local_manifest_text(manifest_id)
            try:
                locator = json.loads(str(row["locator"] or "{}"))
            except json.JSONDecodeError:
                continue
            start = int(locator.get("chunkStart") or 0)
            end = int(locator.get("chunkEnd") or 0)
            excerpt = manifest_text[manifest_id][start:end].strip()
            if not excerpt:
                continue
            item = grouped.setdefault(
                source_asset_id,
                {
                    "sourceObjectId": source_asset_id,
                    "sourceObjectKind": "source_asset",
                    "sourceVersion": max(1, int(row["source_version"] or 1)),
                    "knowledgeDocumentId": str(row["knowledge_document_id"] or ""),
                    "documentVersionId": str(row["document_version_id"] or ""),
                    "title": str(row["title"] or "本机资料")[:300],
                    "chunks": [],
                },
            )
            item["chunks"].append(
                {"ordinal": int(row["ordinal"] or 0), "excerpt": excerpt[:1_500]}
            )

        documents: list[dict[str, Any]] = []
        for source_asset_id, item in grouped.items():
            chunks = list(item.pop("chunks", []))
            # Cover the document's opening plus its most substantial remaining
            # section.  This keeps the prompt bounded while avoiding a first-page-only
            # profile when the key fact appears later in the file.
            selected = chunks[:1]
            if len(chunks) > 1:
                selected.append(max(chunks[1:], key=lambda chunk: len(chunk["excerpt"])))
            local = self.document_text(source_asset_id)
            state = state_documents.get(source_asset_id) or {}
            documents.append(
                {
                    **item,
                    "contentHash": str(local.get("contentHash") or ""),
                    "title": str(item.get("title") or state.get("title") or "本机资料")[:300],
                    "excerpt": "\n".join(chunk["excerpt"] for chunk in selected)[:3_000],
                }
            )
        documents.sort(key=lambda item: (str(item.get("title") or ""), str(item["sourceObjectId"])))
        return {
            "clientId": project_id,
            "eligibleDocumentCount": len(grouped),
            "scannedDocumentCount": len(documents),
            "documents": documents,
        }

    def knowledge_presentation(self, project_id: str) -> dict[str, Any]:
        """Read client-scoped memory lanes and proven relationship cards."""
        context = self._context()
        memory_kind_map = {
            "favorite_memory": "favorite",
            "answer_favorite": "favorite",
            "correction_memory": "correction",
            "answer_correction": "correction",
            "user_correction": "correction",
            "user_supplement": "correction",
            "system_inference": "system_inference",
            "inferred_memory": "system_inference",
        }
        with self.runtime._connection() as connection:
            scope_id = self.runtime._local_object_scope_id(
                connection,
                context.sandbox_id,
            )
            local_document_count = int(
                connection.execute(
                    """
                    SELECT COUNT(*)
                    FROM knowledge_documents
                    WHERE scope_id=? AND client_id=? AND sandbox_id=?
                      AND document_kind='local_private_wiki'
                      AND lifecycle_state='active' AND parse_state='ready'
                    """,
                    (scope_id, project_id, context.sandbox_id),
                ).fetchone()[0]
            )
            memory_rows = connection.execute(
                """
                SELECT d.id, d.title, d.document_kind, d.current_version,
                       d.publication_state, d.updated_at, d.version,
                       v.id AS document_version_id, v.content_hash,
                       v.object_manifest_id,
                       (
                         SELECT sm.source_object_id
                         FROM derivation_lineage AS dl
                         JOIN source_set_members AS sm
                           ON sm.scope_id=dl.scope_id
                          AND sm.source_set_id=dl.source_set_id
                          AND sm.lifecycle_state='active'
                         WHERE dl.scope_id=d.scope_id
                           AND dl.derivative_object_id=d.id
                           AND dl.invalidated_at IS NULL
                           AND sm.source_object_kind='ai_answer'
                         ORDER BY sm.ordinal, sm.id
                         LIMIT 1
                       ) AS source_answer_id
                FROM knowledge_documents AS d
                JOIN document_versions AS v
                  ON v.scope_id=d.scope_id AND v.document_id=d.id
                 AND v.version=d.current_version
                WHERE d.scope_id=? AND d.client_id=? AND d.sandbox_id=?
                  AND d.lifecycle_state='active'
                  AND d.document_kind IN (
                    'favorite_memory', 'answer_favorite',
                    'correction_memory', 'answer_correction',
                    'user_correction', 'user_supplement',
                    'system_inference', 'inferred_memory'
                  )
                ORDER BY d.updated_at DESC, d.id
                """,
                (scope_id, project_id, context.sandbox_id),
            ).fetchall()
            correction_rows = connection.execute(
                """
                SELECT fact.id, fact.version, fact.updated_at,
                       fact.fact_hash, fact.fact_object_manifest_id,
                       fact.confirmed_by_membership_id, sources.purpose_kind,
                       (
                         SELECT member.source_object_id
                         FROM source_set_members AS member
                         WHERE member.scope_id=fact.scope_id
                           AND member.source_set_id=fact.source_set_id
                           AND member.source_object_kind='ai_answer'
                           AND member.lifecycle_state='active'
                         ORDER BY member.ordinal, member.id
                         LIMIT 1
                       ) AS source_answer_id
                FROM atomic_facts AS fact
                JOIN source_sets AS sources
                  ON sources.scope_id=fact.scope_id
                 AND sources.id=fact.source_set_id
                 AND sources.client_id=?
                 AND sources.purpose_kind IN ('answer_correction', 'answer_remember')
                 AND sources.lifecycle_state='active'
                WHERE fact.scope_id=? AND fact.lifecycle_state='active'
                  AND fact.verification_state='verified'
                ORDER BY fact.updated_at DESC, fact.id
                """,
                (project_id, scope_id),
            ).fetchall()
            relation_rows = connection.execute(
                """
                SELECT r.id, r.predicate, r.verification_state, r.confidence,
                       r.version, sd.title AS subject_title,
                       od.title AS object_title
                FROM relationship_triples AS r
                JOIN atomic_facts AS sf
                  ON sf.scope_id=r.scope_id AND sf.id=r.subject_fact_id
                 AND sf.lifecycle_state='active'
                JOIN content_chunks AS sc
                  ON sc.scope_id=sf.scope_id AND sc.id=sf.chunk_id
                 AND sc.lifecycle_state='active'
                JOIN document_versions AS sv
                  ON sv.scope_id=sc.scope_id AND sv.id=sc.document_version_id
                JOIN knowledge_documents AS sd
                  ON sd.scope_id=sv.scope_id AND sd.id=sv.document_id
                 AND sd.client_id=? AND sd.sandbox_id=?
                 AND sd.lifecycle_state='active'
                JOIN atomic_facts AS ofact
                  ON ofact.scope_id=r.scope_id AND ofact.id=r.object_fact_id
                 AND ofact.lifecycle_state='active'
                JOIN content_chunks AS oc
                  ON oc.scope_id=ofact.scope_id AND oc.id=ofact.chunk_id
                 AND oc.lifecycle_state='active'
                JOIN document_versions AS ov
                  ON ov.scope_id=oc.scope_id AND ov.id=oc.document_version_id
                JOIN knowledge_documents AS od
                  ON od.scope_id=ov.scope_id AND od.id=ov.document_id
                 AND od.client_id=? AND od.sandbox_id=?
                 AND od.lifecycle_state='active'
                WHERE r.scope_id=? AND r.lifecycle_state='active'
                ORDER BY CASE r.verification_state WHEN 'verified' THEN 0 ELSE 1 END,
                         r.updated_at DESC, r.id
                LIMIT 24
                """,
                (
                    project_id,
                    context.sandbox_id,
                    project_id,
                    context.sandbox_id,
                    scope_id,
                ),
            ).fetchall()

        memories: list[dict[str, Any]] = []
        for row in memory_rows:
            summary = ""
            manifest_id = str(row["object_manifest_id"] or "")
            if manifest_id:
                try:
                    raw_memory = self._read_local_manifest_text(manifest_id)
                    try:
                        parsed_memory = json.loads(raw_memory)
                    except (TypeError, ValueError, json.JSONDecodeError):
                        parsed_memory = None
                    summary = (
                        str(parsed_memory.get("content") or "")
                        if isinstance(parsed_memory, Mapping)
                        else raw_memory
                    )[:2_000]
                except LocalRuntimeError:
                    summary = ""
            memories.append(
                {
                    "id": str(row["id"]),
                    "documentVersionId": str(row["document_version_id"]),
                    "title": str(row["title"] or "已存记忆"),
                    "summary": summary,
                    "memoryKind": memory_kind_map.get(
                        str(row["document_kind"] or "").lower(),
                        "system_inference",
                    ),
                    "contentHash": str(row["content_hash"] or ""),
                    "publicationState": str(row["publication_state"] or "draft"),
                    "sourceAnswerId": str(row["source_answer_id"] or "") or None,
                    "version": int(row["version"] or 1),
                    "updatedAt": row["updated_at"],
                    "authority": "current_device",
                }
            )
        for row in correction_rows:
            manifest_id = str(row["fact_object_manifest_id"] or "")
            if not manifest_id:
                continue
            try:
                raw_correction = self._read_local_manifest_text(manifest_id)
                parsed_correction = json.loads(raw_correction)
            except (LocalRuntimeError, TypeError, ValueError, json.JSONDecodeError):
                continue
            if not isinstance(parsed_correction, Mapping):
                continue
            summary = str(parsed_correction.get("content") or "").strip()
            if not summary:
                continue
            correction_kind = str(
                parsed_correction.get("correctionKind") or "correction"
            )
            is_remember = correction_kind == "remember" or str(
                row["purpose_kind"] or ""
            ) == "answer_remember"
            superseded_text = (
                str(parsed_correction.get("selectedText") or "").strip()
                if correction_kind == "correction"
                else ""
            )
            memories.append(
                {
                    "id": str(row["id"]),
                    "documentVersionId": manifest_id,
                    "title": (
                        "明确记住"
                        if is_remember
                        else ("人工纠错" if correction_kind == "correction" else "人工补充")
                    ),
                    "summary": summary[:2_000],
                    "memoryKind": "explicit_memory" if is_remember else "correction",
                    "correctionKind": correction_kind,
                    # This stays in the local presentation only.  It lets the
                    # answer agent suppress an explicitly rejected wording
                    # without uploading the original answer text to cloud.
                    "supersededText": (
                        None if is_remember else (superseded_text[:2_000] or None)
                    ),
                    "contentHash": str(row["fact_hash"] or ""),
                    "publicationState": "published",
                    "sourceAnswerId": str(row["source_answer_id"] or "") or None,
                    "version": int(row["version"] or 1),
                    "updatedAt": row["updated_at"],
                    "authority": "organization_cloud",
                }
            )
        relationship_cards = [
            {
                "id": str(row["id"]),
                "subject": str(row["subject_title"] or "未命名事实"),
                "predicate": str(row["predicate"] or "相关"),
                "object": str(row["object_title"] or "未命名事实"),
                "verificationState": str(row["verification_state"] or "candidate"),
                "confidence": row["confidence"],
                "version": int(row["version"] or 1),
                "authority": "current_device",
            }
            for row in relation_rows
        ]
        return {
            "clientId": project_id,
            "localOriginalCount": local_document_count,
            "savedMemories": memories,
            "relationshipCards": relationship_cards,
            "updatedAt": utc_now(),
        }

    def pending_cloud_materials(self, project_id: str) -> list[dict[str, Any]]:
        state = self._load_project_state(project_id)
        return [
            {**dict(item), "documentId": str(document_id)}
            for document_id, item in dict(state.get("documents") or {}).items()
            if isinstance(item, Mapping)
            and str(item.get("cloudMetadataState") or "") != "ready"
        ]

    def delete_document_local(
        self,
        project_id: str,
        document_id: str,
    ) -> dict[str, Any]:
        state = self._load_project_state(project_id)
        documents = dict(state.get("documents") or {})
        entry = documents.get(document_id)
        if not isinstance(entry, Mapping):
            raise LocalRuntimeError(
                404,
                "local_document_missing",
                "当前设备没有该项目资料",
            )
        normalized = dict(entry)
        now = utc_now()
        context = self._context()
        for object_key in ("localSourceId", "localSummaryId"):
            object_id = str(normalized.get(object_key) or "")
            if object_id:
                self.runtime.local_storage_object_set_lifecycle(
                    object_id=object_id,
                    lifecycle_state="deleted",
                )
        # A file disappearing from the project-state JSON is not enough: the
        # strict 88-table Wiki projection must be tombstoned as well, otherwise
        # search and AI recall can keep consuming an apparently deleted file.
        # Preserve rows for audit/recovery, but invalidate every derivative.
        with self.runtime._connection() as connection:
            scope_id = self.runtime._local_object_scope_id(
                connection,
                context.sandbox_id,
            )
            knowledge_rows = connection.execute(
                "SELECT id FROM knowledge_documents "
                "WHERE scope_id=? AND client_id=? AND source_asset_id=? "
                "AND lifecycle_state='active'",
                (scope_id, project_id, document_id),
            ).fetchall()
            knowledge_ids = [str(row["id"]) for row in knowledge_rows]
            version_ids: list[str] = []
            chunk_ids: list[str] = []
            lineage_ids: list[str] = []
            source_set_ids: list[str] = []
            manifest_ids: list[str] = []
            if knowledge_ids:
                placeholders = ",".join("?" for _ in knowledge_ids)
                version_rows = connection.execute(
                    f"SELECT id, object_manifest_id FROM document_versions "
                    f"WHERE scope_id=? AND document_id IN ({placeholders})",
                    (scope_id, *knowledge_ids),
                ).fetchall()
                version_ids = [str(row["id"]) for row in version_rows]
                manifest_ids.extend(
                    str(row["object_manifest_id"])
                    for row in version_rows
                    if str(row["object_manifest_id"] or "")
                )
            if version_ids:
                placeholders = ",".join("?" for _ in version_ids)
                chunk_rows = connection.execute(
                    f"SELECT id, object_manifest_id FROM content_chunks "
                    f"WHERE scope_id=? AND document_version_id IN ({placeholders})",
                    (scope_id, *version_ids),
                ).fetchall()
                chunk_ids = [str(row["id"]) for row in chunk_rows]
                manifest_ids.extend(
                    str(row["object_manifest_id"])
                    for row in chunk_rows
                    if str(row["object_manifest_id"] or "")
                )
                lineage_rows = connection.execute(
                    f"SELECT id, source_set_id FROM derivation_lineage "
                    f"WHERE scope_id=? AND derivative_object_id IN ({placeholders}) "
                    "AND invalidated_at IS NULL",
                    (scope_id, *version_ids),
                ).fetchall()
                lineage_ids = [str(row["id"]) for row in lineage_rows]
                source_set_ids = sorted(
                    {
                        str(row["source_set_id"])
                        for row in lineage_rows
                        if str(row["source_set_id"] or "")
                    }
                )
            connection.execute("BEGIN IMMEDIATE")
            source_row = connection.execute(
                "SELECT object_manifest_id, version FROM source_assets "
                "WHERE id=? AND scope_id=? AND client_id=?",
                (document_id, scope_id, project_id),
            ).fetchone()
            if source_row is not None:
                source_manifest_id = str(source_row["object_manifest_id"] or "")
                if source_manifest_id:
                    manifest_ids.append(source_manifest_id)
                connection.execute(
                    "UPDATE source_assets SET lifecycle_state='deleted', "
                    "availability_state='deleted', deleted_at=?, updated_at=?, "
                    "version=version+1 WHERE id=? AND scope_id=?",
                    (now, now, document_id, scope_id),
                )
                connection.execute(
                    "UPDATE secured_resources SET lifecycle_state='deleted', "
                    "deleted_at=?, updated_at=?, version=version+1 "
                    "WHERE id=? AND scope_id=?",
                    (now, now, document_id, scope_id),
                )
            if knowledge_ids:
                placeholders = ",".join("?" for _ in knowledge_ids)
                connection.execute(
                    f"UPDATE knowledge_documents SET lifecycle_state='deleted', "
                    f"deleted_at=?, updated_at=?, version=version+1 "
                    f"WHERE scope_id=? AND id IN ({placeholders})",
                    (now, now, scope_id, *knowledge_ids),
                )
            if chunk_ids:
                placeholders = ",".join("?" for _ in chunk_ids)
                connection.execute(
                    f"UPDATE atomic_facts SET lifecycle_state='deleted', "
                    f"deleted_at=?, updated_at=?, version=version+1 "
                    f"WHERE scope_id=? AND chunk_id IN ({placeholders}) "
                    "AND lifecycle_state='active'",
                    (now, now, scope_id, *chunk_ids),
                )
                connection.execute(
                    f"UPDATE content_chunks SET lifecycle_state='deleted', "
                    f"deleted_at=?, updated_at=?, version=version+1 "
                    f"WHERE scope_id=? AND id IN ({placeholders}) "
                    "AND lifecycle_state='active'",
                    (now, now, scope_id, *chunk_ids),
                )
            if lineage_ids:
                placeholders = ",".join("?" for _ in lineage_ids)
                connection.execute(
                    f"UPDATE search_index_manifests SET status='invalidated', "
                    f"invalidated_at=? WHERE scope_id=? AND lineage_id IN ({placeholders}) "
                    "AND invalidated_at IS NULL",
                    (now, scope_id, *lineage_ids),
                )
                connection.execute(
                    f"UPDATE vector_index_manifests SET status='invalidated', "
                    f"invalidated_at=? WHERE scope_id=? AND lineage_id IN ({placeholders}) "
                    "AND invalidated_at IS NULL",
                    (now, scope_id, *lineage_ids),
                )
                connection.execute(
                    f"UPDATE derivation_lineage SET invalidated_at=? "
                    f"WHERE scope_id=? AND id IN ({placeholders}) "
                    "AND invalidated_at IS NULL",
                    (now, scope_id, *lineage_ids),
                )
            if source_set_ids:
                placeholders = ",".join("?" for _ in source_set_ids)
                connection.execute(
                    f"UPDATE source_set_members SET lifecycle_state='deleted', "
                    f"removed_at=?, deleted_at=?, updated_at=?, version=version+1 "
                    f"WHERE scope_id=? AND source_set_id IN ({placeholders}) "
                    "AND lifecycle_state='active'",
                    (now, now, now, scope_id, *source_set_ids),
                )
                connection.execute(
                    f"UPDATE source_sets SET lifecycle_state='deleted', "
                    f"deleted_at=?, updated_at=?, version=version+1 "
                    f"WHERE scope_id=? AND id IN ({placeholders}) "
                    "AND lifecycle_state='active'",
                    (now, now, scope_id, *source_set_ids),
                )
            unique_manifest_ids = sorted(set(manifest_ids))
            if unique_manifest_ids:
                placeholders = ",".join("?" for _ in unique_manifest_ids)
                connection.execute(
                    f"UPDATE object_manifests SET lifecycle_state='deleted', "
                    f"availability_state='deleted', deleted_at=? "
                    f"WHERE scope_id=? AND id IN ({placeholders})",
                    (now, scope_id, *unique_manifest_ids),
                )
            lifecycle_id = new_id()
            lifecycle_integrity = hashlib.sha256(
                canonical_json(
                    {
                        "id": lifecycle_id,
                        "scopeId": scope_id,
                        "resourceId": document_id,
                        "from": "active",
                        "to": "deleted",
                        "reason": "member_local_source_deleted",
                        "occurredAt": now,
                    }
                ).encode("utf-8")
            ).hexdigest()
            connection.execute(
                "INSERT INTO lifecycle_events ("
                "id,scope_id,operation_id,secured_resource_id,from_state,"
                "to_state,tombstone_version,actor_id,reason_code,occurred_at,"
                "origin_instance_id,created_at,integrity_hash"
                ") VALUES (?,?,NULL,?,'active','deleted',?,?,?, ?,?,?,?)",
                (
                    lifecycle_id,
                    scope_id,
                    document_id,
                    int((source_row or {"version": 0})["version"] or 0) + 1,
                    context.principal_id,
                    "member_local_source_deleted",
                    now,
                    context.sandbox_id,
                    now,
                    lifecycle_integrity,
                ),
            )
            connection.commit()
        documents.pop(document_id, None)
        cloud_document_id = str(normalized.get("cloudDocumentId") or "")
        pending_deletes = dict(state.get("pendingCloudDeletes") or {})
        if cloud_document_id:
            pending_deletes[cloud_document_id] = {
                "documentId": cloud_document_id,
                "state": "pending",
                "updatedAt": now,
            }
        state["documents"] = documents
        state["pendingCloudDeletes"] = pending_deletes
        self._write_project_state(project_id, state)
        return {
            "documentId": document_id,
            "fileName": normalized.get("fileName") or normalized.get("title"),
            "cloudDocumentId": cloud_document_id or None,
            "localState": "deleted",
            "cloudMetadataState": (
                "pending" if cloud_document_id else "not_applicable"
            ),
        }

    def pending_cloud_deletes(self, project_id: str) -> list[dict[str, Any]]:
        state = self._load_project_state(project_id)
        return [
            dict(item)
            for item in dict(state.get("pendingCloudDeletes") or {}).values()
            if isinstance(item, Mapping)
        ]

    def complete_cloud_delete(
        self,
        project_id: str,
        cloud_document_id: str,
    ) -> None:
        state = self._load_project_state(project_id)
        pending = dict(state.get("pendingCloudDeletes") or {})
        pending.pop(cloud_document_id, None)
        state["pendingCloudDeletes"] = pending
        self._write_project_state(project_id, state)

    def save_link_import_run(
        self,
        project_id: str,
        run: Mapping[str, Any],
    ) -> dict[str, Any]:
        run_id = str(run.get("runId") or "")
        if not run_id:
            raise LocalRuntimeError(
                422,
                "link_import_run_id_required",
                "链接导入回执缺少运行标识",
            )
        state = self._load_project_state(project_id)
        runs = dict(state.get("linkImportRuns") or {})
        normalized = {
            "runId": run_id,
            "clientId": project_id,
            "sourcePlatform": str(run.get("sourcePlatform") or ""),
            "sourceUrl": str(run.get("sourceUrl") or ""),
            "title": str(run.get("title") or ""),
            "status": str(run.get("status") or "failed"),
            "state": str(run.get("state") or run.get("status") or "failed"),
            "stage": str(run.get("stage") or ""),
            "progress": int(run.get("progress") or 0),
            "documentId": run.get("documentId"),
            "documentPath": run.get("documentPath"),
            "mediaCacheStatus": str(
                run.get("mediaCacheStatus") or "not_downloaded"
            ),
            "error": run.get("error"),
            "errorCode": run.get("errorCode"),
            "retryable": bool(run.get("retryable")),
            "pollingEnabled": bool(run.get("pollingEnabled")),
            "createdAt": str(run.get("createdAt") or utc_now()),
            "updatedAt": str(run.get("updatedAt") or utc_now()),
            "sourceScope": "local_private",
        }
        runs[run_id] = normalized
        if len(runs) > 100:
            retained = sorted(
                runs.values(),
                key=lambda item: (
                    str(item.get("updatedAt") or ""),
                    str(item.get("runId") or ""),
                ),
                reverse=True,
            )[:100]
            runs = {str(item["runId"]): item for item in retained}
        state["linkImportRuns"] = runs
        self._write_project_state(project_id, state)
        return normalized

    def link_import_runs(
        self,
        project_id: str,
        *,
        limit: int = 20,
        run_id: str | None = None,
    ) -> list[dict[str, Any]]:
        state = self._load_project_state(project_id)
        runs = [
            dict(item)
            for item in dict(state.get("linkImportRuns") or {}).values()
            if isinstance(item, Mapping)
        ]
        if run_id:
            runs = [
                item for item in runs if str(item.get("runId") or "") == run_id
            ]
        return sorted(
            runs,
            key=lambda item: (
                str(item.get("updatedAt") or ""),
                str(item.get("runId") or ""),
            ),
            reverse=True,
        )[: max(1, min(int(limit), 100))]

    def cancel_link_import_run(
        self,
        project_id: str,
        run_id: str,
    ) -> dict[str, Any]:
        runs = self.link_import_runs(project_id, run_id=run_id)
        if not runs:
            raise LocalRuntimeError(404, "import_run_missing", "链接导入任务不存在")
        run = runs[0]
        if run["status"] not in {"queued", "running", "processing"}:
            return run
        return self.save_link_import_run(
            project_id,
            {
                **run,
                "status": "canceled",
                "state": "cancelled",
                "stage": "canceled",
                "pollingEnabled": False,
                "updatedAt": utc_now(),
            },
        )

    def meetings(self, project_id: str) -> list[dict[str, Any]]:
        state = self._load_project_state(project_id)
        return [
            dict(item)
            for item in (state.get("meetings") or {}).values()
            if isinstance(item, Mapping)
        ]

    def meeting(self, project_id: str, meeting_id: str) -> dict[str, Any]:
        state = self._load_project_state(project_id)
        item = (state.get("meetings") or {}).get(meeting_id)
        if not isinstance(item, Mapping):
            raise LocalRuntimeError(404, "meeting_missing", "当前设备没有该会议草稿")
        return dict(item)

    def save_meeting(
        self,
        project_id: str,
        meeting: Mapping[str, Any],
    ) -> dict[str, Any]:
        state = self._load_project_state(project_id)
        meeting_id = str(meeting.get("id") or "").strip()
        if not meeting_id:
            raise LocalRuntimeError(422, "meeting_id_required", "会议 ID 缺失")
        normalized = {
            **dict(meeting),
            "id": meeting_id,
            "clientId": project_id,
            "sourceScope": "local_private",
            "updatedAt": utc_now(),
        }
        meetings = dict(state.get("meetings") or {})
        meetings[meeting_id] = normalized
        state["meetings"] = meetings
        self._write_project_state(project_id, state)
        return normalized

    def report_draft(
        self,
        report_id: str,
    ) -> dict[str, Any]:
        context = self._context()
        rows = self.runtime.local_storage_objects_by_media_type(
            media_type=self.PROJECT_STATE_MEDIA_TYPE,
        )
        for row in rows:
            if str(row["sandbox_id"]) != context.sandbox_id:
                continue
            try:
                state = json.loads(
                    self._managed_path(
                        str(row["storage_key"])
                    ).read_text(encoding="utf-8")
                )
            except (OSError, UnicodeDecodeError, json.JSONDecodeError):
                continue
            draft = (state.get("reportDrafts") or {}).get(report_id)
            if isinstance(draft, Mapping):
                return dict(draft)
        raise LocalRuntimeError(
            404,
            "report_draft_missing",
            "当前设备没有该未保存报告草稿",
        )

    def report_drafts(
        self,
        project_id: str,
        *,
        event_line_id: str | None = None,
        include_saved: bool = False,
    ) -> list[dict[str, Any]]:
        """List private report drafts for one project without crossing sandboxes."""

        state = self._load_project_state(project_id)
        expected_event_line_id = str(event_line_id or "").strip()
        items: list[dict[str, Any]] = []
        for value in (state.get("reportDrafts") or {}).values():
            if not isinstance(value, Mapping):
                continue
            item = dict(value)
            if str(item.get("client_id") or item.get("clientId") or project_id) != project_id:
                continue
            item_event_line_id = str(
                item.get("event_line_id") or item.get("eventLineId") or ""
            ).strip()
            if expected_event_line_id and item_event_line_id != expected_event_line_id:
                continue
            if not include_saved and str(item.get("status") or "") == "saved":
                continue
            items.append(item)
        return sorted(
            items,
            key=lambda item: str(item.get("updated_at") or item.get("created_at") or ""),
            reverse=True,
        )

    def document_project_id(self, document_id: str) -> str:
        """Resolve only the project identity before opening local document text."""

        project_id, _state, _entry = self._document_entry(document_id)
        if not project_id:
            raise LocalRuntimeError(
                409,
                "local_document_project_missing",
                "本机资料缺少项目归属，暂不能读取",
            )
        return project_id

    def save_report_draft(
        self,
        project_id: str,
        draft: Mapping[str, Any],
    ) -> dict[str, Any]:
        state = self._load_project_state(project_id)
        report_id = str(draft.get("id") or "").strip()
        if not report_id:
            raise LocalRuntimeError(422, "report_draft_id_required", "报告草稿 ID 缺失")
        normalized = {
            **dict(draft),
            "id": report_id,
            "client_id": project_id,
            "sourceScope": "local_private_draft",
            "updated_at": utc_now(),
        }
        drafts = dict(state.get("reportDrafts") or {})
        drafts[report_id] = normalized
        state["reportDrafts"] = drafts
        self._write_project_state(project_id, state)
        return normalized

    def _document_entry(
        self,
        document_id: str,
    ) -> tuple[str, dict[str, Any], dict[str, Any]]:
        context = self._context()
        rows = self.runtime.local_storage_objects_by_media_type(
            media_type=self.PROJECT_STATE_MEDIA_TYPE,
        )
        for row in rows:
            if str(row["sandbox_id"]) != context.sandbox_id:
                continue
            path = self._managed_path(str(row["storage_key"]))
            try:
                state = json.loads(path.read_text(encoding="utf-8"))
            except (OSError, UnicodeDecodeError, json.JSONDecodeError):
                continue
            entry = (state.get("documents") or {}).get(document_id)
            if isinstance(entry, dict):
                state["_localStorageVersion"] = int(row["version"])
                state["_localSandboxId"] = context.sandbox_id
                return str(state.get("projectId") or ""), state, entry
        raise LocalRuntimeError(
            404,
            "local_document_missing",
            "当前设备没有该资料的源文件；可查看组织共享摘要或在本机重新导入",
        )

    def _source_path(
        self,
        entry: Mapping[str, Any],
        *,
        sandbox_id: str | None = None,
    ) -> tuple[Path, dict[str, Any]]:
        sandbox_id = sandbox_id or self._context().sandbox_id
        source_id = str(entry.get("localSourceId") or "")
        row = self.runtime.local_storage_object_get(
            sandbox_id=sandbox_id,
            object_id=source_id,
        )
        if row is None or str(row["lifecycle_state"]) != "active":
            raise LocalRuntimeError(
                404,
                "local_document_source_missing",
                "当前设备的资料源文件不存在或已移除",
            )
        path = self._managed_path(str(row["storage_key"]))
        try:
            data = path.read_bytes()
        except OSError as exc:
            raise LocalRuntimeError(
                404,
                "local_document_source_missing",
                "当前设备的资料源文件无法读取",
            ) from exc
        if (
            len(data) != int(row["byte_size"])
            or hashlib.sha256(data).hexdigest() != str(row["content_hash"])
        ):
            raise LocalRuntimeError(
                409,
                "local_document_source_corrupt",
                "当前设备的资料源文件校验失败",
            )
        return path, dict(row)

    @staticmethod
    def _docx_text(path: Path) -> str:
        try:
            with zipfile.ZipFile(path) as package:
                document_xml = package.read("word/document.xml")
        except (OSError, KeyError, zipfile.BadZipFile) as exc:
            raise LocalRuntimeError(
                422,
                "local_document_format_invalid",
                "Word 文档结构无效，无法读取",
            ) from exc
        root = ElementTree.fromstring(document_xml)
        namespace = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        paragraphs = []
        for paragraph in root.iter(f"{namespace}p"):
            text = "".join(
                node.text or "" for node in paragraph.iter(f"{namespace}t")
            )
            if text:
                paragraphs.append(text)
        return "\n\n".join(paragraphs)

    @staticmethod
    def _add_docx_markdown_runs(paragraph: Any, text: str) -> None:
        pattern = re.compile(r"(\*\*(.+?)\*\*|`([^`]+)`)")
        cursor = 0
        for match in pattern.finditer(text):
            if match.start() > cursor:
                paragraph.add_run(text[cursor : match.start()])
            if match.group(2) is not None:
                run = paragraph.add_run(match.group(2))
                run.bold = True
            else:
                run = paragraph.add_run(match.group(3) or "")
                run.font.name = "Menlo"
            cursor = match.end()
        if cursor < len(text):
            paragraph.add_run(text[cursor:])

    @classmethod
    def _render_markdown_into_docx(
        cls,
        document: Any,
        *,
        title: str,
        content: str,
    ) -> None:
        body = document._element.body
        for child in list(body):
            tag = str(child.tag).rsplit("}", 1)[-1]
            if tag != "sectPr":
                body.remove(child)

        style_names = {str(style.name) for style in document.styles}

        def paragraph(style: str | None = None) -> Any:
            if style and style in style_names:
                return document.add_paragraph(style=style)
            return document.add_paragraph()

        normalized_title = title.strip()
        if normalized_title:
            title_paragraph = paragraph("Title")
            title_paragraph.add_run(normalized_title)

        lines = content.replace("\r\n", "\n").replace("\r", "\n").split("\n")
        index = 0
        while index < len(lines):
            raw = lines[index]
            line = raw.strip()
            if not line:
                index += 1
                continue

            if (
                line.startswith("|")
                and index + 1 < len(lines)
                and re.fullmatch(
                    r"\s*\|?\s*:?-{3,}:?\s*"
                    r"(?:\|\s*:?-{3,}:?\s*)+\|?\s*",
                    lines[index + 1],
                )
            ):
                headers = [
                    value.strip() for value in line.strip("|").split("|")
                ]
                table_rows: list[list[str]] = []
                index += 2
                while index < len(lines) and lines[index].strip().startswith(
                    "|"
                ):
                    values = [
                        value.strip()
                        for value in lines[index].strip().strip("|").split("|")
                    ]
                    if len(values) < len(headers):
                        values.extend([""] * (len(headers) - len(values)))
                    table_rows.append(values[: len(headers)])
                    index += 1
                table = document.add_table(
                    rows=1 + len(table_rows),
                    cols=max(1, len(headers)),
                )
                if "Table Grid" in style_names:
                    table.style = "Table Grid"
                for column, value in enumerate(headers):
                    cell_paragraph = table.rows[0].cells[column].paragraphs[0]
                    cls._add_docx_markdown_runs(cell_paragraph, value)
                    for run in cell_paragraph.runs:
                        run.bold = True
                for row_index, values in enumerate(table_rows, start=1):
                    for column, value in enumerate(values):
                        cell_paragraph = table.rows[row_index].cells[
                            column
                        ].paragraphs[0]
                        cls._add_docx_markdown_runs(cell_paragraph, value)
                continue

            heading = re.match(r"^(#{1,4})\s+(.+)$", line)
            if heading:
                level = len(heading.group(1))
                item = paragraph(f"Heading {level}")
                cls._add_docx_markdown_runs(item, heading.group(2).strip())
                index += 1
                continue

            bullet = re.match(r"^[-*+]\s+(.+)$", line)
            if bullet:
                item = paragraph("List Bullet")
                cls._add_docx_markdown_runs(item, bullet.group(1).strip())
                index += 1
                continue

            numbered = re.match(r"^\d+[.)]\s+(.+)$", line)
            if numbered:
                item = paragraph("List Number")
                cls._add_docx_markdown_runs(item, numbered.group(1).strip())
                index += 1
                continue

            item = paragraph("Normal")
            cls._add_docx_markdown_runs(item, line)
            index += 1

    @classmethod
    def _render_docx_roundtrip(
        cls,
        source_data: bytes,
        *,
        title: str,
        content: str,
    ) -> bytes:
        document = cls._open_docx_bytes(
            source_data,
            code="local_document_format_invalid",
            message="Word 文档结构无效，无法保存",
        )
        cls._render_markdown_into_docx(
            document,
            title=title,
            content=content,
        )
        return cls._save_docx_bytes(
            document,
            code="docx_roundtrip_failed",
            message="Word 文档保存失败，原文件未被修改",
        )

    def document_text(self, document_id: str) -> dict[str, Any]:
        context = self._context()
        project_id, state, entry = self._document_entry(document_id)
        if str(state.get("_localSandboxId") or "") != context.sandbox_id:
            raise LocalRuntimeError(
                409,
                "local_storage_sandbox_changed",
                "本机工作空间已切换，请重试",
            )
        path, row = self._source_path(
            entry,
            sandbox_id=context.sandbox_id,
        )
        media_type = str(entry.get("mediaType") or row["media_type"] or "")
        if path.suffix.lower() == ".docx":
            content = self._docx_text(path)
            kind = "docx"
        elif path.suffix.lower() == ".pdf" or media_type == "application/pdf":
            try:
                from pypdf import PdfReader
            except ImportError as exc:
                raise LocalRuntimeError(
                    503,
                    "pdf_text_executor_not_connected",
                    "PDF 文字读取组件尚未安装，请重新安装当前版本",
                ) from exc
            try:
                reader = PdfReader(str(path))
                if reader.is_encrypted:
                    try:
                        reader.decrypt("")
                    except Exception as exc:
                        raise LocalRuntimeError(
                            415,
                            "local_document_pdf_encrypted",
                            "PDF 已加密，暂时无法读取正文",
                        ) from exc
                content = "\n\n".join(
                    str(page.extract_text() or "").strip()
                    for page in reader.pages
                ).strip()
            except LocalRuntimeError:
                raise
            except Exception as exc:
                raise LocalRuntimeError(
                    415,
                    "local_document_pdf_invalid",
                    "PDF 文件结构异常，无法读取正文",
                ) from exc
            if not content:
                raise LocalRuntimeError(
                    415,
                    "local_document_ocr_required",
                    "PDF 未检测到可读取文字，需要 OCR；自动 OCR 尚未接通",
                )
            kind = "pdf"
        elif media_type.startswith("text/") or path.suffix.lower() in {
            ".md",
            ".txt",
            ".json",
            ".csv",
            ".tsv",
            ".xml",
            ".html",
        }:
            try:
                content = path.read_text(encoding="utf-8")
            except UnicodeDecodeError as exc:
                raise LocalRuntimeError(
                    415,
                    "local_document_encoding_unsupported",
                    "当前文件不是 UTF-8 文本，暂不能在编辑器中打开",
                ) from exc
            kind = path.suffix.lower().lstrip(".") or "text"
        else:
            raise LocalRuntimeError(
                415,
                "local_document_preview_unsupported",
                "该文件格式只能作为本机原文件使用，暂不能以文本编辑器打开",
            )
        return {
            "documentId": document_id,
            "projectId": project_id,
            "content": content,
            "kind": kind,
            "title": entry.get("title") or entry.get("fileName") or path.name,
            "fileName": entry.get("fileName") or path.name,
            "path": str(path),
            "sourceScope": "local_private",
            "contentHash": row["content_hash"],
            "byteSize": int(row["byte_size"]),
            "mediaType": media_type,
            "storageVersion": int(row["version"]),
        }

    def _update_docx_document_text(
        self,
        document_id: str,
        *,
        project_id: str,
        title: str,
        content: str,
        expected_version: int | None,
        idempotency_key: str | None,
    ) -> dict[str, Any]:
        normalized_content = content.replace("\r\n", "\n").strip()
        if not normalized_content:
            raise LocalRuntimeError(
                422,
                "document_content_required",
                "请先输入文档内容",
            )
        context = self._context()
        project_state_id = (
            f"project-state:{self._stable_segment(project_id)}"
        )
        with self.runtime.local_storage_object_lock(
            sandbox_id=context.sandbox_id,
            object_id=project_state_id,
        ):
            state = self._load_project_state(project_id)
            raw_entry = (state.get("documents") or {}).get(document_id)
            if not isinstance(raw_entry, Mapping):
                raise LocalRuntimeError(
                    404,
                    "local_document_missing",
                    "当前设备没有该资料的源文件",
                )
            entry = dict(raw_entry)
            source_id = str(entry.get("localSourceId") or "")
            if not source_id:
                raise LocalRuntimeError(
                    409,
                    "local_document_source_missing",
                    "当前设备的资料源对象缺失",
                )
            with self.runtime.local_storage_object_lock(
                sandbox_id=context.sandbox_id,
                object_id=source_id,
            ):
                row = self.runtime.local_storage_object_get(
                    sandbox_id=context.sandbox_id,
                    object_id=source_id,
                )
                if row is None or str(row["lifecycle_state"]) != "active":
                    raise LocalRuntimeError(
                        404,
                        "local_document_source_missing",
                        "当前设备的资料源文件不存在或已移除",
                    )
                current_version = int(row["version"])
                storage_key = str(row["storage_key"])
                path = self._managed_path(storage_key)
                if path.suffix.casefold() != ".docx":
                    raise LocalRuntimeError(
                        415,
                        "local_document_update_unsupported",
                        "当前资料不是 .docx Word 文档",
                    )
                try:
                    current_data = path.read_bytes()
                except OSError as exc:
                    raise LocalRuntimeError(
                        404,
                        "local_document_source_missing",
                        "当前设备的资料源文件无法读取",
                    ) from exc
                current_hash = hashlib.sha256(current_data).hexdigest()
                if (
                    len(current_data) != int(row["byte_size"])
                    or current_hash != str(row["content_hash"])
                ):
                    raise LocalRuntimeError(
                        409,
                        "local_document_source_corrupt",
                        "当前设备的资料源文件校验失败",
                    )

                normalized_title = (
                    title.strip()
                    or str(entry.get("title") or "").strip()
                    or path.stem
                )
                request_fingerprint = hashlib.sha256(
                    canonical_json(
                        {
                            "documentId": document_id,
                            "title": normalized_title,
                            "contentHash": hashlib.sha256(
                                normalized_content.encode("utf-8")
                            ).hexdigest(),
                        }
                    ).encode("utf-8")
                ).hexdigest()
                operation_identity = (
                    f"explicit:{idempotency_key}"
                    if str(idempotency_key or "").strip()
                    else f"request:{request_fingerprint}"
                )
                operation_hash = hashlib.sha256(
                    operation_identity.encode("utf-8")
                ).hexdigest()
                previous_key_hash = str(
                    entry.get("lastEditorIdempotencyKeyHash") or ""
                )
                previous_fingerprint = str(
                    entry.get("lastEditorRequestFingerprint") or ""
                )
                previous_result_hash = str(
                    entry.get("lastEditorResultContentHash") or ""
                )
                if previous_key_hash == operation_hash:
                    if previous_fingerprint != request_fingerprint:
                        raise LocalRuntimeError(
                            409,
                            "document_edit_idempotency_conflict",
                            "文档保存操作标识已用于不同内容",
                        )
                    if current_hash == previous_result_hash:
                        return {
                            "clientId": project_id,
                            "documentId": document_id,
                            "title": normalized_title,
                            "fileName": (
                                entry.get("fileName") or path.name
                            ),
                            "path": str(path),
                            "sourceScope": "local_private",
                            "contentHash": current_hash,
                            "byteSize": len(current_data),
                            "mediaType": self.DOCX_MEDIA_TYPE,
                            "storageVersion": current_version,
                            "idempotentReplay": True,
                        }
                if (
                    expected_version is not None
                    and int(expected_version) != current_version
                ):
                    raise LocalRuntimeError(
                        409,
                        "local_storage_version_conflict",
                        "Word 文档版本已变化，请刷新后重试",
                    )

                output_data = self._render_docx_roundtrip(
                    current_data,
                    title=normalized_title,
                    content=normalized_content,
                )
                try:
                    stored = self._upsert_object(
                        sandbox_id=context.sandbox_id,
                        object_id=source_id,
                        storage_key=storage_key,
                        media_type=self.DOCX_MEDIA_TYPE,
                        data=output_data,
                        expected_version=current_version,
                    )
                except Exception:
                    temporary = path.with_name(
                        f".{path.name}.{new_id()}.restore"
                    )
                    try:
                        temporary.write_bytes(current_data)
                        temporary.replace(path)
                    finally:
                        if temporary.exists():
                            temporary.unlink()
                    raise

                summary_id = str(entry.get("localSummaryId") or "")
                if summary_id:
                    summary_row = self.runtime.local_storage_object_get(
                        sandbox_id=context.sandbox_id,
                        object_id=summary_id,
                    )
                    if summary_row is not None:
                        summary_payload = {
                            "schema": (
                                "yiyu.project-local-private-knowledge.v1"
                            ),
                            "sourceScope": "local_private",
                            "projectId": project_id,
                            "sourceId": source_id,
                            "contentHash": stored["contentHash"],
                            "summary": normalized_content[:2000],
                            "summaryKind": "text_excerpt",
                            "sourceDescription": (
                                "当前设备受管的本机私有项目 Word 资料"
                            ),
                            "updatedAt": stored["updatedAt"],
                            "fileName": entry.get("fileName") or path.name,
                        }
                        self._upsert_object(
                            sandbox_id=context.sandbox_id,
                            object_id=summary_id,
                            storage_key=str(summary_row["storage_key"]),
                            media_type=self.SUMMARY_MEDIA_TYPE,
                            data=canonical_json(summary_payload).encode(
                                "utf-8"
                            ),
                            expected_version=int(summary_row["version"]),
                        )
                entry.update(
                    {
                        "title": normalized_title,
                        "contentHash": stored["contentHash"],
                        "byteSize": stored["byteSize"],
                        "mediaType": self.DOCX_MEDIA_TYPE,
                        "updatedAt": stored["updatedAt"],
                        "localSourceVersion": stored["version"],
                        "lastEditorRequestFingerprint": request_fingerprint,
                        "lastEditorIdempotencyKeyHash": operation_hash,
                        "lastEditorResultContentHash": stored["contentHash"],
                    }
                )
                state["documents"][document_id] = entry
                self._write_project_state(project_id, state)
                return {
                    "clientId": project_id,
                    "documentId": document_id,
                    "title": normalized_title,
                    "fileName": entry.get("fileName") or path.name,
                    "path": str(path),
                    "sourceScope": "local_private",
                    "contentHash": stored["contentHash"],
                    "byteSize": stored["byteSize"],
                    "mediaType": self.DOCX_MEDIA_TYPE,
                    "storageVersion": stored["version"],
                    "idempotentReplay": False,
                }

    def update_document_text(
        self,
        document_id: str,
        *,
        title: str,
        content: str,
        expected_version: int | None = None,
        idempotency_key: str | None = None,
    ) -> dict[str, Any]:
        context = self._context()
        project_id, state, entry = self._document_entry(document_id)
        if str(state.get("_localSandboxId") or "") != context.sandbox_id:
            raise LocalRuntimeError(
                409,
                "local_storage_sandbox_changed",
                "本机工作空间已切换，请重试",
            )
        path, row = self._source_path(
            entry,
            sandbox_id=context.sandbox_id,
        )
        if path.suffix.lower() == ".docx":
            return self._update_docx_document_text(
                document_id,
                project_id=project_id,
                title=title,
                content=content,
                expected_version=expected_version,
                idempotency_key=idempotency_key,
            )
        if (
            expected_version is not None
            and int(expected_version) != int(row["version"])
        ):
            raise LocalRuntimeError(
                409,
                "local_storage_version_conflict",
                "本机文档版本已变化，请刷新后重试",
            )
        media_type = str(entry.get("mediaType") or row["media_type"] or "")
        if not (
            media_type.startswith("text/")
            or path.suffix.lower()
            in {".md", ".txt", ".json", ".csv", ".tsv", ".xml", ".html"}
        ):
            raise LocalRuntimeError(
                415,
                "local_document_update_unsupported",
                "该文件格式不能在文本编辑器中覆盖保存",
            )
        data = content.encode("utf-8")
        stored = self._upsert_object(
            sandbox_id=context.sandbox_id,
            object_id=str(entry["localSourceId"]),
            storage_key=str(row["storage_key"]),
            media_type=media_type or "text/plain",
            data=data,
            expected_version=int(row["version"]),
        )
        summary_id = str(entry.get("localSummaryId") or "")
        if summary_id:
            summary_row = self.runtime.local_storage_object_get(
                sandbox_id=context.sandbox_id,
                object_id=summary_id
            )
            if summary_row is not None:
                summary_payload = {
                    "schema": "yiyu.project-local-private-knowledge.v1",
                    "sourceScope": "local_private",
                    "projectId": project_id,
                    "sourceId": entry["localSourceId"],
                    "contentHash": stored["contentHash"],
                    "summary": content[:2000],
                    "summaryKind": "text_excerpt",
                    "sourceDescription": "当前设备受管的本机私有项目文本资料",
                    "updatedAt": stored["updatedAt"],
                    "fileName": entry.get("fileName") or path.name,
                }
                self._upsert_object(
                    sandbox_id=context.sandbox_id,
                    object_id=summary_id,
                    storage_key=str(summary_row["storage_key"]),
                    media_type=self.SUMMARY_MEDIA_TYPE,
                    data=canonical_json(summary_payload).encode("utf-8"),
                    expected_version=int(summary_row["version"]),
                )
        entry.update(
            {
                "title": title.strip() or entry.get("title") or path.stem,
                "contentHash": stored["contentHash"],
                "byteSize": stored["byteSize"],
                "updatedAt": stored["updatedAt"],
            }
        )
        state["documents"][document_id] = entry
        self._write_project_state(project_id, state)
        return {
            "clientId": project_id,
            "documentId": document_id,
            "title": entry["title"],
            "fileName": entry.get("fileName") or path.name,
            "path": str(path),
            "sourceScope": "local_private",
            "contentHash": stored["contentHash"],
            "byteSize": stored["byteSize"],
            "mediaType": media_type or "text/plain",
            "storageVersion": stored["version"],
        }

    def optimization_candidates(
        self,
        project_ids: Iterable[str],
    ) -> list[dict[str, Any]]:
        """List current-device documents that can participate in deep reading."""

        context = self._context()
        candidates: list[dict[str, Any]] = []
        for project_id in dict.fromkeys(
            str(value).strip() for value in project_ids if str(value).strip()
        ):
            state = self._load_project_state(project_id)
            for document_id, raw_entry in dict(
                state.get("documents") or {}
            ).items():
                if not isinstance(raw_entry, Mapping):
                    continue
                entry = dict(raw_entry)
                summary_id = str(entry.get("localSummaryId") or "")
                if not summary_id:
                    continue
                summary_row = self.runtime.local_storage_object_get(
                    sandbox_id=context.sandbox_id,
                    object_id=summary_id,
                )
                if summary_row is None:
                    continue
                summary_path = self._managed_path(str(summary_row["storage_key"]))
                try:
                    summary_payload = json.loads(summary_path.read_text("utf-8"))
                except (OSError, UnicodeDecodeError, json.JSONDecodeError):
                    summary_payload = {}
                summary_kind = (
                    str(summary_payload.get("summaryKind") or "")
                    if isinstance(summary_payload, Mapping)
                    else ""
                )
                candidates.append(
                    {
                        "projectId": project_id,
                        "documentId": str(document_id),
                        "title": str(
                            entry.get("title")
                            or entry.get("fileName")
                            or document_id
                        ),
                        "contentHash": str(entry.get("contentHash") or ""),
                        "summaryKind": summary_kind,
                        "deepRead": summary_kind == "ai_summary",
                    }
                )
        return candidates

    def update_ai_summary(
        self,
        document_id: str,
        *,
        summary: str,
        model_name: str,
    ) -> dict[str, Any]:
        """Replace a local-private sidecar with an AI-derived summary."""

        normalized = summary.strip()
        if not normalized:
            raise LocalRuntimeError(
                422,
                "local_ai_summary_empty",
                "深度解析没有生成可保存的摘要",
            )
        context = self._context()
        project_id, state, entry = self._document_entry(document_id)
        if str(state.get("_localSandboxId") or "") != context.sandbox_id:
            raise LocalRuntimeError(
                409,
                "local_storage_sandbox_changed",
                "本机工作空间已切换，请重试",
            )
        summary_id = str(entry.get("localSummaryId") or "")
        if not summary_id:
            raise LocalRuntimeError(
                409,
                "local_summary_missing",
                "该资料尚无本机摘要对象",
            )
        row = self.runtime.local_storage_object_get(
            sandbox_id=context.sandbox_id,
            object_id=summary_id,
        )
        if row is None:
            raise LocalRuntimeError(
                409,
                "local_summary_missing",
                "该资料的本机摘要对象不存在",
            )
        payload = {
            "schema": "yiyu.project-local-private-knowledge.v1",
            "sourceScope": "local_private",
            "projectId": project_id,
            "sourceId": entry["localSourceId"],
            "contentHash": entry.get("contentHash") or "",
            "summary": normalized[:8000],
            "summaryKind": "ai_summary",
            "sourceDescription": "当前设备受管资料的本机深度解析摘要",
            "modelName": model_name,
            "updatedAt": utc_now(),
            "fileName": entry.get("fileName") or entry.get("title") or "",
        }
        stored = self._upsert_object(
            sandbox_id=context.sandbox_id,
            object_id=summary_id,
            storage_key=str(row["storage_key"]),
            media_type=self.SUMMARY_MEDIA_TYPE,
            data=canonical_json(payload).encode("utf-8"),
            expected_version=int(row["version"]),
        )
        entry["summaryKind"] = "ai_summary"
        entry["updatedAt"] = stored["updatedAt"]
        state["documents"][document_id] = entry
        self._write_project_state(project_id, state)
        return {
            "projectId": project_id,
            "documentId": document_id,
            "summaryKind": "ai_summary",
            "summaryHash": hashlib.sha256(
                normalized.encode("utf-8")
            ).hexdigest(),
            "modelName": model_name,
            "updatedAt": stored["updatedAt"],
        }

    @staticmethod
    def _folder_dto(
        project_id: str,
        folder: Mapping[str, Any],
        documents: Mapping[str, Any],
    ) -> dict[str, Any]:
        folder_id = str(folder.get("id") or "")
        return {
            "id": folder_id,
            "clientId": project_id,
            "label": folder.get("label") or "未命名文件夹",
            "path": "",
            "fileCount": sum(
                str(item.get("folderId") or "") == folder_id
                for item in documents.values()
                if isinstance(item, dict)
            ),
            "lastScannedAt": folder.get("updatedAt"),
            "folderKind": "local_virtual",
            "sourceType": "member_local",
            "isSystem": False,
            "isHidden": bool(folder.get("isHidden")),
            "sortOrder": int(folder.get("sortOrder") or 0),
            "createdByRule": folder.get("createdByRule"),
        }

    def folders(self, project_id: str) -> list[dict[str, Any]]:
        state = self._load_project_state(project_id)
        documents = dict(state.get("documents") or {})
        folders = [
            self._folder_dto(project_id, folder, documents)
            for folder in state.get("folders") or []
        ]
        return sorted(
            folders,
            key=lambda item: (item["sortOrder"], str(item["label"])),
        )

    def documents(self, project_id: str) -> list[dict[str, Any]]:
        """Return only files that physically exist in the current sandbox.

        Organization-cloud knowledge metadata is intentionally excluded. It is
        exposed through the project knowledge context instead of masquerading
        as an openable member-local file.
        """
        state = self._load_project_state(project_id)
        sandbox_id = str(state.get("_localSandboxId") or "")
        result: list[dict[str, Any]] = []
        for document_id, raw_entry in dict(state.get("documents") or {}).items():
            if not isinstance(raw_entry, Mapping):
                continue
            entry = dict(raw_entry)
            try:
                managed_path, _ = self._source_path(
                    entry,
                    sandbox_id=sandbox_id,
                )
            except LocalRuntimeError:
                continue
            raw_original = str(entry.get("originalSourcePath") or "").strip()
            original_path = (
                Path(raw_original).expanduser()
                if raw_original
                else None
            )
            open_path = (
                original_path
                if original_path is not None and original_path.is_file()
                else managed_path
            )
            suffix = managed_path.suffix.lower().lstrip(".")
            processing = self.processing_state(entry)
            result.append(
                {
                    "id": str(document_id),
                    "clientId": project_id,
                    "folderId": entry.get("folderId"),
                    "title": (
                        entry.get("title")
                        or entry.get("fileName")
                        or managed_path.name
                    ),
                    "path": str(open_path),
                    "managedPath": str(managed_path),
                    "originalSourcePath": (
                        str(original_path)
                        if original_path is not None
                        else None
                    ),
                    "kind": suffix or "local_private",
                    "source": "member_local",
                    "excerpt": "",
                    "tags": [
                        "local_private",
                        str(entry.get("cloudMetadataState") or "ready"),
                    ],
                    "cloudMetadataState": (
                        entry.get("cloudMetadataState") or "ready"
                    ),
                    "localState": "ready",
                    "importedAt": entry.get("updatedAt"),
                    **processing,
                }
            )
        return sorted(
            result,
            key=lambda item: str(item.get("importedAt") or ""),
            reverse=True,
        )

    def duplicate_document_groups(self, project_id: str) -> dict[str, Any]:
        """Derive duplicate candidates from current-device source files only."""
        state = self._load_project_state(project_id)
        entries = dict(state.get("documents") or {})
        visible = {str(item["id"]): item for item in self.documents(project_id)}
        by_hash: dict[str, list[dict[str, Any]]] = {}
        by_name_without_hash: dict[str, list[dict[str, Any]]] = {}
        for document_id, document in visible.items():
            raw = entries.get(document_id)
            if not isinstance(raw, Mapping):
                continue
            content_hash = str(raw.get("contentHash") or "").strip()
            file_name = str(raw.get("fileName") or document.get("title") or "").strip()
            item = {
                "id": document_id,
                "documentId": document_id,
                "fileName": file_name,
                "kind": document.get("kind") or "local_private",
                "managedPath": document.get("managedPath") or "",
                "originalPath": document.get("originalSourcePath") or "",
                "contentHash": content_hash,
                "parseStatus": document.get("parseStatus") or document.get("parseState"),
                "sectionCount": 0,
                "chunkCount": 0,
                "importedAt": document.get("importedAt"),
                "fileSizeBytes": int(raw.get("byteSize") or 0),
                "refTaskAttachmentCount": 0,
                "refEvidenceCardCount": 0,
                "refAtomicFactCount": 0,
            }
            if content_hash:
                by_hash.setdefault(content_hash, []).append(item)
            elif file_name:
                by_name_without_hash.setdefault(file_name.casefold(), []).append(item)
        groups: list[dict[str, Any]] = []
        for content_hash, items in sorted(by_hash.items()):
            if len(items) < 2:
                continue
            groups.append(
                {
                    "groupKey": f"hash:{content_hash}",
                    "groupType": "same_content_hash",
                    "fileName": items[0]["fileName"],
                    "contentHash": content_hash,
                    "count": len(items),
                    "documents": items,
                }
            )
        for normalized_name, items in sorted(by_name_without_hash.items()):
            if len(items) < 2:
                continue
            name_hash = hashlib.sha256(normalized_name.encode("utf-8")).hexdigest()
            groups.append(
                {
                    "groupKey": f"name:{name_hash}",
                    "groupType": "same_filename",
                    "fileName": items[0]["fileName"],
                    "contentHash": "",
                    "count": len(items),
                    "documents": items,
                }
            )
        return {"groups": groups, "generatedAt": utc_now()}

    def create_folder(self, project_id: str, payload: Mapping[str, Any]) -> dict[str, Any]:
        label = str(payload.get("label") or "").strip()
        if not label:
            raise LocalRuntimeError(422, "folder_label_required", "请输入文件夹名称")
        state = self._load_project_state(project_id)
        if any(
            str(item.get("label") or "").casefold() == label.casefold()
            for item in state.get("folders") or []
        ):
            raise LocalRuntimeError(409, "folder_label_exists", "已存在同名文件夹")
        now = utc_now()
        folder = {
            "id": new_id(),
            "label": label,
            "isHidden": bool(payload.get("isHidden")),
            "sortOrder": int(
                payload.get("sortOrder")
                if payload.get("sortOrder") is not None
                else len(state.get("folders") or [])
            ),
            "createdByRule": payload.get("createdByRule"),
            "createdAt": now,
            "updatedAt": now,
        }
        state.setdefault("folders", []).append(folder)
        self._write_project_state(project_id, state)
        return self._folder_dto(
            project_id,
            folder,
            dict(state.get("documents") or {}),
        )

    def update_folder(
        self,
        project_id: str,
        folder_id: str,
        payload: Mapping[str, Any],
    ) -> dict[str, Any]:
        state = self._load_project_state(project_id)
        folder = next(
            (
                item
                for item in state.get("folders") or []
                if str(item.get("id") or "") == folder_id
            ),
            None,
        )
        if folder is None:
            raise LocalRuntimeError(404, "folder_missing", "文件夹不存在")
        if "label" in payload:
            label = str(payload.get("label") or "").strip()
            if not label:
                raise LocalRuntimeError(422, "folder_label_required", "请输入文件夹名称")
            folder["label"] = label
        if "isHidden" in payload:
            folder["isHidden"] = bool(payload.get("isHidden"))
        if "sortOrder" in payload:
            folder["sortOrder"] = int(payload.get("sortOrder") or 0)
        folder["updatedAt"] = utc_now()
        self._write_project_state(project_id, state)
        return self._folder_dto(
            project_id,
            folder,
            dict(state.get("documents") or {}),
        )

    def delete_folder(self, project_id: str, folder_id: str) -> dict[str, Any]:
        state = self._load_project_state(project_id)
        original = list(state.get("folders") or [])
        state["folders"] = [
            item
            for item in original
            if str(item.get("id") or "") != folder_id
        ]
        if len(state["folders"]) == len(original):
            raise LocalRuntimeError(404, "folder_missing", "文件夹不存在")
        for entry in (state.get("documents") or {}).values():
            if isinstance(entry, dict) and str(entry.get("folderId") or "") == folder_id:
                entry["folderId"] = None
        self._write_project_state(project_id, state)
        return {"ok": True, "deleted": True, "folderId": folder_id}

    def move_document(
        self,
        project_id: str,
        document_id: str,
        payload: Mapping[str, Any],
    ) -> dict[str, Any]:
        state = self._load_project_state(project_id)
        entry = (state.get("documents") or {}).get(document_id)
        if not isinstance(entry, dict):
            raise LocalRuntimeError(
                404,
                "local_document_missing",
                "当前设备没有该资料的本机文件映射",
            )
        folder_id = str(payload.get("folderId") or "")
        if not folder_id and payload.get("folderLabel"):
            label = str(payload.get("folderLabel") or "").strip()
            folder = next(
                (
                    item
                    for item in state.get("folders") or []
                    if str(item.get("label") or "").casefold() == label.casefold()
                ),
                None,
            )
            if folder is None:
                folder = self.create_folder(project_id, {"label": label})
                state = self._load_project_state(project_id)
                folder_id = str(folder["id"])
            else:
                folder_id = str(folder["id"])
        if folder_id and not any(
            str(item.get("id") or "") == folder_id
            for item in state.get("folders") or []
        ):
            raise LocalRuntimeError(404, "folder_missing", "目标文件夹不存在")
        entry["folderId"] = folder_id or None
        entry["updatedAt"] = utc_now()
        state["documents"][document_id] = entry
        self._write_project_state(project_id, state)
        return {
            "id": document_id,
            "clientId": project_id,
            "folderId": entry["folderId"],
            "title": entry.get("title") or entry.get("fileName") or "",
            "path": "",
            "kind": "local_private",
            "source": "member_local",
            "excerpt": "",
            "tags": [],
            "importedAt": entry.get("updatedAt"),
        }

    def apply_folder_labels(
        self,
        project_id: str,
        labels: Iterable[Any],
    ) -> dict[str, Any]:
        normalized = [
            str(label).strip()
            for label in labels
            if str(label).strip()
        ]
        for label in dict.fromkeys(normalized):
            try:
                self.create_folder(
                    project_id,
                    {"label": label, "createdByRule": "strict_recommendation"},
                )
            except LocalRuntimeError as exc:
                if exc.code != "folder_label_exists":
                    raise
        return {
            "clientId": project_id,
            "folders": self.folders(project_id),
            "state": "completed",
        }

    def resolve_duplicates(
        self,
        project_id: str,
        payload: Mapping[str, Any],
    ) -> dict[str, Any]:
        normalized = self.preflight_duplicate_resolution(project_id, payload)
        group_key = normalized["groupKey"]
        action = normalized["action"]
        keep_ids = list(normalized["keepV2DocumentIds"])
        delete_ids = list(normalized["deleteV2DocumentIds"])
        state = self._load_project_state(project_id)
        deleted = 0
        if action == "delete_others":
            for document_id in delete_ids:
                entry = state["documents"][document_id]
                source_id = str(entry.get("localSourceId") or "")
                if source_id:
                    self.runtime.local_storage_object_set_lifecycle(
                        object_id=source_id,
                        lifecycle_state="deleted",
                    )
                deleted += 1
        state.setdefault("duplicateResolutions", {})[group_key] = {
            "action": action,
            "keptDocumentIds": keep_ids,
            "deletedDocumentIds": delete_ids,
            "migrateReferences": normalized["migrateReferences"],
            "note": normalized["note"],
            "resolvedAt": utc_now(),
        }
        self._write_project_state(project_id, state)
        return {
            "action": action,
            "groupKey": group_key,
            "deletedCount": deleted,
            "recycledTo": "strict-local-storage:deleted",
            "migratedTaskAttachments": 0,
            "migratedEvidenceRefs": 0,
            "migratedAtomicFacts": 0,
            "keptDocumentIds": keep_ids,
        }

    def duplicate_cloud_delete_targets(
        self,
        project_id: str,
        document_ids: Iterable[str],
    ) -> list[dict[str, str]]:
        """Return only cloud metadata identifiers; never expose local paths."""
        state = self._load_project_state(project_id)
        documents = dict(state.get("documents") or {})
        targets: list[dict[str, str]] = []
        for document_id in document_ids:
            entry = documents.get(str(document_id))
            if not isinstance(entry, Mapping):
                raise LocalRuntimeError(
                    409,
                    "duplicate_resolution_not_local",
                    "只能处置当前项目中由本机持有的重复资料",
                )
            cloud_document_id = str(entry.get("cloudDocumentId") or "").strip()
            if cloud_document_id:
                targets.append(
                    {
                        "documentId": str(document_id),
                        "cloudDocumentId": cloud_document_id,
                    }
                )
        return targets

    def preflight_duplicate_resolution(
        self,
        project_id: str,
        payload: Mapping[str, Any],
    ) -> dict[str, Any]:
        group_key = str(payload.get("groupKey") or "").strip()
        action = str(payload.get("action") or "").strip()
        if not group_key or action not in {"delete_others", "keep_all"}:
            raise LocalRuntimeError(
                422,
                "duplicate_resolution_invalid",
                "重复资料处置请求无效",
            )
        keep_ids = [
            str(value)
            for value in payload.get("keepV2DocumentIds") or []
            if str(value)
        ]
        delete_ids = [
            str(value)
            for value in payload.get("deleteV2DocumentIds") or []
            if str(value)
        ]
        if (
            len(set(keep_ids)) != len(keep_ids)
            or len(set(delete_ids)) != len(delete_ids)
            or set(keep_ids) & set(delete_ids)
        ):
            raise LocalRuntimeError(
                422,
                "duplicate_resolution_documents_invalid",
                "重复资料处置清单存在重复或冲突",
            )
        if action == "delete_others" and not delete_ids:
            raise LocalRuntimeError(
                422,
                "duplicate_resolution_documents_required",
                "请选择要移除的重复资料",
            )
        if action == "keep_all" and delete_ids:
            raise LocalRuntimeError(
                422,
                "duplicate_resolution_keep_all_invalid",
                "保留全部资料时不能包含待移除资料",
            )
        state = self._load_project_state(project_id)
        missing = [
            document_id
            for document_id in [*keep_ids, *delete_ids]
            if document_id not in (state.get("documents") or {})
        ]
        if missing:
            raise LocalRuntimeError(
                409,
                "duplicate_resolution_not_local",
                "只能在当前设备处置本机拥有源文件的重复资料",
            )
        sandbox_id = str(state.get("_localSandboxId") or "")
        unavailable = []
        for document_id in delete_ids:
            source_id = str(
                state["documents"][document_id].get("localSourceId") or ""
            )
            source = (
                self.runtime.local_storage_object_get(
                    sandbox_id=sandbox_id,
                    object_id=source_id,
                )
                if source_id
                else None
            )
            if source is None or str(source["lifecycle_state"]) != "active":
                unavailable.append(document_id)
        if unavailable:
            raise LocalRuntimeError(
                409,
                "duplicate_resolution_source_unavailable",
                "待移除资料的本机源文件不存在或已处置",
            )
        return {
            "groupKey": group_key,
            "action": action,
            "keepV2DocumentIds": keep_ids,
            "deleteV2DocumentIds": delete_ids,
            "migrateReferences": bool(payload.get("migrateReferences")),
            "note": str(payload.get("note") or ""),
        }

    @staticmethod
    def _normalize_docx_field_label(value: Any) -> str:
        label = re.sub(r"[\s\u3000]+", " ", str(value or "")).strip()
        label = label.strip(":：-_[]【】")
        return label[:120]

    @staticmethod
    def _require_docx_dependency() -> Any:
        try:
            from docx import Document
        except ImportError as exc:
            raise LocalRuntimeError(
                503,
                "docx_executor_dependency_missing",
                "本机 Word 执行器依赖缺失，请重新安装严格新版",
            ) from exc
        return Document

    @classmethod
    def _open_docx_bytes(
        cls,
        data: bytes,
        *,
        code: str = "template_docx_invalid",
        message: str = "Word 模板结构损坏，无法填写",
    ) -> Any:
        Document = cls._require_docx_dependency()
        try:
            return Document(BytesIO(data))
        except Exception as exc:
            raise LocalRuntimeError(422, code, message) from exc

    @classmethod
    def _save_docx_bytes(
        cls,
        document: Any,
        *,
        code: str = "docx_generation_failed",
        message: str = "Word 文档生成失败，原文件未被修改",
    ) -> bytes:
        stream = BytesIO()
        try:
            document.save(stream)
        except Exception as exc:
            raise LocalRuntimeError(422, code, message) from exc
        data = stream.getvalue()
        cls._open_docx_bytes(
            data,
            code=code,
            message=message,
        )
        return data

    @staticmethod
    def _docx_containers(document: Any) -> list[Any]:
        containers = [document]
        seen = {id(document.element)}
        for section in document.sections:
            for candidate in (
                section.header,
                section.first_page_header,
                section.even_page_header,
                section.footer,
                section.first_page_footer,
                section.even_page_footer,
            ):
                key = id(candidate._element)
                if key not in seen:
                    seen.add(key)
                    containers.append(candidate)
        return containers

    @classmethod
    def _iter_docx_paragraphs(cls, container: Any) -> Iterable[Any]:
        seen: set[int] = set()

        def walk(current: Any) -> Iterable[Any]:
            for paragraph in current.paragraphs:
                key = id(paragraph._p)
                if key not in seen:
                    seen.add(key)
                    yield paragraph
            for table in current.tables:
                for row in table.rows:
                    for cell in row.cells:
                        yield from walk(cell)

        yield from walk(container)

    @classmethod
    def _iter_docx_tables(cls, container: Any) -> Iterable[Any]:
        seen: set[int] = set()

        def walk(current: Any) -> Iterable[Any]:
            for table in current.tables:
                key = id(table._tbl)
                if key in seen:
                    continue
                seen.add(key)
                yield table
                for row in table.rows:
                    for cell in row.cells:
                        yield from walk(cell)

        yield from walk(container)

    @classmethod
    def _docx_template_labels(cls, document: Any) -> list[str]:
        labels: list[str] = []
        seen: set[str] = set()

        def add(value: Any) -> None:
            label = cls._normalize_docx_field_label(value)
            if label and label not in seen:
                seen.add(label)
                labels.append(label)

        for container in cls._docx_containers(document):
            for paragraph in cls._iter_docx_paragraphs(container):
                for match in cls._DOCX_PLACEHOLDER_PATTERN.finditer(
                    str(paragraph.text or "")
                ):
                    add(match.group(1))
            for table in cls._iter_docx_tables(container):
                for row in table.rows:
                    if len(row.cells) < 2:
                        continue
                    target = str(row.cells[1].text or "").strip().casefold()
                    if target in cls._DOCX_EMPTY_MARKERS or any(
                        marker and marker in target
                        for marker in cls._DOCX_EMPTY_MARKERS
                    ):
                        add(row.cells[0].text)
        return labels

    @classmethod
    def _replace_docx_paragraph_placeholders(
        cls,
        paragraph: Any,
        values: Mapping[str, str],
    ) -> int:
        runs = list(paragraph.runs)
        text = "".join(str(run.text or "") for run in runs)
        matches = list(cls._DOCX_PLACEHOLDER_PATTERN.finditer(text))
        if not runs or not matches:
            return 0

        def locate(position: int) -> tuple[int, int]:
            cursor = 0
            for index, run in enumerate(runs):
                run_text = str(run.text or "")
                end = cursor + len(run_text)
                if position < end:
                    return index, position - cursor
                cursor = end
            return len(runs) - 1, len(str(runs[-1].text or ""))

        applied = 0
        for match in reversed(matches):
            label = cls._normalize_docx_field_label(match.group(1))
            replacement = str(values.get(label) or "").strip()
            if not replacement:
                continue
            start_index, start_offset = locate(match.start())
            end_index, end_offset = locate(match.end() - 1)
            start_run = runs[start_index]
            end_run = runs[end_index]
            if start_index == end_index:
                current = str(start_run.text or "")
                start_run.text = (
                    current[:start_offset]
                    + replacement
                    + current[end_offset + 1 :]
                )
            else:
                prefix = str(start_run.text or "")[:start_offset]
                suffix = str(end_run.text or "")[end_offset + 1 :]
                start_run.text = prefix + replacement
                for index in range(start_index + 1, end_index):
                    runs[index].text = ""
                end_run.text = suffix
            applied += 1
        return applied

    @classmethod
    def _apply_docx_template_values(
        cls,
        document: Any,
        values: Mapping[str, str],
    ) -> int:
        applied = 0
        for container in cls._docx_containers(document):
            for paragraph in cls._iter_docx_paragraphs(container):
                applied += cls._replace_docx_paragraph_placeholders(
                    paragraph,
                    values,
                )
            for table in cls._iter_docx_tables(container):
                for row in table.rows:
                    if len(row.cells) < 2:
                        continue
                    label = cls._normalize_docx_field_label(row.cells[0].text)
                    replacement = str(values.get(label) or "").strip()
                    if not replacement:
                        continue
                    current = str(row.cells[1].text or "").strip()
                    normalized = current.casefold()
                    if not (
                        normalized in cls._DOCX_EMPTY_MARKERS
                        or any(
                            marker and marker in normalized
                            for marker in cls._DOCX_EMPTY_MARKERS
                        )
                    ):
                        continue
                    row.cells[1].text = replacement
                    applied += 1
        return applied

    @classmethod
    def _docx_attachment_checklist(cls, document: Any) -> list[str]:
        attachments: list[str] = []
        for container in cls._docx_containers(document):
            for table in cls._iter_docx_tables(container):
                if not table.rows:
                    continue
                headers = [
                    str(cell.text or "").strip()
                    for cell in table.rows[0].cells
                ]
                column = next(
                    (
                        index
                        for index, value in enumerate(headers)
                        if "附件名称" in value
                    ),
                    None,
                )
                if column is None:
                    continue
                for row in table.rows[1:]:
                    if len(row.cells) <= column:
                        continue
                    name = cls._normalize_docx_field_label(
                        row.cells[column].text
                    )
                    if name and name not in attachments:
                        attachments.append(name)
        return attachments

    def _ensure_template_source_object(
        self,
        *,
        sandbox_id: str,
        project_id: str,
        data: bytes,
        content_hash: str,
    ) -> dict[str, Any]:
        object_id = (
            "template-source:"
            f"{self._stable_segment(project_id)}:{content_hash}"
        )
        storage_key = (
            "local-template-fill/"
            f"{self._stable_segment(sandbox_id)}/"
            f"{self._stable_segment(project_id)}/sources/"
            f"{content_hash}.docx"
        )
        with self.runtime.local_storage_object_lock(
            sandbox_id=sandbox_id,
            object_id=object_id,
        ):
            existing = self.runtime.local_storage_object_get(
                sandbox_id=sandbox_id,
                object_id=object_id,
            )
            if existing is None:
                return self._upsert_object(
                    sandbox_id=sandbox_id,
                    object_id=object_id,
                    storage_key=storage_key,
                    media_type=self.DOCX_MEDIA_TYPE,
                    data=data,
                    expected_version=0,
                )
            if (
                str(existing.get("storage_key") or "") != storage_key
                or str(existing.get("content_hash") or "") != content_hash
                or str(existing.get("media_type") or "")
                != self.DOCX_MEDIA_TYPE
                or str(existing.get("lifecycle_state") or "") != "active"
            ):
                raise LocalRuntimeError(
                    409,
                    "template_source_identity_conflict",
                    "本机模板源对象与当前请求不一致",
                )
            path = self._managed_path(storage_key)
            try:
                stored_data = path.read_bytes()
            except OSError as exc:
                raise LocalRuntimeError(
                    409,
                    "template_source_missing",
                    "本机受管模板源文件缺失",
                ) from exc
            if (
                len(stored_data) != int(existing["byte_size"])
                or hashlib.sha256(stored_data).hexdigest() != content_hash
            ):
                raise LocalRuntimeError(
                    409,
                    "template_source_corrupt",
                    "本机受管模板源文件校验失败",
                )
            return {
                "objectId": object_id,
                "storageKey": storage_key,
                "contentHash": content_hash,
                "mediaType": self.DOCX_MEDIA_TYPE,
                "byteSize": len(stored_data),
                "version": int(existing["version"]),
                "updatedAt": str(existing["updated_at"]),
                "path": str(path),
            }

    def _validated_template_fill_run(
        self,
        project_id: str,
        run: Mapping[str, Any],
    ) -> dict[str, Any]:
        context = self._context()
        output_key = str(run.get("outputStorageKey") or "")
        output_id = str(run.get("outputObjectId") or "")
        if not output_key or not output_id:
            raise LocalRuntimeError(
                409,
                "template_fill_receipt_invalid",
                "模板填充回执缺少本机输出对象",
            )
        output_path = self._managed_path(output_key)
        row = self.runtime.local_storage_object_get(
            sandbox_id=context.sandbox_id,
            object_id=output_id,
            storage_key=output_key,
        )
        if row is None or str(row["lifecycle_state"]) != "active":
            raise LocalRuntimeError(
                409,
                "template_fill_output_missing",
                "模板填充输出已从当前设备移除",
            )
        if str(row["media_type"]) != self.DOCX_MEDIA_TYPE:
            raise LocalRuntimeError(
                409,
                "template_fill_output_format_invalid",
                "模板填充输出格式无效",
            )
        try:
            data = output_path.read_bytes()
        except OSError as exc:
            raise LocalRuntimeError(
                409,
                "template_fill_output_missing",
                "模板填充输出文件缺失",
            ) from exc
        if (
            len(data) != int(row["byte_size"])
            or hashlib.sha256(data).hexdigest() != str(row["content_hash"])
        ):
            raise LocalRuntimeError(
                409,
                "template_fill_output_corrupt",
                "模板填充输出校验失败",
            )
        self._open_docx_bytes(
            data,
            code="template_fill_output_corrupt",
            message="模板填充输出结构损坏",
        )
        return {
            **dict(run),
            "clientId": project_id,
            "outputPath": str(output_path),
            "outputContentHash": str(row["content_hash"]),
            "outputStorageVersion": int(row["version"]),
        }

    def _existing_template_output(
        self,
        *,
        sandbox_id: str,
        object_id: str,
        storage_key: str,
    ) -> dict[str, Any] | None:
        row = self.runtime.local_storage_object_get(
            sandbox_id=sandbox_id,
            object_id=object_id,
        )
        if row is None:
            return None
        if (
            str(row.get("storage_key") or "") != storage_key
            or str(row.get("media_type") or "") != self.DOCX_MEDIA_TYPE
            or str(row.get("lifecycle_state") or "") != "active"
        ):
            raise LocalRuntimeError(
                409,
                "template_fill_output_identity_conflict",
                "本机模板输出对象与当前请求不一致",
            )
        path = self._managed_path(storage_key)
        try:
            data = path.read_bytes()
        except OSError as exc:
            raise LocalRuntimeError(
                409,
                "template_fill_output_missing",
                "本机模板输出文件缺失",
            ) from exc
        content_hash = hashlib.sha256(data).hexdigest()
        if (
            len(data) != int(row["byte_size"])
            or content_hash != str(row["content_hash"])
        ):
            raise LocalRuntimeError(
                409,
                "template_fill_output_corrupt",
                "本机模板输出校验失败",
            )
        self._open_docx_bytes(
            data,
            code="template_fill_output_corrupt",
            message="本机模板输出结构损坏",
        )
        return {
            "objectId": object_id,
            "storageKey": storage_key,
            "contentHash": content_hash,
            "mediaType": self.DOCX_MEDIA_TYPE,
            "byteSize": len(data),
            "version": int(row["version"]),
            "updatedAt": str(row["updated_at"]),
            "path": str(path),
        }

    def start_template_fill(
        self,
        project_id: str,
        *,
        template_path: str,
        values: Mapping[str, Any],
        idempotency_key: str | None = None,
    ) -> dict[str, Any]:
        raw_path = str(template_path or "").strip()
        try:
            source = Path(raw_path).expanduser().resolve(strict=True)
        except (OSError, RuntimeError) as exc:
            raise LocalRuntimeError(
                422,
                "template_file_missing",
                "模板文件不存在",
            ) from exc
        if not source.is_file():
            raise LocalRuntimeError(
                422,
                "template_file_missing",
                "模板文件不存在",
            )
        if source.suffix.casefold() != ".docx":
            raise LocalRuntimeError(
                415,
                "template_format_unsupported",
                "请选择 .docx Word 模板",
            )
        try:
            source_data = source.read_bytes()
        except OSError as exc:
            raise LocalRuntimeError(
                422,
                "template_file_unreadable",
                "模板文件无法读取",
            ) from exc
        document = self._open_docx_bytes(source_data)
        normalized_values = {
            self._normalize_docx_field_label(key): str(value or "").strip()
            for key, value in values.items()
            if self._normalize_docx_field_label(key)
        }
        labels = self._docx_template_labels(document)
        source_hash = hashlib.sha256(source_data).hexdigest()
        context = self._context()
        request_fingerprint = hashlib.sha256(
            canonical_json(
                {
                    "projectId": project_id,
                    "sourceContentHash": source_hash,
                    "values": normalized_values,
                }
            ).encode("utf-8")
        ).hexdigest()
        operation_identity = (
            f"explicit:{idempotency_key}"
            if str(idempotency_key or "").strip()
            else f"request:{request_fingerprint}"
        )
        operation_hash = hashlib.sha256(
            operation_identity.encode("utf-8")
        ).hexdigest()
        state_object_id = (
            f"project-state:{self._stable_segment(project_id)}"
        )
        with self.runtime.local_storage_object_lock(
            sandbox_id=context.sandbox_id,
            object_id=state_object_id,
        ):
            state = self._load_project_state(project_id)
            runs = dict(state.get("templateFillRuns") or {})
            for candidate in runs.values():
                if (
                    isinstance(candidate, Mapping)
                    and str(candidate.get("idempotencyKeyHash") or "")
                    == operation_hash
                ):
                    if (
                        str(candidate.get("requestFingerprint") or "")
                        != request_fingerprint
                    ):
                        raise LocalRuntimeError(
                            409,
                            "template_fill_idempotency_conflict",
                            "模板填充操作标识已用于不同请求",
                        )
                    return self._validated_template_fill_run(
                        project_id,
                        candidate,
                    )

            source_object = self._ensure_template_source_object(
                sandbox_id=context.sandbox_id,
                project_id=project_id,
                data=source_data,
                content_hash=source_hash,
            )
            self._apply_docx_template_values(
                document,
                normalized_values,
            )
            output_data = self._save_docx_bytes(document)
            run_id = f"template-fill-{operation_hash[:24]}"
            safe_stem = self._safe_name(source.stem)
            output_name = f"{safe_stem}-已填充-{operation_hash[:8]}.docx"
            output_key = (
                "local-template-fill/"
                f"{self._stable_segment(context.sandbox_id)}/"
                f"{self._stable_segment(project_id)}/outputs/{output_name}"
            )
            output_object_id = f"template-output:{operation_hash}"
            output = self._existing_template_output(
                sandbox_id=context.sandbox_id,
                object_id=output_object_id,
                storage_key=output_key,
            )
            if output is None:
                output = self._upsert_object(
                    sandbox_id=context.sandbox_id,
                    object_id=output_object_id,
                    storage_key=output_key,
                    media_type=self.DOCX_MEDIA_TYPE,
                    data=output_data,
                    expected_version=0,
                )
            fields = []
            for label in labels:
                value = normalized_values.get(label, "")
                status = "filled" if value else "missing"
                fields.append(
                    {
                        "label": label,
                        "value": value,
                        "status": status,
                        "evidenceTitles": [],
                        "fieldType": "general",
                        "valueKind": "summary" if value else "missing",
                        "confidence": 1.0 if value else 0.0,
                        "basisSummary": (
                            "由当前严格项目元数据填写"
                            if value
                            else "当前严格项目上下文没有同名字段"
                        ),
                        "followUpQuestion": (
                            None if value else f"请补充“{label}”"
                        ),
                        "reviewRequired": not bool(value),
                    }
                )
            now = utc_now()
            filled = sum(
                field["status"] == "filled" for field in fields
            )
            run = {
                "id": run_id,
                "clientId": project_id,
                "templateName": source.name,
                "templatePath": str(source),
                "status": "completed",
                "phase": "completed",
                "progress": 100,
                "stageLabel": "本机严格 Word 模板填充完成",
                "elapsedMs": 0,
                "fieldCount": len(fields),
                "processedCount": len(fields),
                "filledCount": filled,
                "missingCount": len(fields) - filled,
                "reviewFieldCount": len(fields) - filled,
                "currentFieldLabel": None,
                "evidenceTitles": [],
                "attachmentChecklist": self._docx_attachment_checklist(
                    document
                ),
                "fields": fields,
                "templateSourceObjectId": source_object["objectId"],
                "templateSourceStorageKey": source_object["storageKey"],
                "templateSourceContentHash": source_hash,
                "outputObjectId": output["objectId"],
                "outputStorageKey": output["storageKey"],
                "outputStorageVersion": output["version"],
                "outputContentHash": output["contentHash"],
                "outputPath": output["path"],
                "requestFingerprint": request_fingerprint,
                "idempotencyKeyHash": operation_hash,
                "sourceScope": "local_private",
                "persistedToOrganizationCloud": False,
                "errorMessage": None,
                "createdAt": now,
                "updatedAt": now,
            }
            runs[run_id] = run
            state["templateFillRuns"] = runs
            self._write_project_state(project_id, state)
            return self._validated_template_fill_run(project_id, run)

    def template_fill_run(self, project_id: str, run_id: str) -> dict[str, Any]:
        run = (
            self._load_project_state(project_id)
            .get("templateFillRuns", {})
            .get(run_id)
        )
        if not isinstance(run, dict):
            raise LocalRuntimeError(
                404,
                "template_fill_run_missing",
                "模板填充记录不存在",
            )
        return self._validated_template_fill_run(project_id, run)

    def render_report(
        self,
        *,
        report_id: str,
        report_version: int,
        title: str,
        markdown: str,
        output_format: str,
    ) -> dict[str, Any]:
        if output_format not in {"md", "docx"}:
            raise LocalRuntimeError(
                415,
                "report_render_format_unsupported",
                "当前本机严格渲染器支持 Markdown 和 Word；PDF 需先生成 Word 后由系统导出",
            )
        context = self._context()
        safe_title = self._safe_name(title or "报告")
        object_id = f"report-render:{report_id}:{report_version}:{output_format}"
        prefix = (
            "local-report-renders/"
            f"{self._stable_segment(context.sandbox_id)}/"
            f"{self._stable_segment(report_id)}/"
        )
        if output_format == "md":
            file_name = f"{safe_title}-v{report_version}.md"
            data = markdown.encode("utf-8")
            media_type = "text/markdown"
        else:
            file_name = f"{safe_title}-v{report_version}.docx"
            document_xml = (
                '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                '<w:document xmlns:w="http://schemas.openxmlformats.org/'
                'wordprocessingml/2006/main"><w:body>'
                + "".join(
                    (
                        "<w:p><w:r><w:t xml:space=\"preserve\">"
                        + escape(line)
                        + "</w:t></w:r></w:p>"
                    )
                    for line in markdown.splitlines()
                )
                + (
                    '<w:sectPr><w:pgSz w:w="11906" w:h="16838"/>'
                    '<w:pgMar w:top="1440" w:right="1440" w:bottom="1440" '
                    'w:left="1440" w:header="708" w:footer="708" w:gutter="0"/>'
                    "</w:sectPr></w:body></w:document>"
                )
            )
            content_types = (
                '<?xml version="1.0" encoding="UTF-8"?>'
                '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
                '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
                '<Default Extension="xml" ContentType="application/xml"/>'
                '<Override PartName="/word/document.xml" '
                'ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
                "</Types>"
            )
            relationships = (
                '<?xml version="1.0" encoding="UTF-8"?>'
                '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
                '<Relationship Id="rId1" '
                'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" '
                'Target="word/document.xml"/>'
                "</Relationships>"
            )
            target = self._managed_path(f"{prefix}.{object_id}.building")
            target.parent.mkdir(parents=True, exist_ok=True)
            try:
                with zipfile.ZipFile(
                    target,
                    mode="w",
                    compression=zipfile.ZIP_DEFLATED,
                ) as package:
                    package.writestr("[Content_Types].xml", content_types)
                    package.writestr("_rels/.rels", relationships)
                    package.writestr("word/document.xml", document_xml)
                data = target.read_bytes()
            finally:
                target.unlink(missing_ok=True)
            media_type = (
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document"
            )
        stored = self._upsert_object(
            sandbox_id=context.sandbox_id,
            object_id=object_id,
            storage_key=f"{prefix}{file_name}",
            media_type=media_type,
            data=data,
        )
        return {
            "artifact_id": report_id,
            "artifact_version": report_version,
            "format": output_format,
            "file_path": stored["path"],
            "file_name": file_name,
            "sourceScope": "current_device_managed_storage",
        }

    def _register_file(
        self,
        *,
        sandbox_id: str,
        project_id: str,
        source: Path,
    ) -> dict[str, Any]:
        if not source.is_file():
            raise LocalRuntimeError(
                422,
                "material_file_missing",
                f"文件不存在：{source}",
            )
        source_id = new_id()
        prefix = (
            "local-project-materials/"
            f"{self._stable_segment(sandbox_id)}/"
            f"{self._stable_segment(project_id)}/"
        )
        storage_key = f"{prefix}{source_id}-{self._safe_name(source.name)}"
        target = self._managed_path(storage_key)
        with self.runtime.local_storage_object_lock(
            sandbox_id=sandbox_id,
            object_id=source_id,
        ):
            target.parent.mkdir(parents=True, exist_ok=True)
            temporary = target.with_name(f".{target.name}.{new_id()}.tmp")
            digest = hashlib.sha256()
            byte_size = 0
            with (
                source.open("rb") as input_stream,
                temporary.open("wb") as output_stream,
            ):
                while True:
                    chunk = input_stream.read(1024 * 1024)
                    if not chunk:
                        break
                    digest.update(chunk)
                    byte_size += len(chunk)
                    output_stream.write(chunk)
            temporary.replace(target)
            media_type = (
                mimetypes.guess_type(source.name)[0]
                or "application/octet-stream"
            )
            stored = self.runtime.local_storage_object_put(
                sandbox_id=sandbox_id,
                object_id=source_id,
                storage_key=storage_key,
                content_hash=digest.hexdigest(),
                media_type=media_type,
                byte_size=byte_size,
                expected_version=0,
                original_path=str(source),
            )
        now = str(stored["updatedAt"])
        summary_kind = "metadata_only"
        summary = (
            f"本机资料《{source.name}》，媒体类型 {media_type}，"
            f"大小 {byte_size} 字节；当前仅登记元数据。"
        )
        if media_type.startswith("text/") or source.suffix.lower() in {
            ".md",
            ".json",
            ".csv",
            ".tsv",
        }:
            try:
                excerpt = target.read_text(encoding="utf-8")[:2000].strip()
            except (OSError, UnicodeDecodeError):
                excerpt = ""
            if excerpt:
                summary_kind = "text_excerpt"
                summary = excerpt
        sidecar_payload = {
            "schema": "yiyu.project-local-private-knowledge.v1",
            "sourceScope": "local_private",
            "projectId": project_id,
            "sourceId": source_id,
            "contentHash": digest.hexdigest(),
            "summary": summary,
            "summaryKind": summary_kind,
            "sourceDescription": "当前设备受管的本机私有项目资料",
            "updatedAt": now,
            "fileName": source.name,
        }
        sidecar_id = new_id()
        sidecar = self._upsert_object(
            sandbox_id=sandbox_id,
            object_id=sidecar_id,
            storage_key=f"{prefix}{source_id}.summary.json",
            media_type=self.SUMMARY_MEDIA_TYPE,
            data=canonical_json(sidecar_payload).encode("utf-8"),
        )
        return {
            "localSourceId": source_id,
            "localSummaryId": sidecar_id,
            "fileName": source.name,
            "managedPath": str(target),
            "originalSourcePath": str(source),
            "mediaType": media_type,
            "byteSize": byte_size,
            "contentHash": digest.hexdigest(),
            "summaryKind": summary_kind,
            "summaryPath": sidecar["path"],
            "updatedAt": now,
        }

    def import_paths(
        self,
        *,
        project_id: str,
        mode: str,
        paths: Iterable[Any],
        idempotency_key: str | None = None,
    ) -> dict[str, Any]:
        context = self._context()
        selected: list[Path] = []
        for raw in paths:
            candidate = Path(str(raw)).expanduser().resolve()
            if candidate.is_dir():
                selected.extend(
                    path for path in sorted(candidate.rglob("*")) if path.is_file()
                )
            else:
                selected.append(candidate)
        unique_files = list(dict.fromkeys(selected))
        if not unique_files:
            raise LocalRuntimeError(422, "materials_required", "请选择要导入的资料")
        if len(unique_files) > 1000:
            raise LocalRuntimeError(
                422,
                "material_limit_exceeded",
                "单次最多导入 1000 个文件",
            )
        normalized_mode = mode if mode in {"folder", "file"} else "file"
        operation_key = idempotency_key or new_id()
        request_fingerprint = hashlib.sha256(
            canonical_json(
                {
                    "projectId": project_id,
                    "mode": normalized_mode,
                    "paths": [str(path) for path in unique_files],
                }
            ).encode("utf-8")
        ).hexdigest()
        object_id, storage_key = self._import_operation_identity(
            context.sandbox_id,
            operation_key,
        )
        with self.runtime.local_storage_object_lock(
            sandbox_id=context.sandbox_id,
            object_id=object_id,
        ):
            receipt = self._load_import_operation(
                sandbox_id=context.sandbox_id,
                object_id=object_id,
                storage_key=storage_key,
                request_fingerprint=request_fingerprint,
            )
            if receipt is not None:
                return receipt
            materials = [
                self._register_file(
                    sandbox_id=context.sandbox_id,
                    project_id=project_id,
                    source=path,
                )
                for path in unique_files
            ]
            result = {
                "sandboxId": context.sandbox_id,
                "projectId": project_id,
                "mode": normalized_mode,
                "materials": materials,
            }
            self._write_import_operation(
                sandbox_id=context.sandbox_id,
                object_id=object_id,
                storage_key=storage_key,
                project_id=project_id,
                operation_kind="paths",
                request_fingerprint=request_fingerprint,
                result=result,
            )
            return result

    def import_text(
        self,
        *,
        project_id: str,
        title: str,
        content: str,
        idempotency_key: str | None = None,
    ) -> dict[str, Any]:
        normalized_content = content.strip()
        if not normalized_content:
            raise LocalRuntimeError(
                422,
                "document_content_required",
                "请输入文档内容",
            )
        data = normalized_content.encode("utf-8")
        if len(data) > 5 * 1024 * 1024:
            raise LocalRuntimeError(
                422,
                "document_content_too_large",
                "文本资料不能超过 5MB",
            )
        context = self._context()
        normalized_title = title.strip() or "新建文档"
        operation_key = idempotency_key or new_id()
        request_fingerprint = hashlib.sha256(
            canonical_json(
                {
                    "projectId": project_id,
                    "title": normalized_title,
                    "contentHash": hashlib.sha256(data).hexdigest(),
                }
            ).encode("utf-8")
        ).hexdigest()
        receipt_id, receipt_key = self._import_operation_identity(
            context.sandbox_id,
            operation_key,
        )
        with self.runtime.local_storage_object_lock(
            sandbox_id=context.sandbox_id,
            object_id=receipt_id,
        ):
            receipt = self._load_import_operation(
                sandbox_id=context.sandbox_id,
                object_id=receipt_id,
                storage_key=receipt_key,
                request_fingerprint=request_fingerprint,
            )
            if receipt is not None:
                return receipt
            result = self._prepare_text_import(
                sandbox_id=context.sandbox_id,
                project_id=project_id,
                normalized_title=normalized_title,
                data=data,
                normalized_content=normalized_content,
            )
            self._write_import_operation(
                sandbox_id=context.sandbox_id,
                object_id=receipt_id,
                storage_key=receipt_key,
                project_id=project_id,
                operation_kind="text",
                request_fingerprint=request_fingerprint,
                result=result,
            )
            return result

    def _prepare_text_import(
        self,
        *,
        sandbox_id: str,
        project_id: str,
        normalized_title: str,
        data: bytes,
        normalized_content: str,
    ) -> dict[str, Any]:
        source_id = new_id()
        file_name = self._safe_name(normalized_title)
        if not file_name.lower().endswith((".md", ".txt")):
            file_name = f"{file_name}.md"
        prefix = (
            "local-project-materials/"
            f"{self._stable_segment(sandbox_id)}/"
            f"{self._stable_segment(project_id)}/"
        )
        source = self._upsert_object(
            sandbox_id=sandbox_id,
            object_id=source_id,
            storage_key=f"{prefix}{source_id}-{file_name}",
            media_type="text/markdown",
            data=data,
        )
        now = str(source["updatedAt"])
        summary_id = new_id()
        summary_payload = {
            "schema": "yiyu.project-local-private-knowledge.v1",
            "sourceScope": "local_private",
            "projectId": project_id,
            "sourceId": source_id,
            "contentHash": source["contentHash"],
            "summary": normalized_content[:2000],
            "summaryKind": "text_excerpt",
            "sourceDescription": "当前设备受管的本机私有项目文本资料",
            "updatedAt": now,
            "fileName": file_name,
        }
        summary = self._upsert_object(
            sandbox_id=sandbox_id,
            object_id=summary_id,
            storage_key=f"{prefix}{source_id}.summary.json",
            media_type=self.SUMMARY_MEDIA_TYPE,
            data=canonical_json(summary_payload).encode("utf-8"),
        )
        return {
            "localSourceId": source_id,
            "localSummaryId": summary_id,
            "fileName": file_name,
            "title": normalized_title,
            "managedPath": source["path"],
            "originalSourcePath": None,
            "mediaType": "text/markdown",
            "byteSize": len(data),
            "contentHash": source["contentHash"],
            "summaryKind": "text_excerpt",
            "summaryPath": summary["path"],
            "updatedAt": now,
        }

    def _session_key(self, sandbox_id: str, session_id: str) -> str:
        return (
            "smart-import/"
            f"{self._stable_segment(sandbox_id)}/sessions/{session_id}.json"
        )

    def _write_session(
        self,
        sandbox_id: str,
        state: Mapping[str, Any],
    ) -> dict[str, Any]:
        session_id = str((state.get("session") or {}).get("id") or "")
        if not session_id:
            raise LocalRuntimeError(
                422,
                "smart_import_session_invalid",
                "导入会话无效",
            )
        expected_version = int(state.get("_localStorageVersion") or 0)
        payload = {
            key: value
            for key, value in dict(state).items()
            if not str(key).startswith("_")
        }
        self._upsert_object(
            sandbox_id=sandbox_id,
            object_id=session_id,
            storage_key=self._session_key(sandbox_id, session_id),
            media_type=self.SMART_SESSION_MEDIA_TYPE,
            data=canonical_json(payload).encode("utf-8"),
            expected_version=expected_version,
        )
        payload["_localStorageVersion"] = expected_version + 1
        payload["_localSandboxId"] = sandbox_id
        return payload

    def _load_session(
        self,
        sandbox_id: str,
        session_id: str,
    ) -> dict[str, Any]:
        storage_key = self._session_key(sandbox_id, session_id)
        row = self.runtime.local_storage_object_get(
            sandbox_id=sandbox_id,
            object_id=session_id,
            storage_key=storage_key,
        )
        if row is not None and row["sandbox_id"] != sandbox_id:
            raise LocalRuntimeError(
                409,
                "local_storage_sandbox_changed",
                "本机工作空间已切换，请重试",
            )
        if row is None or row["lifecycle_state"] != "active":
            raise LocalRuntimeError(
                404,
                "smart_import_session_missing",
                "智能导入会话不存在",
            )
        data = self._managed_path(storage_key).read_bytes()
        if (
            len(data) != int(row["byte_size"])
            or hashlib.sha256(data).hexdigest() != str(row["content_hash"])
        ):
            raise LocalRuntimeError(
                409,
                "smart_import_session_corrupt",
                "智能导入会话文件校验失败",
            )
        payload = json.loads(data.decode("utf-8"))
        if not isinstance(payload, dict):
            raise LocalRuntimeError(
                409,
                "smart_import_session_corrupt",
                "智能导入会话无效",
            )
        payload["_localStorageVersion"] = int(row["version"])
        payload["_localSandboxId"] = sandbox_id
        return payload

    def create_smart_session(self, payload: Mapping[str, Any]) -> dict[str, Any]:
        context = self._context()
        now = utc_now()
        session_id = new_id()
        state = {
            "session": {
                "id": session_id,
                "client_id": payload.get("clientId"),
                "project_event_line_id": payload.get("projectEventLineId"),
                "narrator_user_id": context.membership_id,
                "title": str(payload.get("title") or "智能导入"),
                "status": "drafting",
                "total_chunks": 0,
                "total_files": 0,
                "created_at": now,
                "updated_at": now,
                "imported_at": None,
            },
            "chunks": [],
            "staged_files": [],
        }
        return self._write_session(context.sandbox_id, state)

    def get_smart_session(self, session_id: str) -> dict[str, Any]:
        context = self._context()
        return self._load_session(context.sandbox_id, session_id)

    def update_smart_session(
        self,
        session_id: str,
        payload: Mapping[str, Any],
    ) -> dict[str, Any]:
        context = self._context()
        state = self._load_session(context.sandbox_id, session_id)
        session = state["session"]
        for source, target in (
            ("clientId", "client_id"),
            ("projectEventLineId", "project_event_line_id"),
            ("title", "title"),
        ):
            if source in payload:
                session[target] = payload.get(source)
        session["updated_at"] = utc_now()
        return self._write_session(context.sandbox_id, state)

    def discard_smart_session(self, session_id: str) -> dict[str, Any]:
        context = self._context()
        state = self._load_session(context.sandbox_id, session_id)
        state["session"]["status"] = "discarded"
        state["session"]["updated_at"] = utc_now()
        self._write_session(context.sandbox_id, state)
        return {"ok": True}

    def upload_smart_file(self, session_id: str, upload: Any) -> dict[str, Any]:
        context = self._context()
        state = self._load_session(context.sandbox_id, session_id)
        file_id = new_id()
        file_name = self._safe_name(
            str(getattr(upload, "filename", "") or "file")
        )
        stream = getattr(upload, "file", upload)
        if hasattr(stream, "seek"):
            stream.seek(0)
        data = stream.read()
        if not isinstance(data, bytes):
            raise LocalRuntimeError(
                422,
                "smart_import_file_invalid",
                "上传文件无效",
            )
        media_type = str(
            getattr(upload, "content_type", "")
            or mimetypes.guess_type(file_name)[0]
            or "application/octet-stream"
        )
        storage_key = (
            "smart-import/"
            f"{self._stable_segment(context.sandbox_id)}/files/"
            f"{file_id}-{file_name}"
        )
        stored = self._upsert_object(
            sandbox_id=context.sandbox_id,
            object_id=file_id,
            storage_key=storage_key,
            media_type=media_type,
            data=data,
        )
        record = {
            "id": file_id,
            "session_id": session_id,
            "original_filename": file_name,
            "storage_path": stored["path"],
            "size_bytes": len(data),
            "mime_type": media_type,
            "assigned_chunk_id": None,
            "role_override": None,
            "document_id": None,
            "document_inserted_at": None,
            "upload_at": stored["updatedAt"],
            "_content_hash": stored["contentHash"],
            "_storage_key": storage_key,
        }
        state["staged_files"].append(record)
        state["session"]["total_files"] = len(state["staged_files"])
        state["session"]["updated_at"] = utc_now()
        self._write_session(context.sandbox_id, state)
        return {
            key: value
            for key, value in record.items()
            if not key.startswith("_")
        }

    def _all_sessions(self, sandbox_id: str) -> list[dict[str, Any]]:
        rows = self.runtime.local_storage_objects_by_media_type(
            media_type=self.SMART_SESSION_MEDIA_TYPE,
        )
        result = []
        for row in rows:
            if row["sandbox_id"] != sandbox_id:
                continue
            try:
                result.append(
                    self._load_session(sandbox_id, str(row["object_id"]))
                )
            except (OSError, ValueError, LocalRuntimeError):
                continue
        return result

    def _find_child(
        self,
        *,
        child_key: str,
        child_id: str,
    ) -> tuple[str, dict[str, Any], dict[str, Any]]:
        context = self._context()
        for state in self._all_sessions(context.sandbox_id):
            for item in state.get(child_key) or []:
                if str(item.get("id")) == child_id:
                    return context.sandbox_id, state, item
        raise LocalRuntimeError(
            404,
            "smart_import_item_missing",
            "智能导入条目不存在",
        )

    def delete_smart_file(self, file_id: str) -> dict[str, Any]:
        sandbox_id, state, _ = self._find_child(
            child_key="staged_files",
            child_id=file_id,
        )
        state["staged_files"] = [
            item for item in state["staged_files"] if item["id"] != file_id
        ]
        state["session"]["total_files"] = len(state["staged_files"])
        state["session"]["updated_at"] = utc_now()
        lifecycle = self.runtime.local_storage_object_set_lifecycle(
            object_id=file_id,
            lifecycle_state="deleted",
        )
        if lifecycle["sandboxId"] != sandbox_id:
            raise LocalRuntimeError(
                409,
                "local_storage_sandbox_changed",
                "本机工作空间已切换，请重试",
            )
        self._write_session(sandbox_id, state)
        return {"ok": True}

    def assign_smart_file(
        self,
        file_id: str,
        chunk_id: str | None,
    ) -> dict[str, Any]:
        sandbox_id, state, item = self._find_child(
            child_key="staged_files",
            child_id=file_id,
        )
        if chunk_id and not any(
            chunk["id"] == chunk_id for chunk in state["chunks"]
        ):
            raise LocalRuntimeError(
                404,
                "smart_import_chunk_missing",
                "导入分块不存在",
            )
        item["assigned_chunk_id"] = chunk_id
        state["session"]["updated_at"] = utc_now()
        self._write_session(sandbox_id, state)
        return {"ok": True}

    def add_smart_chunk(
        self,
        session_id: str,
        payload: Mapping[str, Any],
    ) -> dict[str, Any]:
        context = self._context()
        state = self._load_session(context.sandbox_id, session_id)
        now = utc_now()
        chunk = {
            "id": new_id(),
            "session_id": session_id,
            "sequence": len(state["chunks"]) + 1,
            "raw_text": str(payload.get("rawText") or ""),
            "parsed_json": "{}",
            "parsed": {},
            "parse_status": "pending",
            "parse_error": "",
            "user_edited_parsed": 0,
            "created_at": now,
            "updated_at": now,
        }
        state["chunks"].append(chunk)
        for file_id in payload.get("fileIds") or []:
            for item in state["staged_files"]:
                if item["id"] == file_id:
                    item["assigned_chunk_id"] = chunk["id"]
        state["session"]["total_chunks"] = len(state["chunks"])
        state["session"]["updated_at"] = now
        self._write_session(context.sandbox_id, state)
        if bool(payload.get("autoParse")):
            self.parse_smart_chunk(chunk["id"])
            return self._load_session(context.sandbox_id, session_id)
        return state

    def update_smart_chunk(
        self,
        chunk_id: str,
        payload: Mapping[str, Any],
    ) -> dict[str, Any]:
        sandbox_id, state, chunk = self._find_child(
            child_key="chunks",
            child_id=chunk_id,
        )
        chunk["raw_text"] = str(payload.get("rawText") or "")
        chunk["parse_status"] = "pending"
        chunk["parse_error"] = ""
        chunk["updated_at"] = utc_now()
        state["session"]["updated_at"] = chunk["updated_at"]
        self._write_session(sandbox_id, state)
        if bool(payload.get("autoParse")):
            self.parse_smart_chunk(chunk_id)
            return self._load_session(sandbox_id, state["session"]["id"])
        return state

    def delete_smart_chunk(self, chunk_id: str) -> dict[str, Any]:
        sandbox_id, state, _ = self._find_child(
            child_key="chunks",
            child_id=chunk_id,
        )
        state["chunks"] = [
            chunk for chunk in state["chunks"] if chunk["id"] != chunk_id
        ]
        for sequence, chunk in enumerate(state["chunks"], start=1):
            chunk["sequence"] = sequence
        for item in state["staged_files"]:
            if item.get("assigned_chunk_id") == chunk_id:
                item["assigned_chunk_id"] = None
        state["session"]["total_chunks"] = len(state["chunks"])
        state["session"]["updated_at"] = utc_now()
        self._write_session(sandbox_id, state)
        return {"ok": True}

    @staticmethod
    def _parse_text(raw_text: str) -> dict[str, Any]:
        lines = [
            line.strip(" \t-*•")
            for line in raw_text.splitlines()
            if line.strip(" \t-*•")
        ]
        entity_pattern = re.compile(
            r"[\u4e00-\u9fffA-Za-z0-9]{2,24}(?:基金会|公司|机构|中心|项目|团队)"
        )
        entities = [
            {"name": value, "kind": "organization"}
            for value in sorted(set(entity_pattern.findall(raw_text)))
        ]
        events = [{"summary": line[:300]} for line in lines]
        commitments = [
            {"content": line[:300], "status": "candidate"}
            for line in lines
            if any(token in line for token in ("承诺", "负责", "将于", "计划"))
        ]
        risks = [
            {
                "title": line[:80],
                "severity": "medium",
                "description": line[:300],
                "signal_kind": "text_rule",
            }
            for line in lines
            if any(token in line for token in ("风险", "困难", "阻碍", "延期"))
        ]
        return {
            "entities": entities,
            "relationships": [],
            "events": events,
            "opinions": [],
            "files_classified": [],
            "files_suggested_to_attach": [],
            "commitments": commitments,
            "risk_signals": risks,
            "open_questions": (
                [] if lines else ["该分块没有可解析的文字，请补充内容"]
            ),
        }

    def _smart_file_text(
        self,
        *,
        sandbox_id: str,
        staged_file: Mapping[str, Any],
    ) -> str:
        file_id = str(staged_file.get("id") or "")
        row = self.runtime.local_storage_object_get(
            sandbox_id=sandbox_id,
            object_id=file_id,
        )
        if row is None or str(row["lifecycle_state"]) != "active":
            raise LocalRuntimeError(
                404,
                "smart_import_file_missing",
                "智能导入文件已从当前设备移除，请重新选择",
            )
        path = self._managed_path(str(row["storage_key"]))
        try:
            data = path.read_bytes()
        except OSError as exc:
            raise LocalRuntimeError(
                404,
                "smart_import_file_missing",
                "智能导入文件无法从当前设备读取，请重新选择",
            ) from exc
        if (
            len(data) != int(row["byte_size"])
            or hashlib.sha256(data).hexdigest() != str(row["content_hash"])
        ):
            raise LocalRuntimeError(
                409,
                "smart_import_file_corrupt",
                "智能导入文件校验失败，请重新选择",
            )
        media_type = str(
            staged_file.get("mime_type")
            or row["media_type"]
            or ""
        ).lower()
        suffix = Path(
            str(staged_file.get("original_filename") or path.name)
        ).suffix.lower()
        if suffix == ".docx":
            return self._docx_text(path).strip()
        if media_type.startswith("text/") or suffix in {
            ".md",
            ".txt",
            ".json",
            ".csv",
            ".tsv",
            ".xml",
            ".html",
        }:
            try:
                return data.decode("utf-8-sig").strip()
            except UnicodeDecodeError as exc:
                raise LocalRuntimeError(
                    415,
                    "smart_import_file_encoding_unsupported",
                    "智能导入文件不是 UTF-8 文本，请转换编码后重试",
                ) from exc
        raise LocalRuntimeError(
            415,
            "smart_import_file_content_unsupported",
            "当前文件格式尚不能提取正文；请改用文本、Markdown 或 Word 文档",
        )

    def parse_smart_chunk(self, chunk_id: str) -> dict[str, Any]:
        sandbox_id, state, chunk = self._find_child(
            child_key="chunks",
            child_id=chunk_id,
        )
        assigned_ids = {
            str(value)
            for value in chunk.get("file_ids") or []
            if str(value)
        }
        assigned_files = [
            item
            for item in state.get("staged_files") or []
            if str(item.get("id") or "") in assigned_ids
            or str(item.get("assigned_chunk_id") or "") == chunk_id
        ]
        try:
            sections = [str(chunk.get("raw_text") or "").strip()]
            for item in assigned_files:
                file_text = self._smart_file_text(
                    sandbox_id=sandbox_id,
                    staged_file=item,
                )
                if file_text:
                    sections.append(
                        f"[文件：{item.get('original_filename') or '未命名资料'}]\n"
                        f"{file_text}"
                    )
            source_text = "\n\n".join(
                section for section in sections if section
            ).strip()
            if not source_text:
                raise LocalRuntimeError(
                    422,
                    "smart_import_content_required",
                    "该分块没有可解析的文字或文件正文",
                )
            parsed = self._parse_text(source_text)
        except LocalRuntimeError as exc:
            chunk["parse_status"] = "failed"
            chunk["parse_error"] = exc.message
            chunk["updated_at"] = utc_now()
            state["session"]["status"] = "draft"
            state["session"]["updated_at"] = chunk["updated_at"]
            self._write_session(sandbox_id, state)
            raise
        chunk["parsed"] = parsed
        chunk["parsed_json"] = canonical_json(parsed)
        chunk["parse_status"] = "parsed"
        chunk["parse_error"] = ""
        chunk["updated_at"] = utc_now()
        state["session"]["status"] = "ready_for_review"
        state["session"]["updated_at"] = chunk["updated_at"]
        self._write_session(sandbox_id, state)
        return {"ok": True, "parsed": parsed}

    def patch_smart_chunk(
        self,
        chunk_id: str,
        parsed: Mapping[str, Any],
    ) -> dict[str, Any]:
        sandbox_id, state, chunk = self._find_child(
            child_key="chunks",
            child_id=chunk_id,
        )
        chunk["parsed"] = dict(parsed)
        chunk["parsed_json"] = canonical_json(dict(parsed))
        chunk["parse_status"] = "parsed"
        chunk["user_edited_parsed"] = 1
        chunk["updated_at"] = utc_now()
        state["session"]["status"] = "ready_for_review"
        state["session"]["updated_at"] = chunk["updated_at"]
        return self._write_session(sandbox_id, state)

    def smart_preview(self, session_id: str) -> dict[str, Any]:
        state = self.get_smart_session(session_id)
        parsed_chunks = [
            chunk
            for chunk in state["chunks"]
            if chunk["parse_status"] == "parsed"
        ]
        fields = (
            "entities",
            "relationships",
            "events",
            "opinions",
            "commitments",
            "risk_signals",
            "files_classified",
            "files_suggested_to_attach",
            "open_questions",
        )
        aggregated = {
            field: [
                item
                for chunk in parsed_chunks
                for item in (chunk.get("parsed") or {}).get(field, [])
            ]
            for field in fields
        }
        for staged in state["staged_files"]:
            aggregated["files_classified"].append(
                {
                    "original_filename": staged["original_filename"],
                    "role": staged.get("role_override") or "source_material",
                    "confidence": 1.0,
                }
            )
        return {
            "session_id": session_id,
            "chunks_total": len(state["chunks"]),
            "chunks_parsed": len(parsed_chunks),
            "chunks_failed": [
                {
                    "chunk_id": chunk["id"],
                    "sequence": chunk["sequence"],
                }
                for chunk in state["chunks"]
                if chunk["parse_status"] == "failed"
            ],
            **aggregated,
        }

    def mark_smart_imported(
        self,
        session_id: str,
        *,
        document_ids: Mapping[str, str] | None = None,
    ) -> dict[str, Any]:
        context = self._context()
        state = self._load_session(context.sandbox_id, session_id)
        now = utc_now()
        state["session"]["status"] = "imported"
        state["session"]["imported_at"] = now
        state["session"]["updated_at"] = now
        mapping = dict(document_ids or {})
        for item in state["staged_files"]:
            document_id = mapping.get(str(item["id"]))
            if document_id:
                item["document_id"] = document_id
                item["document_inserted_at"] = now
        return self._write_session(context.sandbox_id, state)


def select_relevant_excerpt(
    content: str,
    query: str,
    *,
    max_chars: int = 40_000,
) -> str:
    """Stable helper for handlers whose repository class is replaced in tests."""
    return LocalProjectMaterialsRepository.select_relevant_excerpt(
        content,
        query,
        max_chars=max_chars,
    )
