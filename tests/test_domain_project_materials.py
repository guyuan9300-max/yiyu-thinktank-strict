from __future__ import annotations

import json
from io import BytesIO
from pathlib import Path
from types import SimpleNamespace
from typing import Any, Mapping

import pytest
from cryptography.fernet import Fernet
from docx import Document
from fastapi.testclient import TestClient

from backend.app import link_material_fetcher
from backend.app.project_materials_local import (
    LocalProjectMaterialsRepository,
    _current_project_document_projection,
)
from backend.app.runtime import LocalRuntimeError, WorkspaceRuntime
from backend.app.secret_store import MemorySecretStore
from backend.app.ui_domains import project_materials as project_materials_ui
from backend.app.ui_domains.project_materials import (
    _UNSUPPORTED_ROUTE_SPECS,
    router,
)
from backend.app.ui_domains.registry import build_default_registry
from backend.app.ui_domains.routing import NOT_HANDLED, UiRequest
from cloud_backend.app.config import CloudConfig
from cloud_backend.app.main import create_app
from cloud_backend.app.repositories.project_materials import (
    GC07ProjectMaterialsRepository,
)
from cloud_backend.app.repository import CloudRepository
from strict_common.schema import runtime_connection
from tests.test_gc01_authorization import _auth, _seed_gc01_cloud


def _cloud_client(tmp_path: Path) -> tuple[TestClient, Path]:
    database = tmp_path / "strict-cloud.db"
    config = CloudConfig(
        data_dir=tmp_path,
        database_path=database,
        bootstrap_token="bootstrap-test",
        master_key=Fernet.generate_key().decode(),
        cloud_instance_id=None,
    )
    return TestClient(create_app(config)), database


def _bootstrap(client: TestClient) -> tuple[dict[str, Any], dict[str, str]]:
    response = client.post(
        "/api/v2/auth/bootstrap-organization",
        json={
            "organizationName": "项目资料测试组织",
            "displayName": "管理员",
            "email": "project-materials@example.com",
            "password": "12345678",
            "bootstrapToken": "bootstrap-test",
        },
    )
    assert response.status_code == 201, response.text
    payload = response.json()
    return payload, {"Authorization": f"Bearer {payload['accessToken']}"}


def _default_project(
    client: TestClient,
    headers: dict[str, str],
) -> dict[str, Any]:
    response = client.get(
        "/api/v2/domain/project-materials/projects",
        headers=headers,
    )
    assert response.status_code == 200, response.text
    return next(
        item
        for item in response.json()["projects"]
        if item["isDefaultInternalProject"]
    )


def test_document_ai_action_consumes_real_material_thread_style_and_agent_skill(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    captured: dict[str, Any] = {}

    class FakeStore:
        def document_text(self, document_id: str) -> dict[str, Any]:
            assert document_id == "doc_phase18"
            return {
                "projectId": "client_rici",
                "title": "日慈环节18资料",
                "content": "环节18唯一资料哨兵：心灵魔法学院教师培训。",
                "contentHash": "hash_phase18",
            }

    class FakeRuntime:
        def cloud_query(self, path: str, query: Any = None) -> Any:
            if path == "/api/v2/workbench/libraries/writing_skill":
                return [{"id": "style_clear", "name": "清晰表达", "distilledMd": "先说结论，再给依据。"}]
            if path == "/api/v2/agent-skills/skill_evidence":
                return {
                    "skillId": "skill_evidence",
                    "shortName": "证据—判断—边界",
                    "description": "区分事实与判断",
                    "instructions": ["先列证据", "再给判断", "标明信息缺口"],
                    "outputTemplate": "证据 / 判断 / 边界",
                    "agentKinds": ["project_workspace"],
                    "enabled": True,
                    "version": 3,
                    "contentHash": "skill_hash",
                }
            raise AssertionError(path)

        def project_agent_skill(self, item: Mapping[str, Any]) -> None:
            assert item["skillId"] == "skill_evidence"

        def workbench_chat_history(self, client_id: str, thread_id: str) -> list[dict[str, Any]]:
            assert (client_id, thread_id) == ("client_rici", "thread_phase18")
            return [{"question": "此前讨论什么？", "answerMarkdown": "讨论教师培训项目。"}]

        def private_ai_completion(self, **kwargs: Any) -> dict[str, Any]:
            captured.update(kwargs)
            return {"content": "环节18生成草稿", "sourceScope": "member_local_private_request"}

    monkeypatch.setattr(project_materials_ui, "_local_store", lambda _compatibility: FakeStore())
    response = router.dispatch(
        SimpleNamespace(runtime=FakeRuntime()),
        UiRequest(
            method="POST",
            path="clients/client_rici/documents/ai-action",
            query={},
            body={
                "content": "待编辑正文",
                "action": "rewrite_pro",
                "userRequest": "改成项目报告语气",
                "workingDocumentIds": ["doc_phase18"],
                "activeSkillId": "style_clear",
                "activeSkillIds": ["skill_evidence"],
                "threadId": "thread_phase18",
            },
            idempotency_key="phase18-document-ai",
        ),
    )

    assert response["content"] == "环节18生成草稿"
    assert response["targetScope"] == "cursor_insert"
    assert response["persistedToOrganizationCloud"] is False
    assert response["activeSkillIds"] == ["skill_evidence"]
    assert {item["type"] for item in response["sources"]} == {
        "member_local_document",
        "writing_skill",
        "agent_skill",
        "workbench_thread_context",
    }
    assert "环节18唯一资料哨兵" in captured["prompt"]
    assert "此前讨论什么" in captured["prompt"]
    assert "先说结论" in captured["system_prompt"]
    assert "证据—判断—边界" in captured["system_prompt"]
    assert "标明信息缺口" in captured["system_prompt"]
    assert captured["read_timeout_seconds"] == 100.0


def test_document_ai_action_is_present_in_strict_runtime_registry() -> None:
    assert any(
        route.method == "POST"
        and route.pattern
        == r"clients/(?P<project_id>[^/]+)/documents/ai-action"
        for route in build_default_registry().routes
    )


def test_visible_member_can_register_own_local_material_without_project_edit_rights(
    tmp_path: Path,
) -> None:
    client, database = _cloud_client(tmp_path)
    with client:
        admin, admin_auth = _bootstrap(client)
        project = _default_project(client, admin_auth)
        project_id = project["projectId"]

        department = client.post(
            "/api/v2/organization/departments",
            headers={
                **admin_auth,
                "Idempotency-Key": "local-material-member-department",
            },
            json={"name": "资料协作部", "expectedOrganizationVersion": 1},
        )
        assert department.status_code == 201, department.text
        invite = client.post(
            "/api/v2/organization/invites",
            headers=admin_auth,
            json={
                "inviteKind": "department",
                "targetId": department.json()["id"],
            },
        )
        assert invite.status_code == 201, invite.text
        joined = client.post(
            "/api/v2/auth/join",
            json={
                "inviteCode": invite.json()["inviteCode"],
                "displayName": "资料协作成员",
                "email": "local-material-member@example.com",
                "password": "member-password",
            },
        )
        assert joined.status_code == 201, joined.text
        member = joined.json()
        member_auth = {
            "Authorization": f"Bearer {member['accessToken']}",
        }

        with runtime_connection(database, "cloud") as connection:
            connection.execute(
                """
                UPDATE organization_memberships
                SET visibility_scope = 'organization'
                WHERE membership_id = ?
                """,
                (member["membershipId"],),
            )
            connection.commit()

        visible = client.get(
            f"/api/v2/domain/project-materials/projects/{project_id}",
            headers=member_auth,
        )
        assert visible.status_code == 200, visible.text

        forbidden_update = client.put(
            f"/api/v2/domain/project-materials/projects/{project_id}",
            headers={
                **member_auth,
                "Idempotency-Key": "member-project-update-forbidden",
            },
            json={
                "name": "成员不得修改项目",
                "expectedVersion": project["version"],
            },
        )
        assert forbidden_update.status_code == 403, forbidden_update.text
        assert forbidden_update.json()["error"]["code"] == "project_forbidden"

        registered = client.post(
            f"/api/v2/domain/project-materials/projects/{project_id}"
            "/materials/register-metadata",
            headers={
                **member_auth,
                "Idempotency-Key": "member-local-material-register",
            },
            json={
                "materials": [
                    {
                        "localSourceId": "member-local-source",
                        "fileName": "成员本机资料.txt",
                        "contentHash": "a" * 64,
                        "byteSize": 42,
                        "mediaType": "text/plain",
                        "sourceKind": "local_private_metadata",
                    }
                ]
            },
        )
        assert registered.status_code == 201, registered.text
        assert registered.json()["materialBoundary"] == {
            "sourceFileContentUploaded": False,
            "sourceFilePathUploaded": False,
            "localSummaryUploaded": False,
        }
        document_id = registered.json()["documents"][0]["documentId"]

    with runtime_connection(database, "cloud", read_only=True) as connection:
        row = connection.execute(
            """
            SELECT d.owner_membership_id, d.visibility_scope, d.document_kind,
                   d.current_version, a.storage_object_id, a.source_locator
            FROM knowledge_documents d
            JOIN source_assets a ON a.source_asset_id = d.source_asset_id
            WHERE d.document_id = ?
            """,
            (document_id,),
        ).fetchone()
    assert row is not None
    assert row["owner_membership_id"] == member["membershipId"]
    assert row["visibility_scope"] == "self"
    assert row["document_kind"] == "local_private_metadata"
    assert int(row["current_version"]) == 0
    assert row["storage_object_id"] is None
    assert row["source_locator"] == ""


def test_mobile_recording_publishes_only_safe_summary_into_existing_88_tables(
    tmp_path: Path,
) -> None:
    database = tmp_path / "strict-cloud.db"
    config, tokens = _seed_gc01_cloud(database)
    with TestClient(create_app(config)) as client:
        auth = _auth(tokens["admin"])
        created = client.post(
            "/api/v2/domain/project-materials/projects",
            headers={**auth, "Idempotency-Key": "mobile-recording-project"},
            json={"name": "移动录音测试项目"},
        )
        assert created.status_code == 201, created.text
        project_id = created.json()["project"]["projectId"]
        response = client.post(
            f"/api/v2/domain/project-materials/projects/{project_id}/mobile-recording-summary",
            headers={**auth, "Idempotency-Key": "mobile-recording-summary-once"},
            json={
                "title": "移动现场记录-知识摘要",
                "safeSummaryMarkdown": "## 决定\n- 由项目组在下周完成教师培训方案。",
                "facts": ["项目组将在下周完成教师培训方案"],
                "sourceContentHash": "b" * 64,
            },
        )
        assert response.status_code == 201, response.text
        result = response.json()
        assert result["localOriginalUploaded"] is False
        assert result["fullTranscriptUploaded"] is False

        rejected = client.post(
            f"/api/v2/domain/project-materials/projects/{project_id}/mobile-recording-summary",
            headers={**auth, "Idempotency-Key": "mobile-recording-body-rejected"},
            json={
                "title": "越界请求",
                "safeSummaryMarkdown": "摘要",
                "sourceContentHash": "c" * 64,
                "transcriptText": "逐字稿不得进入组织云",
            },
        )
        assert rejected.status_code == 422

    with runtime_connection(database, "cloud", read_only=True) as connection:
        assert connection.execute("SELECT COUNT(*) FROM sqlite_master WHERE type='table' AND name NOT LIKE 'sqlite_%'").fetchone()[0] == 88
        document = connection.execute(
            "SELECT document_kind,publication_state FROM knowledge_documents WHERE id=?",
            (result["documentId"],),
        ).fetchone()
        manifest = connection.execute(
            "SELECT receipt FROM object_manifests WHERE id=(SELECT object_manifest_id FROM document_versions WHERE id=?)",
            (result["documentVersionId"],),
        ).fetchone()
        assert document["document_kind"] == "shared_summary"
        assert document["publication_state"] == "published"
        receipt = json.loads(manifest["receipt"])
        assert receipt["localOriginalUploaded"] is False
        assert receipt["fullTranscriptUploaded"] is False
        assert "逐字稿不得进入组织云" not in manifest["receipt"]


def test_project_crud_lifecycle_cas_idempotency_and_receipts(
    tmp_path: Path,
) -> None:
    client, database = _cloud_client(tmp_path)
    with client:
        _, auth = _bootstrap(client)
        create_headers = {**auth, "Idempotency-Key": "project-create-1"}
        created = client.post(
            "/api/v2/domain/project-materials/projects",
            headers=create_headers,
            json={
                "name": "乡村儿童支持",
                "alias": "儿童项目",
                "summary": "项目简介",
                "domain": "公益项目",
                "color": "#123456",
                "participantMembershipIds": [],
            },
        )
        assert created.status_code == 201, created.text
        repeated = client.post(
            "/api/v2/domain/project-materials/projects",
            headers=create_headers,
            json={
                "name": "乡村儿童支持",
                "alias": "儿童项目",
                "summary": "项目简介",
                "domain": "公益项目",
                "color": "#123456",
                "participantMembershipIds": [],
            },
        )
        assert repeated.status_code == 201
        assert repeated.json() == created.json()
        project = created.json()["project"]
        project_id = project["projectId"]

        updated = client.put(
            f"/api/v2/domain/project-materials/projects/{project_id}",
            headers={**auth, "Idempotency-Key": "project-update-1"},
            json={
                "name": "乡村儿童成长支持",
                "expectedVersion": 1,
            },
        )
        assert updated.status_code == 200, updated.text
        assert updated.json()["project"]["version"] == 2
        assert updated.json()["project"]["name"] == "乡村儿童成长支持"

        conflict = client.put(
            f"/api/v2/domain/project-materials/projects/{project_id}",
            headers={**auth, "Idempotency-Key": "project-update-stale"},
            json={"name": "过期覆盖", "expectedVersion": 1},
        )
        assert conflict.status_code == 409
        assert conflict.json()["error"]["code"] == "project_version_conflict"

        for key, state, version in (
            ("project-freeze-1", "frozen", 2),
            ("project-unfreeze-1", "active", 3),
            ("project-archive-1", "archived", 4),
        ):
            transitioned = client.post(
                f"/api/v2/domain/project-materials/projects/{project_id}/lifecycle",
                headers={**auth, "Idempotency-Key": key},
                json={"targetState": state, "expectedVersion": version},
            )
            assert transitioned.status_code == 200, transitioned.text
            assert transitioned.json()["project"]["lifecycleState"] == state
            assert transitioned.json()["project"]["version"] == version + 1

    with runtime_connection(database, "cloud", read_only=True) as connection:
        envelope_count = int(
            connection.execute(
                """
                SELECT COUNT(*) FROM command_envelopes
                WHERE aggregate_id = ?
                """,
                (project_id,),
            ).fetchone()[0]
        )
        audit_count = int(
            connection.execute(
                """
                SELECT COUNT(*) FROM audit_events
                WHERE resource_id = ?
                """,
                (project_id,),
            ).fetchone()[0]
        )
        outbox_count = int(
            connection.execute(
                """
                SELECT COUNT(*) FROM delivery_outbox
                WHERE aggregate_id = ?
                """,
                (project_id,),
            ).fetchone()[0]
        )
    assert envelope_count == audit_count == outbox_count == 5


def _seed_materials(
    database: Path,
    *,
    organization_id: str,
    membership_id: str,
    project_id: str,
) -> None:
    now = "2026-07-30T08:00:00Z"
    with runtime_connection(database, "cloud") as connection:
        for index, document_kind, visibility, preview, raw_body in (
            (
                1,
                "shared_summary",
                "organization",
                "这是可共享的项目摘要。",
                "RAW_BODY_ONE_MUST_NOT_LEAVE_CLOUD",
            ),
            (
                2,
                "shared_summary",
                "organization",
                "这是第二份可共享摘要。",
                "RAW_BODY_TWO_MUST_NOT_LEAVE_CLOUD",
            ),
            (
                3,
                "raw_source",
                "self",
                "PRIVATE_PREVIEW_MUST_NOT_LEAVE_CLOUD",
                "PRIVATE_BODY_MUST_NOT_LEAVE_CLOUD",
            ),
        ):
            source_id = f"source_{index}"
            document_id = f"document_{index}"
            version_id = f"document_version_{index}"
            connection.execute(
                """
                INSERT INTO source_assets (
                    source_asset_id, organization_id, project_id,
                    storage_object_id, file_name, media_type, byte_size,
                    content_hash, source_kind, source_locator,
                    lifecycle_state, created_by_membership_id, version,
                    created_at, updated_at
                ) VALUES (?, ?, ?, NULL, ?, 'application/pdf', 128,
                          'same-source-hash', 'uploaded_file', ?,
                          'active', ?, 1, ?, ?)
                """,
                (
                    source_id,
                    organization_id,
                    project_id,
                    f"资料{index}.pdf",
                    f"/member/private/SOURCE_LOCATOR_{index}.pdf",
                    membership_id,
                    now,
                    now,
                ),
            )
            connection.execute(
                """
                INSERT INTO knowledge_documents (
                    document_id, organization_id, project_id,
                    project_assignment_state, source_asset_id,
                    owner_membership_id, department_id, title,
                    document_kind, visibility_scope, parse_state,
                    lifecycle_state, current_version, version,
                    created_at, updated_at
                ) VALUES (?, ?, ?, 'assigned', ?, ?, NULL, ?, ?, ?,
                          'ready', 'active', 1, 1, ?, ?)
                """,
                (
                    document_id,
                    organization_id,
                    project_id,
                    source_id,
                    membership_id,
                    f"资料{index}",
                    document_kind,
                    visibility,
                    now,
                    now,
                ),
            )
            connection.execute(
                """
                INSERT INTO document_versions (
                    document_version_id, organization_id, document_id,
                    version, content_hash, preview_text, markdown_content,
                    section_count, chunk_count, generator_version, created_at
                ) VALUES (?, ?, ?, 1, 'same-document-hash', ?, ?, 2, 3,
                          'test', ?)
                """,
                (
                    version_id,
                    organization_id,
                    document_id,
                    preview,
                    raw_body,
                    now,
                ),
            )
        connection.execute(
            """
            INSERT INTO processing_attempts (
                processing_attempt_id, organization_id, source_asset_id,
                document_id, processing_kind, state, attempt_no,
                error_code, error_message, started_at, finished_at, created_at
            ) VALUES ('link_run_1', ?, 'source_1', 'document_1',
                      'link_import', 'processing', 1, '', '', ?, NULL, ?)
            """,
            (organization_id, now, now),
        )
        connection.execute(
            """
            UPDATE knowledge_documents
            SET parse_state = 'failed'
            WHERE document_id = 'document_3'
            """
        )
        source_payload = {
            "entities": [
                {
                    "name": "益语智库",
                    "type": "company",
                    "aliases": ["益语"],
                    "confidence": 0.98,
                },
                {
                    "name": "益语智库AI",
                    "type": "company",
                    "aliases": [],
                    "confidence": 0.91,
                },
            ],
            "glossary": [
                {
                    "term": "服务人数",
                    "definition": "项目覆盖的服务对象人数",
                    "category": "成效指标",
                }
            ],
            "glossaryAttributes": [
                {
                    "term": "服务人数",
                    "attributeName": "2026目标",
                    "valueText": "1000",
                    "verificationStatus": "verified",
                },
                {
                    "term": "服务人数",
                    "attributeName": "2026目标",
                    "valueText": "1200",
                    "verificationStatus": "pending",
                },
            ],
        }
        connection.execute(
            """
            INSERT INTO intelligence_records (
                intelligence_id, organization_id, project_id, title,
                summary, source_url, record_kind, status, visibility_scope,
                created_by_membership_id, source_payload_json, version,
                created_at, updated_at
            ) VALUES ('intelligence_materials', ?, ?, '已发布项目资料提炼',
                      '组织已发布的项目资料提炼摘要', '', 'material_digest',
                      'accepted', 'organization', ?, ?, 1, ?, ?)
            """,
            (
                organization_id,
                project_id,
                membership_id,
                json.dumps(source_payload, ensure_ascii=False),
                now,
                now,
            ),
        )
        connection.commit()


def test_material_metadata_safe_summary_derivations_and_processing_state(
    tmp_path: Path,
) -> None:
    client, database = _cloud_client(tmp_path)
    with client:
        login, auth = _bootstrap(client)
        session = client.get("/api/v2/session/current", headers=auth).json()
        project = _default_project(client, auth)
        project_id = project["projectId"]
        _seed_materials(
            database,
            organization_id=login["organizationId"],
            membership_id=session["membershipId"],
            project_id=project_id,
        )

        preview = client.get(
            f"/api/v2/domain/project-materials/projects/{project_id}"
            "/documents/document_1/reading-preview",
            headers=auth,
        )
        assert preview.status_code == 200, preview.text
        assert preview.json()["readSummary"] == "这是可共享的项目摘要。"
        assert all(
            value is False
            for value in preview.json()["materialBoundary"].values()
        )

        shared_text = client.get(
            "/api/v2/domain/project-materials/documents/document_1/text",
            headers=auth,
        )
        assert shared_text.status_code == 200
        assert shared_text.json()["content"] == "这是可共享的项目摘要。"
        private_text = client.get(
            "/api/v2/domain/project-materials/documents/document_3/text",
            headers=auth,
        )
        assert private_text.status_code == 403
        assert private_text.json()["error"]["code"] == "source_content_not_shared"

        duplicates = client.get(
            f"/api/v2/domain/project-materials/projects/{project_id}"
            "/duplicate-documents",
            headers=auth,
        )
        assert duplicates.status_code == 200, duplicates.text
        assert duplicates.json()["groups"][0]["count"] == 3
        assert all(
            item["managedPath"] == "" and item["originalPath"] == ""
            for item in duplicates.json()["groups"][0]["documents"]
        )

        status = client.get(
            f"/api/v2/domain/project-materials/projects/{project_id}"
            "/knowledge-status",
            headers=auth,
        )
        assert status.status_code == 200, status.text
        assert status.json()["confirmedFacts"] >= 2
        assert status.json()["derivation"].startswith("strict_v2:")

        fact_bundle = client.get(
            f"/api/v2/domain/project-materials/projects/{project_id}/fact-bundle",
            headers=auth,
        )
        assert fact_bundle.status_code == 200, fact_bundle.text
        assert fact_bundle.json()["counts"]["atomic_facts"] == 2

        entities = client.get(
            f"/api/v2/domain/project-materials/projects/{project_id}/entities",
            headers=auth,
        )
        glossary = client.get(
            f"/api/v2/domain/project-materials/projects/{project_id}/glossary",
            headers=auth,
        )
        contradictions = client.get(
            f"/api/v2/domain/project-materials/projects/{project_id}"
            "/contradictions",
            headers=auth,
        )
        drift = client.get(
            f"/api/v2/domain/project-materials/projects/{project_id}"
            "/glossary-drift-alerts",
            headers=auth,
        )
        assert entities.json()["total"] == 2
        assert glossary.json()["entries"][0]["term"] == "服务人数"
        assert contradictions.json()["total"] == 1
        assert len(drift.json()["alerts"]) == 1

        entity_items = entities.json()["entities"]
        canonical_id = entity_items[0]["id"]
        merged_id = entity_items[1]["id"]
        verified_entity = client.post(
            f"/api/v2/domain/project-materials/entities/{canonical_id}/verify",
            headers={**auth, "Idempotency-Key": "verify-entity-1"},
            json={"status": "canonical", "reason": "人工确认"},
        )
        assert verified_entity.status_code == 200, verified_entity.text
        assert verified_entity.json()["verifiedStatus"] == "canonical"
        merged_entity = client.post(
            f"/api/v2/domain/project-materials/entities/{merged_id}/merge",
            headers={**auth, "Idempotency-Key": "merge-entity-1"},
            json={
                "survivingEntityId": canonical_id,
                "mergeReason": "人工确认同一机构",
            },
        )
        assert merged_entity.status_code == 200, merged_entity.text
        entities_after = client.get(
            f"/api/v2/domain/project-materials/projects/{project_id}/entities",
            headers=auth,
        ).json()
        assert entities_after["total"] == 1
        assert entities_after["entities"][0]["mentionCount"] == 2

        created_glossary = client.post(
            f"/api/v2/domain/project-materials/projects/{project_id}/glossary",
            headers={**auth, "Idempotency-Key": "create-glossary-1"},
            json={
                "term": "项目里程碑",
                "definition": "项目的重要阶段成果",
                "aliases": ["里程碑"],
                "category": "项目管理",
            },
        )
        assert created_glossary.status_code == 201, created_glossary.text
        repeated_glossary = client.post(
            f"/api/v2/domain/project-materials/projects/{project_id}/glossary",
            headers={**auth, "Idempotency-Key": "create-glossary-1"},
            json={
                "term": "项目里程碑",
                "definition": "项目的重要阶段成果",
                "aliases": ["里程碑"],
                "category": "项目管理",
            },
        )
        assert repeated_glossary.status_code == 201
        assert repeated_glossary.json() == created_glossary.json()
        glossary_entry = created_glossary.json()["entry"]
        updated_glossary = client.patch(
            "/api/v2/domain/project-materials/glossary/"
            f"{glossary_entry['id']}",
            headers={**auth, "Idempotency-Key": "update-glossary-1"},
            json={"definition": "经确认的重要阶段成果"},
        )
        assert updated_glossary.status_code == 200, updated_glossary.text
        assert (
            updated_glossary.json()["entry"]["definition"]
            == "经确认的重要阶段成果"
        )
        deleted_glossary = client.delete(
            "/api/v2/domain/project-materials/glossary/"
            f"{glossary_entry['id']}",
            headers={**auth, "Idempotency-Key": "delete-glossary-1"},
        )
        assert deleted_glossary.status_code == 200, deleted_glossary.text
        glossary_after = client.get(
            f"/api/v2/domain/project-materials/projects/{project_id}/glossary",
            headers=auth,
        ).json()
        assert all(
            item["id"] != glossary_entry["id"]
            for item in glossary_after["entries"]
        )

        attributes = client.get(
            f"/api/v2/domain/project-materials/projects/{project_id}"
            "/glossary-attributes",
            headers=auth,
        ).json()["attributes"]
        pending_attribute = next(
            item
            for item in attributes
            if item["verification_status"] == "pending"
        )
        verified_attribute = client.post(
            f"/api/v2/domain/project-materials/projects/{project_id}"
            f"/glossary-attributes/{pending_attribute['id']}/review",
            headers={**auth, "Idempotency-Key": "verify-attribute-1"},
            json={
                "reviewStatus": "verified",
                "verifiedBy": "user",
                "valueText": "1200",
            },
        )
        assert verified_attribute.status_code == 200, verified_attribute.text
        assert verified_attribute.json()["status"] == "verified"

        drift_id = drift.json()["alerts"][0]["id"]
        reviewed_drift = client.post(
            f"/api/v2/domain/project-materials/projects/{project_id}"
            f"/glossary-drift-alerts/{drift_id}/review",
            headers={**auth, "Idempotency-Key": "review-drift-1"},
            json={"action": "dismiss", "note": "保留现有目标"},
        )
        assert reviewed_drift.status_code == 200, reviewed_drift.text
        resolved_drift = client.get(
            f"/api/v2/domain/project-materials/projects/{project_id}"
            "/glossary-drift-alerts?status=dismissed",
            headers=auth,
        )
        assert resolved_drift.status_code == 200
        assert resolved_drift.json()["alerts"][0]["id"] == drift_id

        contradiction_id = contradictions.json()["contradictions"][0]["id"]
        reviewed_contradiction = client.post(
            "/api/v2/domain/project-materials/contradictions/"
            f"{contradiction_id}/review",
            headers={**auth, "Idempotency-Key": "review-contradiction-1"},
            json={
                "reviewStatus": "resolved",
                "acceptedFactId": pending_attribute["id"],
                "resolutionNote": "采用最新人工确认值",
            },
        )
        assert (
            reviewed_contradiction.status_code == 200
        ), reviewed_contradiction.text
        resolved_contradictions = client.get(
            f"/api/v2/domain/project-materials/projects/{project_id}"
            "/contradictions?status=resolved",
            headers=auth,
        )
        assert resolved_contradictions.status_code == 200
        assert (
            resolved_contradictions.json()["contradictions"][0]["id"]
            == contradiction_id
        )

        folder_plan = client.post(
            f"/api/v2/domain/project-materials/projects/{project_id}"
            "/folder-recommendation",
            headers=auth,
            json={},
        )
        repair = client.post(
            f"/api/v2/domain/project-materials/projects/{project_id}"
            "/auto-repair-preview",
            headers=auth,
            json={},
        )
        assert folder_plan.status_code == 200
        assert folder_plan.json()["totalDocumentCount"] >= 3
        assert repair.status_code == 200
        assert len(repair.json()["items"]) >= 3
        assert all(item["sourcePath"] is None for item in repair.json()["items"])
        repair_queue = client.post(
            f"/api/v2/domain/project-materials/projects/{project_id}"
            "/auto-repair-queue",
            headers={**auth, "Idempotency-Key": "queue-auto-repair-1"},
            json={"documentIds": ["document_3"]},
        )
        assert repair_queue.status_code == 409, repair_queue.text
        assert repair_queue.json()["error"]["code"] == (
            "local_material_auto_repair_executor_not_connected"
        )

        runs = client.get(
            f"/api/v2/domain/project-materials/projects/{project_id}"
            "/link-import-runs",
            headers=auth,
        )
        serialized = json.dumps(runs.json(), ensure_ascii=False)
        assert runs.status_code == 200
        assert runs.json()["runs"][0]["status"] == "running"
        assert runs.json()["runs"][0]["sourceUrl"] == ""
        assert "SOURCE_LOCATOR" not in serialized
        cancelled = client.post(
            f"/api/v2/domain/project-materials/projects/{project_id}"
            "/link-import-runs/link_run_1/cancel",
            headers={**auth, "Idempotency-Key": "cancel-link-run-1"},
        )
        assert cancelled.status_code == 200, cancelled.text
        assert cancelled.json()["run"]["status"] == "canceled"
        queued_link = client.post(
            f"/api/v2/domain/project-materials/projects/{project_id}"
            "/link-import-runs",
            headers={**auth, "Idempotency-Key": "start-link-run-1"},
            json={
                "url": "https://www.bilibili.com/video/BV1example",
                "useBrowserCookies": False,
                "cookieBrowser": "firefox",
            },
        )
        assert queued_link.status_code == 409, queued_link.text
        assert queued_link.json()["error"]["code"] == (
            "link_import_executor_not_connected"
        )

        registered = client.post(
            f"/api/v2/domain/project-materials/projects/{project_id}"
            "/materials/register-metadata",
            headers={**auth, "Idempotency-Key": "register-local-metadata-1"},
            json={
                "materials": [
                    {
                        "localSourceId": "local-source-only",
                        "fileName": "本机私有附件.txt",
                        "contentHash": "local-content-hash",
                        "byteSize": 42,
                        "mediaType": "text/plain",
                        "sourceKind": "local_private_metadata",
                    }
                ]
            },
        )
        assert registered.status_code == 201, registered.text
        assert registered.json()["materialBoundary"] == {
            "sourceFileContentUploaded": False,
            "sourceFilePathUploaded": False,
            "localSummaryUploaded": False,
        }
        published = client.post(
            f"/api/v2/domain/project-materials/projects/{project_id}"
            "/smart-import/publish",
            headers={**auth, "Idempotency-Key": "publish-smart-import-1"},
            json={
                "title": "已审阅访谈导入",
                "parsed": {
                    "entities": [{"name": "新伙伴机构", "kind": "company"}],
                    "events": [{"summary": "共同确认下一步"}],
                    "raw_text": "RAW_SMART_TEXT_MUST_NOT_UPLOAD",
                },
            },
        )
        assert published.status_code == 201, published.text
        assert published.json()["rawTextUploaded"] is False
        assert published.json()["publishedKnowledge"]["sourceType"] == (
            "structured_intelligence_summary"
        )
        assert published.json()["knowledgeDocumentVersion"] == 1
        assert published.json()["documentVersionId"]
        knowledge_context = client.get(
            f"/api/v2/projects/{project_id}/knowledge-context",
            headers=auth,
        )
        assert knowledge_context.status_code == 200, knowledge_context.text
        published_document_id = published.json()["knowledgeDocumentId"]
        published_context_item = next(
            item
            for item in knowledge_context.json()["organizationSharedKnowledge"]
            if item["sourceId"] == published_document_id
        )
        assert published_context_item["sourceType"] == "knowledge_summary"
        assert "intelligence_summary" in published_context_item[
            "sourceDescription"
        ]
        fact_bundle_after_publish = client.get(
            f"/api/v2/domain/project-materials/projects/{project_id}"
            "/fact-bundle",
            headers=auth,
        ).json()
        assert any(
            item["source_v2_document_id"] == published_document_id
            for item in fact_bundle_after_publish["atomic_facts"]
        )
        assert "RAW_SMART_TEXT_MUST_NOT_UPLOAD" not in json.dumps(
            knowledge_context.json(),
            ensure_ascii=False,
        )
        empty_publish = client.post(
            f"/api/v2/domain/project-materials/projects/{project_id}"
            "/smart-import/publish",
            headers={**auth, "Idempotency-Key": "publish-smart-import-empty"},
            json={
                "title": "空内容不得发布",
                "parsed": {
                    "entities": [],
                    "relationships": [],
                    "events": [],
                    "opinions": [],
                    "commitments": [],
                    "risk_signals": [],
                    "open_questions": [],
                },
            },
        )
        assert empty_publish.status_code == 422, empty_publish.text
        assert empty_publish.json()["error"]["code"] == (
            "published_summary_content_required"
        )

        archived = client.request(
            "DELETE",
            f"/api/v2/domain/project-materials/projects/{project_id}"
            "/documents/document_1",
            headers={**auth, "Idempotency-Key": "archive-document-1"},
            json={"expectedVersion": 1},
        )
        assert archived.status_code == 200, archived.text
        assert archived.json()["lifecycleState"] == "archived"

        all_serialized = json.dumps(
            {
                "preview": preview.json(),
                "shared": shared_text.json(),
                "duplicates": duplicates.json(),
                "bundle": fact_bundle.json(),
                "entities": entities.json(),
                "glossary": glossary.json(),
            },
            ensure_ascii=False,
        )
        for forbidden in (
            "RAW_BODY",
            "PRIVATE_BODY",
            "PRIVATE_PREVIEW",
            "SOURCE_LOCATOR",
            "source_locator",
            "markdown_content",
        ):
            assert forbidden not in all_serialized

    with runtime_connection(database, "cloud", read_only=True) as connection:
        registered_source = connection.execute(
            """
            SELECT source_locator
            FROM source_assets
            WHERE content_hash = 'local-content-hash'
            """
        ).fetchone()
        smart_payload = connection.execute(
            """
            SELECT source_payload_json
            FROM intelligence_records
            WHERE record_kind = 'smart_import_reviewed'
            """
        ).fetchone()
        published_summary = connection.execute(
            """
            SELECT d.visibility_scope, d.document_kind, v.preview_text,
                   v.markdown_content
            FROM knowledge_documents d
            JOIN document_versions v
              ON v.document_id = d.document_id
             AND v.version = d.current_version
            WHERE d.document_id = ?
            """,
            (published_document_id,),
        ).fetchone()
        queued_repair = connection.execute(
            """
            SELECT state
            FROM processing_attempts
            WHERE processing_kind = 'local_material_auto_repair'
              AND document_id = 'document_3'
            """
        ).fetchone()
        queued_link_source = connection.execute(
            """
            SELECT source_locator, content_hash
            FROM source_assets
            WHERE source_kind = 'bilibili_link_metadata'
            """
        ).fetchone()
        link_outbox = connection.execute(
            """
            SELECT payload_json, status
            FROM delivery_outbox
            WHERE event_type = 'project_material.link_import_queued'
            """
        ).fetchone()
        smart_audit = connection.execute(
            """
            SELECT summary_json
            FROM audit_events
            WHERE action = 'smart_import.published'
            """
        ).fetchone()
        smart_outbox_rows = connection.execute(
            """
            SELECT operation_id, aggregate_type, aggregate_id,
                   aggregate_version, event_type, payload_json
            FROM delivery_outbox
            WHERE event_type IN (
                'smart_import.published',
                'project_knowledge.summary_published'
            )
            ORDER BY event_type
            """
        ).fetchall()
    assert registered_source is not None
    assert registered_source["source_locator"] == ""
    assert smart_payload is not None
    assert "RAW_SMART_TEXT_MUST_NOT_UPLOAD" not in smart_payload["source_payload_json"]
    assert published_summary is not None
    assert published_summary["visibility_scope"] == "organization"
    assert published_summary["document_kind"] == "intelligence_summary"
    assert "RAW_SMART_TEXT_MUST_NOT_UPLOAD" not in (
        published_summary["preview_text"]
        + published_summary["markdown_content"]
    )
    assert queued_repair is None
    assert queued_link_source is None
    assert link_outbox is None
    assert smart_audit is not None
    smart_audit_payload = json.loads(smart_audit["summary_json"])
    assert set(smart_audit_payload) == {
        "projectId",
        "counts",
        "knowledgeDocumentId",
        "knowledgeDocumentVersion",
        "contentHash",
    }
    assert "parsed" not in smart_audit["summary_json"]
    assert len(smart_outbox_rows) == 2
    assert len({row["operation_id"] for row in smart_outbox_rows}) == 1
    assert {
        row["aggregate_type"] for row in smart_outbox_rows
    } == {"intelligence", "knowledge_document"}
    assert all(
        published_document_id in row["payload_json"]
        for row in smart_outbox_rows
    )


def test_duplicate_resolution_is_cloud_atomic_on_late_cas_conflict(
    tmp_path: Path,
) -> None:
    client, database = _cloud_client(tmp_path)
    with client:
        _, auth = _bootstrap(client)
        project_id = _default_project(client, auth)["projectId"]
        registered = client.post(
            f"/api/v2/domain/project-materials/projects/{project_id}"
            "/materials/register-metadata",
            headers={**auth, "Idempotency-Key": "duplicates-register"},
            json={
                "materials": [
                    {
                        "localSourceId": f"local-duplicate-{index}",
                        "fileName": "同名重复资料.md",
                        "contentHash": f"{index}" * 64,
                        "byteSize": index,
                        "mediaType": "text/markdown",
                    }
                    for index in range(1, 4)
                ]
            },
        )
        assert registered.status_code == 201, registered.text
        documents = registered.json()["documents"]
        groups = client.get(
            f"/api/v2/domain/project-materials/projects/{project_id}"
            "/duplicate-documents",
            headers=auth,
        )
        assert groups.status_code == 200, groups.text
        group = next(
            item for item in groups.json()["groups"] if item["count"] == 3
        )
        keep_id = documents[0]["documentId"]
        delete_ids = [
            documents[1]["documentId"],
            documents[2]["documentId"],
        ]
        conflict_payload = {
            "groupKey": group["groupKey"],
            "action": "delete_others",
            "keepV2DocumentIds": [keep_id],
            "deleteV2DocumentIds": delete_ids,
            "documents": [
                {"documentId": delete_ids[0], "expectedVersion": 1},
                {"documentId": delete_ids[1], "expectedVersion": 99},
            ],
            "migrateReferences": True,
            "note": "CAS 冲突不得半完成",
        }
        conflict = client.post(
            f"/api/v2/domain/project-materials/projects/{project_id}"
            "/duplicate-documents/resolve",
            headers={**auth, "Idempotency-Key": "duplicates-conflict"},
            json=conflict_payload,
        )
        assert conflict.status_code == 409, conflict.text
        assert conflict.json()["error"]["code"] == "document_version_conflict"
        with runtime_connection(database, "cloud", read_only=True) as connection:
            states_after_conflict = connection.execute(
                """
                SELECT lifecycle_state, version
                FROM knowledge_documents
                WHERE document_id IN (?, ?)
                """,
                tuple(delete_ids),
            ).fetchall()
        assert {
            (str(row["lifecycle_state"]), int(row["version"]))
            for row in states_after_conflict
        } == {("active", 1)}

        success_payload = {
            **conflict_payload,
            "documents": [
                {"documentId": document_id, "expectedVersion": 1}
                for document_id in delete_ids
            ],
        }
        headers = {**auth, "Idempotency-Key": "duplicates-resolve"}
        resolved = client.post(
            f"/api/v2/domain/project-materials/projects/{project_id}"
            "/duplicate-documents/resolve",
            headers=headers,
            json=success_payload,
        )
        assert resolved.status_code == 200, resolved.text
        assert resolved.json()["deletedCount"] == 2
        replay = client.post(
            f"/api/v2/domain/project-materials/projects/{project_id}"
            "/duplicate-documents/resolve",
            headers=headers,
            json=success_payload,
        )
        assert replay.status_code == 200
        assert replay.json() == resolved.json()

    with runtime_connection(database, "cloud", read_only=True) as connection:
        archived = connection.execute(
            """
            SELECT lifecycle_state, version
            FROM knowledge_documents
            WHERE document_id IN (?, ?)
            """,
            tuple(delete_ids),
        ).fetchall()
        command = connection.execute(
            """
            SELECT operation_id
            FROM command_envelopes
            WHERE idempotency_key = 'duplicates-resolve'
              AND command_type = 'knowledge_documents.duplicates_resolved'
            """
        ).fetchall()
        audits = connection.execute(
            """
            SELECT COUNT(*) AS count
            FROM audit_events
            WHERE action = 'knowledge_documents.duplicates_resolved'
            """
        ).fetchone()
        outbox = connection.execute(
            """
            SELECT COUNT(*) AS count
            FROM delivery_outbox
            WHERE operation_id = ?
            """,
            (command[0]["operation_id"],),
        ).fetchone()
    assert {
        (str(row["lifecycle_state"]), int(row["version"]))
        for row in archived
    } == {("archived", 2)}
    assert len(command) == 1
    assert int(audits["count"]) == 1
    assert int(outbox["count"]) == 3


def test_local_import_and_smart_import_use_strict_storage_objects(
    tmp_path: Path,
) -> None:
    assert LocalProjectMaterialsRepository.__module__ == (
        "backend.app.project_materials_local"
    )
    assert not hasattr(GC07ProjectMaterialsRepository, "import_paths")
    assert not hasattr(GC07ProjectMaterialsRepository, "create_smart_session")
    runtime = WorkspaceRuntime(
        tmp_path / "local" / "strict-local.db",
        MemorySecretStore(),
    )
    sandbox_id = runtime.current()["sandbox"]["sandboxId"]
    runtime._current_context = lambda require_ready=True: SimpleNamespace(  # type: ignore[method-assign]
        sandbox_id=sandbox_id,
        cloud_instance_id="cloud-local-test",
        organization_id="organization-local-test",
        cloud_api_url="https://local-test.invalid",
        principal_id="principal-local-test",
        membership_id="membership-local-test",
    )
    store = LocalProjectMaterialsRepository(runtime)
    source = tmp_path / "项目资料.md"
    source.write_text("益语智库团队计划推进项目。\n当前风险是时间不足。", encoding="utf-8")

    imported = store.import_paths(
        project_id="project-local-test",
        mode="file",
        paths=[str(source)],
    )
    assert len(imported["materials"]) == 1
    assert imported["materials"][0]["summaryKind"] == "text_excerpt"
    local_context = runtime._local_project_materials(
        sandbox_id=sandbox_id,
        project_id="project-local-test",
    )
    assert local_context["invalidCount"] == 0
    assert local_context["items"][0]["summary"].startswith("益语智库团队")
    text_material = store.import_text(
        project_id="project-local-test",
        title="任务交接背景",
        content="这是只保存在当前设备上的任务交接背景。",
    )
    assert text_material["fileName"].endswith(".md")
    assert text_material["summaryKind"] == "text_excerpt"
    assert Path(text_material["managedPath"]).read_text(
        encoding="utf-8"
    ) == "这是只保存在当前设备上的任务交接背景。"
    store.bind_cloud_documents(
        project_id="project-local-test",
        local_materials=[imported["materials"][0], text_material],
        cloud_documents=[
            {
                "localSourceId": imported["materials"][0]["localSourceId"],
                "documentId": "cloud-file-document",
            },
            {
                "localSourceId": text_material["localSourceId"],
                "documentId": "cloud-text-document",
            }
        ],
    )
    local_documents = {
        item["id"]: item
        for item in store.documents("project-local-test")
    }
    assert set(local_documents) == {
        "cloud-file-document",
        "cloud-text-document",
    }
    assert local_documents["cloud-file-document"]["path"] == str(source.resolve())
    assert local_documents["cloud-file-document"]["originalSourcePath"] == str(
        source.resolve()
    )
    assert Path(local_documents["cloud-text-document"]["path"]).is_file()
    assert all(
        item["source"] == "member_local"
        for item in local_documents.values()
    )
    local_text = store.document_text("cloud-text-document")
    assert local_text["sourceScope"] == "local_private"
    assert local_text["editableInPlace"] is True
    assert "任务交接背景" in local_text["content"]
    updated_text = store.update_document_text(
        "cloud-text-document",
        title="更新后的背景",
        content="这是更新后仍只保存在当前设备上的背景。",
    )
    assert updated_text["title"] == "更新后的背景"
    assert updated_text["fileName"] == "更新后的背景.md"
    assert Path(updated_text["path"]).name.endswith("-更新后的背景.md")
    assert Path(updated_text["path"]).is_file()
    assert not Path(text_material["managedPath"]).exists()
    assert store.document_text("cloud-text-document")["content"].startswith(
        "这是更新后"
    )
    candidates = store.optimization_candidates(["project-local-test"])
    local_candidate = next(
        item
        for item in candidates
        if item["documentId"] == "cloud-text-document"
    )
    assert local_candidate["deepRead"] is False


    summarized = store.update_ai_summary(
        "cloud-text-document",
        summary="AI 仅根据当前设备正文形成的深度摘要。",
        model_name="test-model",
    )
    assert summarized["summaryKind"] == "ai_summary"
    candidates = store.optimization_candidates(["project-local-test"])
    assert next(
        item
        for item in candidates
        if item["documentId"] == "cloud-text-document"
    )["deepRead"] is True
    local_context = runtime._local_project_materials(
        sandbox_id=sandbox_id,
        project_id="project-local-test",
    )
    assert any(
        item["summary"].startswith("AI 仅根据")
        for item in local_context["items"]
    )
    folder = store.create_folder(
        "project-local-test",
        {"label": "项目背景"},
    )
    moved = store.move_document(
        "project-local-test",
        "cloud-text-document",
        {"folderId": folder["id"]},
    )
    assert moved["folderId"] == folder["id"]
    assert store.folders("project-local-test")[0]["fileCount"] == 1
    template = tmp_path / "项目简报.docx"
    template_document = Document()
    template_document.add_heading("{{项目名称}}", level=1)
    template_document.add_paragraph("{{项目简介}}")
    template_document.add_paragraph("{{待补充字段}}")
    template_document.save(template)
    fill_run = store.start_template_fill(
        "project-local-test",
        template_path=str(template),
        values={"项目名称": "严格新版项目", "项目简介": "真实闭环"},
    )
    assert fill_run["status"] == "completed"
    assert fill_run["filledCount"] == 2
    assert fill_run["missingCount"] == 1
    output_document = Document(str(fill_run["outputPath"]))
    assert output_document.paragraphs[0].text == "严格新版项目"
    assert output_document.paragraphs[1].text == "真实闭环"
    assert store.template_fill_run(
        "project-local-test",
        str(fill_run["id"]),
    )["outputPath"] == fill_run["outputPath"]

    state = store.create_smart_session(
        {"clientId": "project-local-test", "title": "访谈材料导入"}
    )
    session_id = state["session"]["id"]
    uploaded = store.upload_smart_file(
        session_id,
        SimpleNamespace(
            filename="访谈附件.txt",
            content_type="text/plain",
            file=BytesIO(b"local smart import attachment"),
        ),
    )
    assert uploaded["session_id"] == session_id
    state = store.add_smart_chunk(
        session_id,
        {
            "rawText": "益语智库团队计划推进项目。\n风险是时间不足。",
            "fileIds": [uploaded["id"]],
            "autoParse": True,
        },
    )
    chunk_id = state["chunks"][0]["id"]
    assert state["chunks"][0]["parse_status"] == "parsed"
    preview = store.smart_preview(session_id)
    assert preview["chunks_parsed"] == 1
    assert preview["events"]
    assert preview["risk_signals"]
    assert preview["files_classified"][0]["original_filename"] == "访谈附件.txt"
    patched = store.patch_smart_chunk(
        chunk_id,
        {
            **state["chunks"][0]["parsed"],
            "open_questions": ["还需确认负责人"],
        },
    )
    assert patched["chunks"][0]["user_edited_parsed"] == 1
    imported_state = store.mark_smart_imported(
        session_id,
        document_ids={uploaded["id"]: "cloud-document-1"},
    )
    assert imported_state["session"]["status"] == "imported"
    assert imported_state["staged_files"][0]["document_id"] == "cloud-document-1"

    file_only_state = store.create_smart_session(
        {"clientId": "project-local-test", "title": "仅附件导入"}
    )
    file_only_id = file_only_state["session"]["id"]
    file_only_upload = store.upload_smart_file(
        file_only_id,
        SimpleNamespace(
            filename="日慈项目背景.txt",
            content_type="text/plain",
            file=BytesIO(
                "日慈基金会计划推进儿童支持项目。\n当前风险是资料口径不一致。".encode(
                    "utf-8"
                )
            ),
        ),
    )
    file_only_state = store.add_smart_chunk(
        file_only_id,
        {
            "rawText": "",
            "fileIds": [file_only_upload["id"]],
            "autoParse": True,
        },
    )
    assert file_only_state["chunks"][0]["parse_status"] == "parsed"
    file_only_preview = store.smart_preview(file_only_id)
    assert any(
        item["summary"] == "日慈基金会计划推进儿童支持项目。"
        for item in file_only_preview["events"]
    )
    assert any(
        item["description"] == "当前风险是资料口径不一致。"
        for item in file_only_preview["risk_signals"]
    )

    restarted_runtime = WorkspaceRuntime(
        runtime.database_path,
        MemorySecretStore(),
    )
    restarted_runtime._current_context = (  # type: ignore[method-assign]
        lambda require_ready=True: SimpleNamespace(
            sandbox_id=sandbox_id,
            membership_id="membership-local-test",
        )
    )
    restarted_preview = LocalProjectMaterialsRepository(
        restarted_runtime
    ).smart_preview(file_only_id)
    assert restarted_preview["events"] == file_only_preview["events"]
    assert restarted_preview["risk_signals"] == file_only_preview["risk_signals"]

    with runtime_connection(
        runtime.database_path,
        "local",
        read_only=True,
    ) as connection:
        rows = connection.execute(
            """
            SELECT media_type, lifecycle_state
            FROM storage_objects
            WHERE sandbox_id = ?
            """,
            (sandbox_id,),
        ).fetchall()
        storage_commands = int(
            connection.execute(
                """
                SELECT COUNT(*) FROM command_envelopes
                WHERE sandbox_id = ?
                  AND command_type = 'local.storage_object.put'
                """,
                (sandbox_id,),
            ).fetchone()[0]
        )
        storage_audits = int(
            connection.execute(
                """
                SELECT COUNT(*) FROM audit_events
                WHERE sandbox_id = ?
                  AND action = 'local.storage_object.updated'
                """,
                (sandbox_id,),
            ).fetchone()[0]
        )
        storage_outbox = int(
            connection.execute(
                """
                SELECT COUNT(*) FROM delivery_outbox
                WHERE sandbox_id = ?
                  AND event_type = 'local.storage_object.updated'
                """,
                (sandbox_id,),
            ).fetchone()[0]
        )
    media_types = {str(row["media_type"]) for row in rows}
    assert "application/vnd.yiyu.project-knowledge-summary+json" in media_types
    assert "application/vnd.yiyu.smart-import-session+json" in media_types
    assert "application/vnd.yiyu.project-local-state+json" in media_types
    assert all(row["lifecycle_state"] == "active" for row in rows)
    assert storage_commands == storage_audits == storage_outbox
    assert storage_commands >= len(rows)


def test_workbench_exposes_only_latest_transcript_projection_per_task() -> None:
    documents = _current_project_document_projection([
        {
            "id": "audio-document",
            "localSourceId": "audio-source",
            "title": "项目访谈.wav",
            "mediaType": "audio/wav",
            "taskId": "task-transcript",
            "importedAt": "2026-08-24T09:59:00+00:00",
        },
        {
            "id": "transcript-document-v1",
            "title": "项目访谈-录音转写.md",
            "mediaType": "text/markdown",
            "taskId": "task-transcript",
            "importedAt": "2026-08-24T10:00:00+00:00",
        },
        {
            "id": "transcript-document-orphan",
            "title": "项目访谈-录音转写.md",
            "mediaType": "text/markdown",
            "taskId": None,
            "importedAt": "2026-08-24T10:00:30+00:00",
        },
        {
            "id": "transcript-document-v2",
            "title": "项目访谈-录音转写.md",
            "mediaType": "text/markdown",
            "taskId": "task-transcript",
            "importedAt": "2026-08-24T10:01:00+00:00",
        },
    ])
    transcript_documents = [
        item for item in documents if item.get("isTranscriptProjection")
    ]
    assert [item["id"] for item in transcript_documents] == [
        "transcript-document-v2"
    ]
    assert transcript_documents[0]["recordingSourceId"] == "audio-source"


def test_local_import_receipts_reuse_sources_after_cloud_failure_and_restart(
    tmp_path: Path,
) -> None:
    database = tmp_path / "local" / "strict-local.db"
    runtime = WorkspaceRuntime(database, MemorySecretStore())
    sandbox_id = runtime.current()["sandbox"]["sandboxId"]
    runtime._current_context = lambda require_ready=True: SimpleNamespace(  # type: ignore[method-assign]
        sandbox_id=sandbox_id,
        cloud_instance_id="cloud-local-retry",
        organization_id="organization-local-retry",
        cloud_api_url="https://local-retry.invalid",
        principal_id="principal-local-retry",
        membership_id="membership-local-test",
    )
    source = tmp_path / "待重试资料.md"
    source.write_text("云端第一次失败后，本机不得重复复制。", encoding="utf-8")
    cloud_calls = 0

    def flaky_cloud_command(
        method: str,
        path: str,
        *,
        payload: dict[str, Any],
        idempotency_key: str,
        refresh_business: bool = True,
    ) -> dict[str, Any]:
        del method, path, idempotency_key, refresh_business
        nonlocal cloud_calls
        cloud_calls += 1
        if cloud_calls == 1:
            raise LocalRuntimeError(
                503,
                "cloud_temporarily_unavailable",
                "组织云暂时不可用",
            )
        return {
            "importRunId": "import-run-retry",
            "documents": [
                {
                    "localSourceId": item["localSourceId"],
                    "documentId": f"cloud-{index}",
                }
                for index, item in enumerate(payload["materials"], start=1)
            ],
            "importedCount": len(payload["materials"]),
            "skippedCount": 0,
            "materialBoundary": {
                "sourceFileContentUploaded": False,
                "sourceFilePathUploaded": False,
                "localSummaryUploaded": False,
            },
        }

    runtime.cloud_command = flaky_cloud_command  # type: ignore[method-assign]
    compatibility = SimpleNamespace(runtime=runtime)
    request = UiRequest(
        method="POST",
        path="imports",
        query={},
        body={
            "clientId": "project-local-retry",
            "mode": "file",
            "paths": [str(source)],
        },
        idempotency_key="import-retry-stable",
    )
    partial = router.dispatch(compatibility, request)
    assert partial[0]["status"] == "partial"
    assert partial[0]["localState"] == "ready"
    assert partial[0]["cloudMetadataState"] == "failed_retryable"
    assert partial[0]["documents"][0]["documentId"].startswith(
        "local-pending:"
    )
    completed = router.dispatch(compatibility, request)
    assert completed[0]["documents"][0]["documentId"] == "cloud-1"
    assert cloud_calls == 2

    with runtime_connection(database, "local", read_only=True) as connection:
        source_rows = connection.execute(
            """
            SELECT object_id
            FROM storage_objects
            WHERE sandbox_id = ? AND media_type = 'text/markdown'
            """,
            (sandbox_id,),
        ).fetchall()
        operation_rows = connection.execute(
            """
            SELECT object_id
            FROM storage_objects
            WHERE sandbox_id = ? AND media_type = ?
            """,
            (
                sandbox_id,
                LocalProjectMaterialsRepository.IMPORT_OPERATION_MEDIA_TYPE,
            ),
        ).fetchall()
    assert len(source_rows) == 1
    assert len(operation_rows) == 1

    text_store = LocalProjectMaterialsRepository(runtime)
    first_text = text_store.import_text(
        project_id="project-local-retry",
        title="跨重启文本",
        content="文本正文只复制一次。",
        idempotency_key="text-retry-stable",
    )
    reopened = WorkspaceRuntime(database, MemorySecretStore())
    reopened._current_context = lambda require_ready=True: SimpleNamespace(  # type: ignore[method-assign]
        sandbox_id=sandbox_id,
        membership_id="membership-local-test",
    )
    replayed_text = LocalProjectMaterialsRepository(reopened).import_text(
        project_id="project-local-retry",
        title="跨重启文本",
        content="文本正文只复制一次。",
        idempotency_key="text-retry-stable",
    )
    assert replayed_text == first_text
    with pytest.raises(LocalRuntimeError) as conflict:
        LocalProjectMaterialsRepository(reopened).import_text(
            project_id="project-local-retry",
            title="跨重启文本",
            content="同一幂等键不能替换成另一份正文。",
            idempotency_key="text-retry-stable",
        )
    assert conflict.value.code == "local_import_idempotency_conflict"


def test_local_material_write_keeps_captured_sandbox_during_switch(
    tmp_path: Path,
) -> None:
    runtime = WorkspaceRuntime(
        tmp_path / "local" / "strict-local.db",
        MemorySecretStore(),
    )
    sandbox_a = runtime.current()["sandbox"]["sandboxId"]
    with runtime_connection(runtime.database_path, "local") as connection:
        device_id = str(
            connection.execute(
                "SELECT device_id FROM device_registry WHERE status = 'active'"
            ).fetchone()["device_id"]
        )
        connection.execute(
            """
            INSERT INTO workspace_sandboxes (
                sandbox_id, device_id, sandbox_kind, replica_epoch,
                runtime_status, display_name, is_active, version,
                created_at, updated_at
            ) VALUES ('sandbox-b', ?, 'organization', 1, 'ready',
                      '另一组织', 0, 1,
                      '2026-07-30T00:00:00.000Z',
                      '2026-07-30T00:00:00.000Z')
            """,
            (device_id,),
        )
        connection.commit()
    calls = 0

    def switched_context(require_ready: bool = True) -> SimpleNamespace:
        del require_ready
        nonlocal calls
        calls += 1
        return SimpleNamespace(
            sandbox_id=sandbox_a if calls == 1 else "sandbox-b",
            membership_id="membership-local-test",
        )

    runtime._current_context = switched_context  # type: ignore[method-assign]
    source = tmp_path / "切换竞态.md"
    source.write_text("只能落入发起请求时捕获的空间", encoding="utf-8")
    LocalProjectMaterialsRepository(runtime).import_paths(
        project_id="project-a",
        mode="file",
        paths=[str(source)],
    )
    with runtime_connection(
        runtime.database_path,
        "local",
        read_only=True,
    ) as connection:
        rows = connection.execute(
            """
            SELECT DISTINCT sandbox_id
            FROM storage_objects
            WHERE media_type IN (
                'text/markdown',
                'application/vnd.yiyu.project-knowledge-summary+json'
            )
            """
        ).fetchall()
        command_sandboxes = connection.execute(
            """
            SELECT DISTINCT sandbox_id
            FROM command_envelopes
            WHERE command_type = 'local.storage_object.put'
            """
        ).fetchall()
    assert {str(row["sandbox_id"]) for row in rows} == {sandbox_a}
    assert {str(row["sandbox_id"]) for row in command_sandboxes} == {sandbox_a}


def test_member_local_edit_updates_only_cloud_metadata(
    tmp_path: Path,
) -> None:
    client, database = _cloud_client(tmp_path)
    with client:
        _, headers = _bootstrap(client)
        project = _default_project(client, headers)
        created = client.post(
            f"/api/v2/domain/project-materials/projects/{project['projectId']}"
            "/materials/register-metadata",
            headers={**headers, "Idempotency-Key": "local-meta-create"},
            json={
                "materials": [
                    {
                        "localSourceId": "device-source-1",
                        "fileName": "项目背景.md",
                        "contentHash": "a" * 64,
                        "byteSize": 18,
                        "mediaType": "text/markdown",
                        "sourceKind": "local_private_text",
                    }
                ]
            },
        )
        assert created.status_code == 201, created.text
        document_id = created.json()["documents"][0]["documentId"]
        updated = client.patch(
            f"/api/v2/domain/project-materials/projects/{project['projectId']}"
            f"/documents/{document_id}/local-metadata",
            headers={**headers, "Idempotency-Key": "local-meta-update"},
            json={
                "expectedVersion": 1,
                "title": "项目背景（已更新）",
                "fileName": "项目背景.md",
                "contentHash": "b" * 64,
                "byteSize": 33,
                "mediaType": "text/markdown",
            },
        )
        assert updated.status_code == 200, updated.text
        assert updated.json()["version"] == 2
        assert updated.json()["materialBoundary"] == {
            "sourceFileContentUploaded": False,
            "sourceFilePathUploaded": False,
            "storageLocatorUploaded": False,
        }
        replay = client.patch(
            f"/api/v2/domain/project-materials/projects/{project['projectId']}"
            f"/documents/{document_id}/local-metadata",
            headers={**headers, "Idempotency-Key": "local-meta-update"},
            json={
                "expectedVersion": 1,
                "title": "项目背景（已更新）",
                "fileName": "项目背景.md",
                "contentHash": "b" * 64,
                "byteSize": 33,
                "mediaType": "text/markdown",
            },
        )
        assert replay.status_code == 200
        assert replay.json() == updated.json()
        published = client.post(
            f"/api/v2/domain/project-materials/projects/{project['projectId']}"
            f"/documents/{document_id}/publish-local-summary",
            headers={
                **headers,
                "Idempotency-Key": "local-summary-publish",
            },
            json={
                "expectedVersion": 2,
                "sourceContentHash": "b" * 64,
                "summary": "这是由当前设备正文提炼、允许组织共享的项目摘要。",
                "generatorVersion": "test-summary-model",
            },
        )
        assert published.status_code == 200, published.text
        assert published.json()["parseState"] == "ready"
        assert published.json()["visibilityScope"] == "organization"
        assert published.json()["version"] == 3
        published_replay = client.post(
            f"/api/v2/domain/project-materials/projects/{project['projectId']}"
            f"/documents/{document_id}/publish-local-summary",
            headers={
                **headers,
                "Idempotency-Key": "local-summary-publish",
            },
            json={
                "expectedVersion": 2,
                "sourceContentHash": "b" * 64,
                "summary": "这是由当前设备正文提炼、允许组织共享的项目摘要。",
                "generatorVersion": "test-summary-model",
            },
        )
        assert published_replay.status_code == 200
        assert published_replay.json() == published.json()
    with runtime_connection(database, "cloud", read_only=True) as connection:
        row = connection.execute(
            """
            SELECT d.title, d.version, d.current_version, d.parse_state,
                   d.visibility_scope, d.document_kind,
                   a.content_hash, a.byte_size, a.source_locator
            FROM knowledge_documents d
            JOIN source_assets a ON a.source_asset_id = d.source_asset_id
            WHERE d.document_id = ?
            """,
            (document_id,),
        ).fetchone()
        command = connection.execute(
            """
            SELECT payload_json FROM command_envelopes
            WHERE idempotency_key = 'local-meta-update'
            """
        ).fetchone()
        summary_version = connection.execute(
            """
            SELECT preview_text, markdown_content, generator_version
            FROM document_versions
            WHERE document_id = ? AND version = 1
            """,
            (document_id,),
        ).fetchone()
        publish_command = connection.execute(
            """
            SELECT payload_json
            FROM command_envelopes
            WHERE idempotency_key = 'local-summary-publish'
            """
        ).fetchone()
    assert row is not None
    assert row["title"] == "项目背景（已更新）"
    assert row["version"] == 3
    assert row["current_version"] == 1
    assert row["parse_state"] == "ready"
    assert row["visibility_scope"] == "organization"
    assert row["document_kind"] == "project_material_summary"
    assert row["content_hash"] == "b" * 64
    assert row["byte_size"] == 33
    assert row["source_locator"] == ""
    assert command is not None
    assert "当前设备" not in str(command["payload_json"])


def test_local_metadata_edit_releases_legacy_deleted_content_hash(
    tmp_path: Path,
) -> None:
    client, database = _cloud_client(tmp_path)
    with client:
        _, headers = _bootstrap(client)
        project = _default_project(client, headers)
        created = client.post(
            f"/api/v2/domain/project-materials/projects/{project['projectId']}"
            "/materials/register-metadata",
            headers={**headers, "Idempotency-Key": "local-meta-collision-create"},
            json={
                "materials": [
                    {
                        "localSourceId": "legacy-deleted-source",
                        "fileName": "旧文件.md",
                        "contentHash": "c" * 64,
                        "byteSize": 12,
                        "mediaType": "text/markdown",
                    },
                    {
                        "localSourceId": "current-source",
                        "fileName": "新文件.md",
                        "contentHash": "d" * 64,
                        "byteSize": 16,
                        "mediaType": "text/markdown",
                    },
                ]
            },
        )
        assert created.status_code == 201, created.text
        first, current = created.json()["documents"]
        # Reproduce tombstones created by releases before deletion began
        # releasing the nullable deduplication key.
        with runtime_connection(database, "cloud") as connection:
            connection.execute(
                "UPDATE source_assets SET lifecycle_state='deleted', "
                "availability_state='deleted', deleted_at=updated_at "
                "WHERE id=?",
                (first["documentId"],),
            )
            connection.commit()
        updated = client.patch(
            f"/api/v2/domain/project-materials/projects/{project['projectId']}"
            f"/documents/{current['documentId']}/local-metadata",
            headers={**headers, "Idempotency-Key": "local-meta-collision-update"},
            json={
                "expectedVersion": 1,
                "title": "新文件（改名）",
                "contentHash": "c" * 64,
                "byteSize": 12,
                "mediaType": "text/markdown",
            },
        )
        assert updated.status_code == 200, updated.text
        assert updated.json()["releasedTombstoneId"] == first["documentId"]
    with runtime_connection(database, "cloud", read_only=True) as connection:
        rows = connection.execute(
            "SELECT id,content_hash,lifecycle_state FROM source_assets "
            "WHERE id IN (?,?) ORDER BY id",
            (first["documentId"], current["documentId"]),
        ).fetchall()
    by_id = {str(row["id"]): row for row in rows}
    assert by_id[first["documentId"]]["content_hash"] is None
    assert by_id[first["documentId"]]["lifecycle_state"] == "deleted"
    assert by_id[current["documentId"]]["content_hash"] == "c" * 64
    assert "这是" not in str(command["payload_json"])
    assert summary_version is not None
    assert "允许组织共享" in summary_version["markdown_content"]
    assert summary_version["generator_version"] == "test-summary-model"
    assert publish_command is not None
    assert "允许组织共享" not in str(publish_command["payload_json"])


class _FakeRuntime:
    def __init__(self) -> None:
        self.calls: list[tuple[str, str]] = []
        self.command_payloads: list[dict[str, Any]] = []

    @staticmethod
    def _project() -> dict[str, Any]:
        return {
            "projectId": "project-1",
            "name": "项目一",
            "alias": "P1",
            "summary": "简介",
            "domain": "公益",
            "color": "#123456",
            "lifecycleState": "active",
            "version": 1,
            "participantMembershipIds": ["member-1"],
            "documentCount": 2,
            "taskCount": 3,
        }

    def require_project_capability(self, project_id: str, capability: str) -> dict[str, Any]:
        assert project_id == "project-1"
        assert capability == "read"
        return self._project()

    def cloud_query(
        self,
        path: str,
        *,
        query: dict[str, str] | None = None,
    ) -> dict[str, Any]:
        self.calls.append(("GET", path))
        if path.endswith("/projects"):
            return {"projects": [self._project()]}
        if path.endswith("/delete-preview"):
            return {
                "name": "项目一",
                "version": 1,
                "documentCount": 2,
                "taskCount": 3,
                "eventLineCount": 1,
                "narrativeCount": 1,
            }
        if path.endswith("/knowledge-status"):
            return {"clientId": "project-1", "confirmedFacts": 2}
        if path.endswith("/fact-bundle"):
            return {"client": {"id": "project-1"}, "counts": {}}
        if path.endswith("/duplicate-documents"):
            return {"groups": [{"groupKey": "hash:x"}]}
        if path.endswith("/entities"):
            return {"entities": [], "total": 0}
        if path.endswith("/entity-merge-candidates"):
            return {"candidates": []}
        if path.endswith("/glossary"):
            return {"entries": [], "total": 0}
        if path.endswith("/glossary-attributes"):
            return {"attributes": []}
        if path.endswith("/glossary-drift-alerts"):
            return {"alerts": []}
        if path.endswith("/contradictions"):
            return {"contradictions": [], "total": 0}
        if path.endswith("/reading-preview"):
            return {
                "title": "摘要",
                "parseState": "ready",
                "publishedSummary": True,
                "readSummary": "摘要正文",
                "aggregateVersion": 1,
            }
        if path.endswith("/text"):
            return {
                "title": "摘要",
                "kind": "shared_summary",
                "content": "摘要正文",
            }
        if "/link-import-runs/" in path:
            return {"runs": [{"runId": "run-1", "status": "running"}]}
        if path.endswith("/link-import-runs"):
            return {"runs": [{"runId": "run-1", "status": "running"}]}
        return {"project": self._project()}

    def cloud_command(
        self,
        method: str,
        path: str,
        *,
        payload: dict[str, Any],
        idempotency_key: str,
        refresh_business: bool = True,
    ) -> dict[str, Any]:
        self.calls.append((method, path))
        self.command_payloads.append(dict(payload))
        if path.endswith("/cancel"):
            return {"run": {"runId": "run-1", "status": "canceled"}}
        if path.endswith("/materials/register-metadata"):
            return {
                "documents": [
                    {
                        "localSourceId": "local-text-1",
                        "documentId": "document-text-1",
                    }
                ],
                "materialBoundary": {
                    "sourceFileContentUploaded": False,
                    "sourceFilePathUploaded": False,
                    "localSummaryUploaded": False,
                },
            }
        if path.endswith("/link-import-runs"):
            return {"run": {"runId": "run-queued", "status": "queued"}}
        if path.endswith("/folder-recommendation"):
            return {"folders": [], "totalDocumentCount": 0}
        if path.endswith("/auto-repair-preview"):
            return {"previewId": "preview-1", "items": []}
        if path.endswith("/auto-repair-queue"):
            return {"jobId": "job-1", "status": "queued", "queuedCount": 1}
        if "/documents/" in path and method == "DELETE":
            return {"deleted": True, "documentId": "document-1"}
        project = self._project()
        if path.endswith("/lifecycle"):
            project["lifecycleState"] = str(payload["targetState"])
            project["version"] = 2
        return {"project": project}

    def project_knowledge_context(self, project_id: str) -> dict[str, Any]:
        self.calls.append(("GET", f"local-knowledge:{project_id}"))
        return {"project": {"projectId": project_id}, "state": {"overall": "ready"}}


class _FakeCompatibility:
    def __init__(self) -> None:
        self.runtime = _FakeRuntime()

    @staticmethod
    def _not_connected(path: str) -> None:
        raise LocalRuntimeError(501, "capability_not_connected", path)


@pytest.mark.parametrize(
    ("method", "path", "body"),
    [
        ("GET", "clients", {}),
        ("POST", "clients", {"name": "项目一"}),
        ("PUT", "clients/project-1", {"name": "项目一"}),
        ("DELETE", "clients/project-1", {}),
        ("GET", "clients/project-1/delete-preview", {}),
        ("POST", "clients/project-1/freeze", {}),
        ("POST", "clients/project-1/unfreeze", {}),
        ("GET", "clients/project-1/knowledge-context", {}),
        ("GET", "clients/project-1/knowledge-status", {}),
        ("GET", "clients/project-1/fact-bundle", {}),
        ("GET", "clients/project-1/duplicate-documents", {}),
        ("GET", "clients/project-1/entities", {}),
        ("GET", "clients/project-1/entity-merge-candidates", {}),
        ("GET", "clients/project-1/glossary", {}),
        ("GET", "clients/project-1/glossary-attributes", {}),
        ("GET", "clients/project-1/glossary-drift-alerts", {}),
        ("GET", "clients/project-1/contradictions", {}),
        ("POST", "clients/project-1/folders/recommend", {}),
        ("POST", "clients/project-1/documents/auto-repair/preview", {}),
        ("POST", "clients/project-1/documents/auto-repair/apply", {}),
        (
            "POST",
            "clients/project-1/glossary",
            {"term": "项目词条"},
        ),
        ("PATCH", "glossary/glossary-1", {"definition": "定义"}),
        ("DELETE", "glossary/glossary-1", {}),
        ("POST", "entities/entity-1/verify", {"status": "canonical"}),
        (
            "POST",
            "entities/entity-2/merge",
            {"survivingEntityId": "entity-1"},
        ),
        (
            "POST",
            "clients/project-1/glossary-attributes/attr-1/verify",
            {},
        ),
        (
            "POST",
            "clients/project-1/glossary-attributes/attr-1/reject",
            {},
        ),
        (
            "POST",
            "clients/project-1/glossary-drift-alerts/alert-1/resolve",
            {"action": "dismiss"},
        ),
        (
            "POST",
            "contradictions/contradiction-1/review",
            {"reviewStatus": "resolved"},
        ),
        (
            "GET",
            "clients/project-1/documents/document-1/reading-preview",
            {},
        ),
        ("DELETE", "clients/project-1/documents/document-1", {}),
        ("GET", "documents/document-1/text", {}),
    ],
)
def test_each_connected_ui_handler_dispatches_to_strict_runtime(
    method: str,
    path: str,
    body: dict[str, Any],
) -> None:
    compatibility = _FakeCompatibility()
    result = router.dispatch(
        compatibility,
        UiRequest(
            method=method,
            path=path,
            query={},
            body=body,
            idempotency_key=f"test-{method}-{path}",
        ),
    )
    assert result is not NOT_HANDLED
    assert compatibility.runtime.calls


def test_link_import_fetches_body_locally_and_registers_only_metadata(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    compatibility = _FakeCompatibility()

    class ImmediateThread:
        def __init__(self, *, target: Any, kwargs: dict[str, Any], **_: Any) -> None:
            self.target = target
            self.kwargs = kwargs

        def start(self) -> None:
            self.target(**self.kwargs)

    class FakeStore:
        def __init__(self) -> None:
            self.runs: dict[str, dict[str, Any]] = {}
            self.bound = False

        def save_link_import_run(
            self,
            _project_id: str,
            run: Mapping[str, Any],
        ) -> dict[str, Any]:
            value = dict(run)
            self.runs[str(value["runId"])] = value
            return value

        def link_import_runs(
            self,
            _project_id: str,
            *,
            limit: int = 20,
            run_id: str | None = None,
        ) -> list[dict[str, Any]]:
            values = list(self.runs.values())
            if run_id:
                values = [
                    item for item in values if item["runId"] == run_id
                ]
            return values[:limit]

        def cancel_link_import_run(
            self,
            _project_id: str,
            run_id: str,
        ) -> dict[str, Any]:
            return self.runs[run_id]

        def import_text(self, **_: Any) -> dict[str, Any]:
            return {
                "localSourceId": "local-link-source",
                "localSummaryId": "local-link-summary",
                "fileName": "链接资料.md",
                "title": "链接资料",
                "contentHash": "link-content-hash",
                "byteSize": 128,
                "mediaType": "text/markdown",
                "managedPath": "/strict/local/link-material.md",
                "updatedAt": "2026-07-30T00:00:00Z",
            }

        def bind_cloud_documents(self, **_: Any) -> None:
            self.bound = True

        def bind_pending_materials(self, **_: Any) -> None:
            return None

        def update_ai_summary(self, *_args: Any, **_kwargs: Any) -> dict[str, Any]:
            return {"summaryKind": "ai_summary"}

    store = FakeStore()
    compatibility.runtime.private_ai_completion = lambda **_kwargs: {
        "content": "网页资料的组织共享摘要",
        "modelName": "test-model",
    }
    monkeypatch.setattr(project_materials_ui, "_local_store", lambda _: store)
    monkeypatch.setattr(project_materials_ui, "_LinkImportThread", ImmediateThread)
    monkeypatch.setattr(
        project_materials_ui,
        "fetch_link_material",
        lambda _url, **_kwargs: {
            "platform": "bilibili",
            "sourceUrl": "https://www.bilibili.com/video/example",
            "title": "链接资料",
            "text": "LINK_BODY_MUST_STAY_LOCAL",
            "metadata": {"accessMode": "anonymous"},
        },
    )
    request = UiRequest(
        method="POST",
        path="clients/project-1/link-materials/import/start",
        query={},
        body={"url": "https://www.bilibili.com/video/example"},
        idempotency_key="link-import-local-1",
    )
    created = router.dispatch(compatibility, request)
    replay = router.dispatch(compatibility, request)
    assert created["status"] == "queued"
    assert replay["status"] == "completed"
    assert replay["documentId"] == "document-text-1"
    assert store.bound is True
    serialized = json.dumps(
        compatibility.runtime.command_payloads[-1],
        ensure_ascii=False,
    )
    assert "LINK_BODY_MUST_STAY_LOCAL" not in serialized
    assert "link-content-hash" in serialized

    cloud_calls_before_local_run_read = list(compatibility.runtime.calls)
    listed = router.dispatch(
        compatibility,
        UiRequest(
            method="GET",
            path="clients/project-1/link-materials/import-runs",
            query={},
            body={},
            idempotency_key="list-link-runs",
        ),
    )
    fetched = router.dispatch(
        compatibility,
        UiRequest(
            method="GET",
            path=f"clients/project-1/link-materials/import-runs/{replay['runId']}",
            query={},
            body={},
            idempotency_key="get-link-run",
        ),
    )
    cancelled = router.dispatch(
        compatibility,
        UiRequest(
            method="POST",
            path=(
                "clients/project-1/link-materials/import-runs/"
                f"{replay['runId']}/cancel"
            ),
            query={},
            body={},
            idempotency_key="cancel-link-run",
        ),
    )
    assert listed[0]["runId"] == replay["runId"]
    assert fetched["runId"] == replay["runId"]
    assert cancelled["status"] == "completed"
    assert compatibility.runtime.calls == cloud_calls_before_local_run_read

    def blocked_fetch(*_args: Any, **_kwargs: Any) -> dict[str, Any]:
        raise link_material_fetcher.LinkMaterialFetchError(
            409,
            "link_import_authentication_required",
            "该内容需要平台登录",
            state="blocked",
            retryable=False,
        )

    monkeypatch.setattr(
        project_materials_ui,
        "fetch_link_material",
        blocked_fetch,
    )
    blocked_started = router.dispatch(
        compatibility,
        UiRequest(
            method="POST",
            path="clients/project-1/link-materials/import/start",
            query={},
            body={"url": "https://www.bilibili.com/video/private"},
            idempotency_key="link-import-local-blocked",
        ),
    )
    blocked_run = store.runs[str(blocked_started["runId"])]
    assert blocked_run["status"] == "failed"
    assert blocked_run["state"] == "blocked"
    assert blocked_run["stage"] == "链接访问受阻"
    assert blocked_run["retryable"] is False
    assert blocked_run["mediaCacheStatus"] == "failed"


def test_link_fetcher_rejects_private_resolution_and_extracts_readable_html(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setattr(
        link_material_fetcher.socket,
        "getaddrinfo",
        lambda *_args, **_kwargs: [
            (2, 1, 6, "", ("127.0.0.1", 443)),
        ],
    )
    with pytest.raises(LocalRuntimeError) as blocked:
        link_material_fetcher._validated_url(  # noqa: SLF001
            "https://www.bilibili.com/video/example"
        )
    assert blocked.value.code == "link_import_private_address_forbidden"

    parser = link_material_fetcher._ReadableHtml()  # noqa: SLF001
    parser.feed(
        "<html><head><title>项目访谈</title>"
        "<meta name='description' content='访谈摘要'></head>"
        "<body><script>SECRET_SCRIPT</script><p>正文事实</p>"
        "<p>正文事实</p><p>下一步行动</p></body></html>"
    )
    assert parser.title == "项目访谈"
    assert parser.description == "访谈摘要"
    assert parser.readable_text() == "项目访谈\n正文事实\n下一步行动"


def test_from_text_handler_keeps_body_local_and_registers_only_metadata(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    compatibility = _FakeCompatibility()
    local_store = SimpleNamespace(
        import_text=lambda **_: {
            "localSourceId": "local-text-1",
            "fileName": "任务交接背景.md",
            "title": "任务交接背景",
            "contentHash": "local-text-hash",
            "byteSize": 36,
            "mediaType": "text/markdown",
            "managedPath": "/local/managed/任务交接背景.md",
        }
    )
    monkeypatch.setattr(
        project_materials_ui,
        "_local_store",
        lambda _: local_store,
    )
    result = router.dispatch(
        compatibility,
        UiRequest(
            method="POST",
            path="clients/project-1/documents/from-text",
            query={},
            body={
                "title": "任务交接背景",
                "content": "LOCAL_SOURCE_BODY_MUST_STAY_ON_DEVICE",
            },
            idempotency_key="from-text-handler-1",
        ),
    )
    assert result["documentId"] == "document-text-1"
    assert result["sourceScope"] == "local_private"
    assert compatibility.runtime.calls[-1] == (
        "POST",
        "/api/v2/domain/project-materials/projects/project-1"
        "/materials/register-metadata",
    )
    serialized_payload = json.dumps(
        compatibility.runtime.command_payloads[-1],
        ensure_ascii=False,
    )
    assert "LOCAL_SOURCE_BODY_MUST_STAY_ON_DEVICE" not in serialized_payload
    assert "local-text-hash" in serialized_payload


def test_document_text_prefers_local_body_without_cloud_probe(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    compatibility = _FakeCompatibility()
    compatibility.runtime.database_path = Path("/tmp/strict-local.db")
    compatibility.runtime.cloud_query = lambda *_args, **_kwargs: (_ for _ in ()).throw(
        AssertionError("local document text must not probe organization cloud")
    )
    monkeypatch.setattr(
        project_materials_ui,
        "_local_store",
        lambda _: SimpleNamespace(
            document_text=lambda _document_id: {
                "projectId": "project-1",
                "title": "本机资料",
                "content": "LOCAL_SENTINEL_BODY",
                "sourceScope": "local_private",
            }
        ),
    )
    result = router.dispatch(
        compatibility,
        UiRequest(
            method="GET",
            path="documents/document-local/text",
            query={},
            body={},
            idempotency_key="local-text-first",
        ),
    )
    assert result["content"] == "LOCAL_SENTINEL_BODY"
    assert result["sourceScope"] == "local_private"


def test_feishu_document_import_preflights_authorities_without_fake_import(
) -> None:
    member_linked = False
    expected_code = "feishu_member_authorization_required"
    class Runtime:
        def __init__(self) -> None:
            self.calls: list[tuple[str, dict[str, str] | None]] = []

        def cloud_query(
            self,
            path: str,
            *,
            query: dict[str, str] | None = None,
        ) -> dict[str, Any]:
            self.calls.append((path, query))
            if path.endswith("/projects/project-1"):
                return {"project": {"projectId": "project-1"}}
            if path != "/api/v2/platform-integrations/query":
                raise AssertionError(path)
            if query == {
                "resourcePath": "org-integrations/feishu",
                "authorizationScope": "organization",
            }:
                return {"resource": {"state": "ready"}}
            if query == {
                "resourcePath": "me/feishu-authorization",
                "authorizationScope": "personal",
            }:
                return {
                    "resource": {
                        "linked": member_linked,
                        "blockedReason": (
                            None
                            if member_linked
                            else "oauth_grant_authority_not_connected"
                        ),
                    }
                }
            raise AssertionError(query)

    compatibility = SimpleNamespace(runtime=Runtime())
    with pytest.raises(LocalRuntimeError) as blocked:
        router.dispatch(
            compatibility,
            UiRequest(
                method="POST",
                path="clients/project-1/feishu-doc-import/import",
                query={},
                body={
                    "items": [
                        {
                            "token": "doc-token-1",
                            "type": "docx",
                            "title": "成员飞书原文",
                            "url": "https://example.feishu.cn/docx/doc-token-1",
                            "content": "RAW_FEISHU_BODY_MUST_NOT_REACH_CLOUD",
                        }
                    ]
                },
                idempotency_key="feishu-document-import-1",
            ),
        )
    assert blocked.value.code == expected_code
    assert all(
        call[0].endswith("/projects/project-1")
        or call[0] == "/api/v2/platform-integrations/query"
        for call in compatibility.runtime.calls
    )
    assert "RAW_FEISHU_BODY_MUST_NOT_REACH_CLOUD" not in json.dumps(
        compatibility.runtime.calls,
        ensure_ascii=False,
    )


def test_feishu_document_import_keeps_body_local_and_publishes_only_summary(
    tmp_path: Path,
) -> None:
    database = tmp_path / "local" / "strict-local.db"
    runtime = WorkspaceRuntime(database, MemorySecretStore())
    sandbox_id = runtime.current()["sandbox"]["sandboxId"]
    runtime._current_context = lambda require_ready=True: SimpleNamespace(  # type: ignore[method-assign]
        sandbox_id=sandbox_id,
        cloud_instance_id="cloud-feishu-import",
        organization_id="organization-feishu-import",
        cloud_api_url="https://feishu-import.invalid",
        principal_id="principal-feishu-import",
        membership_id="membership-feishu-import",
    )
    raw_body = "FEISHU_RAW_BODY_LOCAL_STORAGE_ONLY_9247"
    cloud_payloads: list[tuple[str, dict[str, Any]]] = []

    def cloud_query(
        path: str,
        *,
        query: dict[str, str] | None = None,
    ) -> dict[str, Any]:
        if path.endswith("/projects/project-feishu-import"):
            return {"project": {"projectId": "project-feishu-import"}}
        if path == "/api/v2/platform-integrations/query":
            if query == {
                "resourcePath": "org-integrations/feishu",
                "authorizationScope": "organization",
            }:
                return {"resource": {"state": "ready"}}
            if query == {
                "resourcePath": "me/feishu-authorization",
                "authorizationScope": "personal",
            }:
                return {"resource": {"linked": True, "state": "ready"}}
        raise AssertionError((path, query))

    def cloud_command(
        method: str,
        path: str,
        *,
        payload: dict[str, Any],
        idempotency_key: str,
        refresh_business: bool = True,
    ) -> dict[str, Any]:
        del method, idempotency_key, refresh_business
        cloud_payloads.append((path, payload))
        if path == "/api/v2/platform-integrations/command":
            resource_path = payload["resourcePath"]
            if resource_path == "feishu-doc-import/fetch":
                assert raw_body not in json.dumps(payload, ensure_ascii=False)
                return {
                    "result": {
                        "state": "ready",
                        "items": [
                            {
                                "token": "docx-token-local",
                                "type": "docx",
                                "title": "飞书项目背景",
                                "url": (
                                    "https://example.feishu.cn/docx/"
                                    "docx-token-local"
                                ),
                                "content": raw_body,
                            }
                        ],
                        "failedItems": [],
                    }
                }
            if resource_path == "feishu-doc-import/register-mapping":
                return {
                    "result": {
                        "state": "succeeded",
                        "status": "synced",
                    }
                }
            raise AssertionError(resource_path)
        if path.endswith("/materials/register-metadata"):
            material = payload["materials"][0]
            assert raw_body not in json.dumps(payload, ensure_ascii=False)
            return {
                "documents": [
                    {
                        "documentId": "cloud-document-feishu",
                        "localSourceId": material["localSourceId"],
                        "version": 1,
                    }
                ],
                "materialBoundary": {
                    "sourceFileContentUploaded": False,
                    "sourceFilePathUploaded": False,
                    "localSummaryUploaded": False,
                },
            }
        if path.endswith("/publish-local-summary"):
            assert raw_body not in json.dumps(payload, ensure_ascii=False)
            assert payload["summary"] == "组织共享的飞书项目背景摘要"
            return {
                "documentId": "cloud-document-feishu",
                "version": 2,
            }
        raise AssertionError(path)

    runtime.cloud_query = cloud_query  # type: ignore[method-assign]
    runtime.cloud_command = cloud_command  # type: ignore[method-assign]
    runtime.private_ai_completion = lambda **_: {  # type: ignore[method-assign]
        "content": "组织共享的飞书项目背景摘要",
        "modelName": "strict-summary-test",
    }
    result = router.dispatch(
        SimpleNamespace(runtime=runtime),
        UiRequest(
            method="POST",
            path=(
                "clients/project-feishu-import/"
                "feishu-doc-import/import"
            ),
            query={},
            body={
                "items": [
                    {
                        "token": "docx-token-local",
                        "type": "docx",
                        "title": "飞书项目背景",
                        "url": (
                            "https://example.feishu.cn/docx/"
                            "docx-token-local"
                        ),
                        "content": "UNTRUSTED_RENDERER_BODY_MUST_BE_IGNORED",
                    }
                ]
            },
            idempotency_key="feishu-local-import-1",
        ),
    )
    assert result["importedCount"] == 1
    assert result["failedCount"] == 0
    assert result["items"][0]["documentId"] == "cloud-document-feishu"
    assert result["items"][0]["sharedKnowledgeState"] == "ready"
    assert "UNTRUSTED_RENDERER_BODY_MUST_BE_IGNORED" not in json.dumps(
        cloud_payloads,
        ensure_ascii=False,
    )
    store = LocalProjectMaterialsRepository(runtime)
    local_document = store.document_text("cloud-document-feishu")
    assert local_document["content"] == raw_body
    with runtime_connection(database, "local", read_only=True) as connection:
        source_objects = connection.execute(
            """
            SELECT COUNT(*)
            FROM storage_objects
            WHERE sandbox_id = ? AND media_type = 'text/markdown'
            """,
            (sandbox_id,),
        ).fetchone()[0]
    assert source_objects == 1
    restarted = WorkspaceRuntime(database, MemorySecretStore())
    restarted._current_context = lambda require_ready=True: SimpleNamespace(  # type: ignore[method-assign]
        sandbox_id=sandbox_id,
        membership_id="membership-feishu-import",
    )
    assert (
        LocalProjectMaterialsRepository(restarted)
        .document_text("cloud-document-feishu")["content"]
        == raw_body
    )


def test_duplicate_resolution_preflights_local_then_blocks_before_any_delete(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    compatibility = _FakeCompatibility()
    events: list[str] = []

    class LocalStore:
        @staticmethod
        def preflight_duplicate_resolution(
            project_id: str,
            payload: dict[str, Any],
        ) -> dict[str, Any]:
            events.append("local_preflight")
            assert project_id == "project-1"
            return {
                "groupKey": str(payload["groupKey"]),
                "action": "delete_others",
                "keepV2DocumentIds": ["document-keep"],
                "deleteV2DocumentIds": ["document-delete-1", "document-delete-2"],
                "migrateReferences": True,
                "note": "",
            }

    monkeypatch.setattr(
        project_materials_ui,
        "_local_store",
        lambda _: LocalStore(),
    )
    with pytest.raises(LocalRuntimeError) as blocked:
        router.dispatch(
            compatibility,
            UiRequest(
                method="POST",
                path="clients/project-1/duplicate-documents/resolve",
                query={},
                body={
                    "groupKey": "hash:abc",
                    "action": "delete_others",
                    "keepV2DocumentIds": ["document-keep"],
                    "deleteV2DocumentIds": [
                        "document-delete-1",
                        "document-delete-2",
                    ],
                    "migrateReferences": True,
                },
                idempotency_key="duplicates-handler-1",
            ),
        )
    assert blocked.value.status_code == 409
    assert blocked.value.code == "duplicate_resolution_cloud_metadata_not_connected"
    assert events == ["local_preflight"]
    assert compatibility.runtime.calls == []


def test_inventory_denominator_is_fully_connected() -> None:
    assert len(router.routes) == 69
    assert _UNSUPPORTED_ROUTE_SPECS == ()


def test_gc10_source_groups_keep_local_and_cloud_sources_distinct() -> None:
    groups = project_materials_ui._source_groups(
        hits=[
            {"sourceType": "local_document"},
            {"sourceType": "organization_knowledge"},
            {"sourceType": "official_website"},
            {"sourceType": "explicit_memory"},
            {"sourceType": "favorite"},
        ],
        organization_state="ready",
    )
    assert [item["key"] for item in groups] == [
        "local_original",
        "organization_knowledge",
        "official_website",
        "explicit_memory",
        "favorite",
        "system_inference",
    ]
    assert [item["count"] for item in groups] == [1, 1, 1, 1, 1, 0]
    assert groups[-1]["state"] == "not_connected"


def test_gc10_cloud_manifest_summary_exposes_only_declared_safe_text() -> None:
    assert CloudRepository._knowledge_manifest_summary(
        json.dumps(
            {
                "summary": "组织已发布摘要",
                "localPath": "/Users/member/private.docx",
            }
        )
    ) == "组织已发布摘要"
    assert CloudRepository._knowledge_manifest_summary(
        json.dumps({"localPath": "/Users/member/private.docx"})
    ) == ""


def test_local_wiki_chunker_is_bounded_and_deterministic() -> None:
    content = "# 日慈项目背景\n\n" + ("张真是日慈基金会秘书长。" * 2_000)
    first = LocalProjectMaterialsRepository._chunk_wiki_text(content)
    second = LocalProjectMaterialsRepository._chunk_wiki_text(content)
    assert first == second
    assert len(first) > 1
    assert max(len(str(item["content"])) for item in first) <= 12_000
    assert first[0]["start"] == 0
    assert first[-1]["end"] == len(content)
    assert all(
        int(item["start"]) < int(item["end"])
        for item in first
    )


def test_local_wiki_retrieval_tokens_and_fact_locators_are_rebuildable() -> None:
    content = "张真是日慈基金会秘书长。心灵魔法学院关注杰克逊指标。"
    candidates = LocalProjectMaterialsRepository._fact_candidates(
        content,
        document_offset=120,
    )
    assert [item["content"] for item in candidates] == [
        "张真是日慈基金会秘书长。",
        "心灵魔法学院关注杰克逊指标。",
    ]
    assert candidates[0]["documentStart"] == 120
    assert candidates[0]["documentEnd"] == 120 + len(candidates[0]["content"])

    query_tokens = LocalProjectMaterialsRepository._retrieval_tokens(
        "请问日慈研究团队为什么忙"
    )
    matching_tokens = LocalProjectMaterialsRepository._retrieval_tokens(
        "日慈研究团队正在处理大量研究工作"
    )
    unrelated_tokens = LocalProjectMaterialsRepository._retrieval_tokens(
        "苹果香蕉天气"
    )
    query_vector = LocalProjectMaterialsRepository._sparse_vector(query_tokens)
    assert "为什么" not in query_tokens
    assert LocalProjectMaterialsRepository._sparse_cosine(
        query_vector,
        LocalProjectMaterialsRepository._sparse_vector(matching_tokens),
    ) > LocalProjectMaterialsRepository._sparse_cosine(
        query_vector,
        LocalProjectMaterialsRepository._sparse_vector(unrelated_tokens),
    )


def test_local_pdf_text_and_scanned_pdf_state_are_explicit(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from pypdf import PdfWriter
    from pypdf.generic import (
        DecodedStreamObject,
        DictionaryObject,
        NameObject,
    )

    runtime = WorkspaceRuntime(
        tmp_path / "local-pdf" / "strict-local.db",
        MemorySecretStore(),
    )
    sandbox_id = runtime.current()["sandbox"]["sandboxId"]
    runtime._current_context = lambda require_ready=True: SimpleNamespace(  # type: ignore[method-assign]
        sandbox_id=sandbox_id,
        cloud_instance_id="cloud-pdf",
        organization_id="organization-pdf",
        cloud_api_url="https://pdf.invalid",
        principal_id="principal-pdf",
        membership_id="membership-pdf",
    )
    store = LocalProjectMaterialsRepository(runtime)

    text_pdf = tmp_path / "native-text.pdf"
    writer = PdfWriter()
    page = writer.add_blank_page(width=612, height=792)
    font = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
        }
    )
    stream = DecodedStreamObject()
    stream.set_data(
        b"BT /F1 12 Tf 72 720 Td (PDF_TEXT_SENTINEL) Tj ET"
    )
    page[NameObject("/Resources")] = DictionaryObject(
        {
            NameObject("/Font"): DictionaryObject(
                {NameObject("/F1"): writer._add_object(font)}
            )
        }
    )
    page[NameObject("/Contents")] = writer._add_object(stream)
    with text_pdf.open("wb") as handle:
        writer.write(handle)

    blank_pdf = tmp_path / "scanned-image-only.pdf"
    blank_writer = PdfWriter()
    blank_writer.add_blank_page(width=612, height=792)
    with blank_pdf.open("wb") as handle:
        blank_writer.write(handle)

    imported = store.import_paths(
        project_id="project-pdf",
        mode="file",
        paths=[str(text_pdf), str(blank_pdf)],
    )
    store.bind_cloud_documents(
        project_id="project-pdf",
        local_materials=imported["materials"],
        cloud_documents=[
            {
                "localSourceId": imported["materials"][0]["localSourceId"],
                "documentId": "pdf-text",
            },
            {
                "localSourceId": imported["materials"][1]["localSourceId"],
                "documentId": "pdf-scanned",
            },
        ],
    )
    assert "PDF_TEXT_SENTINEL" in store.document_text("pdf-text")["content"]
    monkeypatch.setattr(
        "backend.app.project_materials_local.extract_ocr_text",
        lambda _path, **_kwargs: "OCR_TEXT_SENTINEL",
    )
    scanned = store.document_text("pdf-scanned")
    assert scanned["kind"] == "pdf_ocr"
    assert scanned["content"] == "OCR_TEXT_SENTINEL"
def test_project_list_keeps_cloud_projects_visible_when_local_state_is_corrupt(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    class CorruptLocalStore:
        def ensure_project_projection(self, _project: Mapping[str, Any]) -> None:
            raise LocalRuntimeError(
                409,
                "local_project_state_corrupt",
                "本机项目资料状态校验失败",
            )

    class Runtime:
        database_path = Path("strict-local.db")

        def cloud_query(self, path: str, query: Any = None) -> dict[str, Any]:
            del query
            assert path.endswith("/projects")
            return {
                "projects": [
                    {
                        "projectId": "project-cloud-authority",
                        "name": "组织云项目",
                        "version": 1,
                        "lifecycleState": "active",
                    }
                ]
            }

        def reconcile_project_projections(self, project_ids: list[str]) -> None:
            assert project_ids == ["project-cloud-authority"]

    monkeypatch.setattr(
        project_materials_ui,
        "_local_store",
        lambda _compatibility: CorruptLocalStore(),
    )
    listed = router.dispatch(
        SimpleNamespace(runtime=Runtime()),
        UiRequest(
            method="GET",
            path="clients",
            query={},
            body={},
            idempotency_key="project-list-local-state-corrupt",
        ),
    )

    assert len(listed) == 1
    assert listed[0]["id"] == "project-cloud-authority"
    assert listed[0]["name"] == "组织云项目"
    assert listed[0]["folderCount"] == 0
    assert listed[0]["folderCapabilityState"] == "local_recovery_required"
