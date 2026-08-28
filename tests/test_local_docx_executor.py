from __future__ import annotations

import hashlib
from pathlib import Path
from types import SimpleNamespace

import pytest
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Inches

from backend.app.project_materials_local import (
    LocalProjectMaterialsRepository,
)
from backend.app.runtime import LocalRuntimeError, WorkspaceRuntime
from backend.app.secret_store import MemorySecretStore
from strict_common.ids import utc_now


def _seed_runtime_scope(runtime: WorkspaceRuntime) -> str:
    now = utc_now()
    sandbox_id = "sandbox-docx-test"
    with runtime._connection() as connection:
        connection.execute(
            "INSERT INTO organizations (id,lifecycle_state,version,updated_at,"
            "record_kind,name,created_at,deleted_at,projection_state,projected_at) "
            "VALUES ('organization-docx','active',1,?,'organization',"
            "'文档测试组织',?,NULL,'current',?)",
            (now, now, now),
        )
        connection.execute(
            "INSERT INTO principals (id,status,identity_version,updated_at,"
            "principal_kind,display_name,version,lifecycle_state,created_at,"
            "deleted_at,projection_state,projected_at) VALUES ("
            "'principal-docx','active',1,?,'person','文档测试成员',1,"
            "'active',?,NULL,'current',?)",
            (now, now, now),
        )
        connection.execute(
            "INSERT INTO authorization_scopes (id,scope_kind,organization_id,"
            "policy_version,created_at,updated_at,status,version,lifecycle_state,"
            "deleted_at,projection_state,projected_at) VALUES ('scope-docx',"
            "'organization','organization-docx',1,?,?,'active',1,'active',NULL,"
            "'current',?)",
            (now, now, now),
        )
        connection.execute(
            "INSERT INTO organization_memberships (id,scope_id,principal_id,"
            "role_key,status,version,record_kind,visibility_scope,lifecycle_state,"
            "created_at,updated_at,deleted_at,projection_state,projected_at) "
            "VALUES ('membership-docx','scope-docx','principal-docx','admin',"
            "'active',1,'membership','organization','active',?,?,NULL,'current',?)",
            (now, now, now),
        )
        connection.execute(
            "INSERT INTO sandboxes (id,scope_id,principal_id,membership_id,"
            "record_kind,cloud_instance_id,database_generation_id,sandbox_kind,"
            "display_name,runtime_status,manifest_hash,version,lifecycle_state,"
            "created_at,updated_at,deleted_at,authority_role,origin_instance_id) "
            "VALUES (?,'scope-docx','principal-docx','membership-docx','sandbox',"
            "'cloud-docx',?,'organization','文档测试工作空间','ready',?,1,"
            "'active',?,?,NULL,'local',?)",
            (
                sandbox_id,
                runtime.identity.database_generation_id,
                runtime.identity.manifest_hash,
                now,
                now,
                runtime.identity.database_generation_id,
            ),
        )
        connection.commit()
    return sandbox_id


def _store(
    tmp_path: Path,
) -> tuple[WorkspaceRuntime, LocalProjectMaterialsRepository, str]:
    runtime = WorkspaceRuntime(
        tmp_path / "local" / "strict-local.db",
        MemorySecretStore(),
    )
    sandbox_id = _seed_runtime_scope(runtime)
    runtime._current_context = lambda require_ready=True: SimpleNamespace(  # type: ignore[method-assign]
        sandbox_id=sandbox_id,
        membership_id="member-docx-test",
    )
    return runtime, LocalProjectMaterialsRepository(runtime), sandbox_id


def _restart(
    database_path: Path,
    sandbox_id: str,
) -> LocalProjectMaterialsRepository:
    runtime = WorkspaceRuntime(database_path, MemorySecretStore())
    runtime._current_context = lambda require_ready=True: SimpleNamespace(  # type: ignore[method-assign]
        sandbox_id=sandbox_id,
        membership_id="member-docx-test",
    )
    return LocalProjectMaterialsRepository(runtime)


def _all_text(document: Document) -> str:
    values = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            values.extend(cell.text for cell in row.cells)
    for section in document.sections:
        values.extend(paragraph.text for paragraph in section.header.paragraphs)
        values.extend(paragraph.text for paragraph in section.footer.paragraphs)
    return "\n".join(values)


def _template(path: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.83)
    section.header.paragraphs[0].add_run("别名：{{项目别名}}")

    paragraph = document.add_paragraph()
    paragraph.add_run("项目：")
    first = paragraph.add_run("{{项目")
    first.bold = True
    second = paragraph.add_run("名称}}")
    second.bold = True
    document.add_paragraph("待确认：{{待补充字段}}")

    fields = document.add_table(rows=1, cols=2)
    fields.style = "Table Grid"
    fields.rows[0].cells[0].text = "项目简介"
    fields.rows[0].cells[1].text = "待填写"

    attachments = document.add_table(rows=2, cols=2)
    attachments.rows[0].cells[0].text = "序号"
    attachments.rows[0].cells[1].text = "附件名称"
    attachments.rows[1].cells[0].text = "1"
    attachments.rows[1].cells[1].text = "项目预算"
    document.save(path)


def _editable_docx(path: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = (
        section.page_height,
        section.page_width,
    )
    section.top_margin = Inches(0.61)
    section.header.paragraphs[0].text = "严格项目页眉"
    document.add_heading("原标题", level=0)
    paragraph = document.add_paragraph("原正文")
    paragraph.style = document.styles["Normal"]
    document.save(path)


def test_docx_editor_preserves_headings_lists_inline_styles_and_tables(
    tmp_path: Path,
) -> None:
    source = tmp_path / "rich-feishu-export-飞书.docx"
    document = Document()
    title = document.add_paragraph()
    title_run = title.add_run("rich-feishu-export")
    title_run.bold = True
    document.add_heading("一级标题", level=1)
    document.add_heading("二级标题", level=2)
    paragraph = document.add_paragraph()
    bold = paragraph.add_run("粗体 ")
    bold.bold = True
    paragraph.add_run("与")
    italic = paragraph.add_run("斜体")
    italic.italic = True
    adjacent = document.add_paragraph()
    adjacent.add_run("选择")
    adjacent_bold = adjacent.add_run("“资助”")
    adjacent_bold.bold = True
    adjacent.add_run("形式继续推进")
    document.add_paragraph("第一项", style="List Number")
    document.add_paragraph("第一项说明")
    document.add_paragraph("第二项", style="List Number")
    document.add_paragraph("第二项说明")
    document.add_paragraph("第三项", style="List Number")
    for value in ("项目符号一", "项目符号二"):
        document.add_paragraph(value, style="List Bullet")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "字段"
    table.cell(0, 1).text = "内容"
    table.cell(1, 0).text = "项目"
    table.cell(1, 1).text = "益语智库"
    document.save(source)

    markdown = LocalProjectMaterialsRepository._docx_editor_markdown(source)

    assert markdown.splitlines()[0] == "# 一级标题"
    assert "rich-feishu-export" not in markdown
    assert "## 二级标题" in markdown
    assert "**粗体**" in markdown
    assert "与*斜体*" in markdown
    assert "选择&#8203;**“资助”**&#8203;形式继续推进" in markdown
    assert "1\\. 第一项" in markdown
    assert "2\\. 第二项" in markdown
    assert "3\\. 第三项" in markdown
    assert "1\\. 第一项\n\n第一项说明\n\n2\\. 第二项" in markdown
    assert "2\\. 第二项\n\n第二项说明\n\n3\\. 第三项" in markdown
    assert "- 项目符号一" in markdown
    assert "| 字段 | 内容 |" in markdown
    assert "| 项目 | 益语智库 |" in markdown


def test_docx_template_fill_is_local_managed_idempotent_and_restartable(
    tmp_path: Path,
) -> None:
    runtime, store, sandbox_id = _store(tmp_path)
    source = tmp_path / "outside-local-db" / "项目模板.docx"
    source.parent.mkdir()
    _template(source)
    original_hash = hashlib.sha256(source.read_bytes()).hexdigest()

    run = store.start_template_fill(
        "project-docx",
        template_path=str(source),
        values={
            "项目名称": "日慈基金会",
            "项目别名": "日慈",
            "项目简介": "面向乡村儿童的社会情感支持项目",
        },
        idempotency_key="fill-project-template",
    )

    assert run["status"] == "completed"
    assert run["fieldCount"] == 4
    assert run["filledCount"] == 3
    assert run["missingCount"] == 1
    assert run["attachmentChecklist"] == ["项目预算"]
    assert run["sourceScope"] == "local_private"
    assert run["persistedToOrganizationCloud"] is False
    assert run["outputPath"].endswith(".docx")
    assert hashlib.sha256(source.read_bytes()).hexdigest() == original_hash

    output_path = Path(run["outputPath"])
    assert runtime.database_path.parent in output_path.parents
    output = Document(output_path)
    output_text = _all_text(output)
    assert "项目：日慈基金会" in output_text
    assert "别名：日慈" in output_text
    assert "面向乡村儿童的社会情感支持项目" in output_text
    assert "{{待补充字段}}" in output_text
    assert abs(output.sections[0].left_margin - Inches(0.72)) < 1000
    assert abs(output.sections[0].right_margin - Inches(0.83)) < 1000
    project_run = next(
        run
        for paragraph in output.paragraphs
        for run in paragraph.runs
        if "日慈基金会" in run.text
    )
    assert project_run.bold is True

    source_row = runtime.local_storage_object_get(
        sandbox_id=sandbox_id,
        object_id=run["templateSourceObjectId"],
    )
    output_row = runtime.local_storage_object_get(
        sandbox_id=sandbox_id,
        object_id=run["outputObjectId"],
    )
    assert source_row is not None
    assert output_row is not None
    assert source_row["media_type"] == store.DOCX_MEDIA_TYPE
    assert output_row["media_type"] == store.DOCX_MEDIA_TYPE
    assert runtime.database_path.parent in store._managed_path(
        source_row["storage_key"]
    ).parents

    replay = store.start_template_fill(
        "project-docx",
        template_path=str(source),
        values={
            "项目名称": "日慈基金会",
            "项目别名": "日慈",
            "项目简介": "面向乡村儿童的社会情感支持项目",
        },
        idempotency_key="fill-project-template",
    )
    assert replay["id"] == run["id"]
    assert replay["outputStorageVersion"] == 1

    with pytest.raises(LocalRuntimeError) as conflict:
        store.start_template_fill(
            "project-docx",
            template_path=str(source),
            values={"项目名称": "不同项目"},
            idempotency_key="fill-project-template",
        )
    assert conflict.value.code == "template_fill_idempotency_conflict"

    state_without_receipt = store._load_project_state("project-docx")
    state_without_receipt["templateFillRuns"].pop(run["id"])
    store._write_project_state("project-docx", state_without_receipt)
    recovered = store.start_template_fill(
        "project-docx",
        template_path=str(source),
        values={
            "项目名称": "日慈基金会",
            "项目别名": "日慈",
            "项目简介": "面向乡村儿童的社会情感支持项目",
        },
        idempotency_key="fill-project-template",
    )
    assert recovered["id"] == run["id"]
    assert recovered["outputStorageVersion"] == 1

    restarted = _restart(runtime.database_path, sandbox_id)
    restored = restarted.template_fill_run("project-docx", run["id"])
    assert restored["outputContentHash"] == run["outputContentHash"]
    assert Path(restored["outputPath"]).is_file()

    state = restarted._load_project_state("project-docx")
    state["templateFillRuns"][run["id"]]["outputStorageKey"] = (
        "../../escaped.docx"
    )
    restarted._write_project_state("project-docx", state)
    with pytest.raises(LocalRuntimeError) as escaped:
        restarted.template_fill_run("project-docx", run["id"])
    assert escaped.value.code == "local_storage_path_invalid"


def test_docx_template_fill_rejects_non_docx_and_broken_packages(
    tmp_path: Path,
) -> None:
    _, store, _ = _store(tmp_path)
    text = tmp_path / "template.txt"
    text.write_text("{{项目名称}}", encoding="utf-8")
    with pytest.raises(LocalRuntimeError) as unsupported:
        store.start_template_fill(
            "project-docx",
            template_path=str(text),
            values={},
        )
    assert unsupported.value.status_code == 415
    assert unsupported.value.code == "template_format_unsupported"

    broken = tmp_path / "broken.docx"
    broken.write_bytes(b"not-a-word-package")
    with pytest.raises(LocalRuntimeError) as invalid:
        store.start_template_fill(
            "project-docx",
            template_path=str(broken),
            values={},
        )
    assert invalid.value.status_code == 422
    assert invalid.value.code == "template_docx_invalid"


def test_docx_editor_roundtrip_preserves_package_cas_and_idempotency(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    runtime, store, sandbox_id = _store(tmp_path)
    source = tmp_path / "editable-source.docx"
    _editable_docx(source)
    original_source_hash = hashlib.sha256(source.read_bytes()).hexdigest()
    material = store.import_paths(
        project_id="project-docx",
        mode="file",
        paths=[str(source)],
        idempotency_key="import-editable-docx",
    )["materials"][0]
    store.bind_cloud_documents(
        project_id="project-docx",
        local_materials=[material],
        cloud_documents=[
            {
                "localSourceId": material["localSourceId"],
                "documentId": "document-docx",
            }
        ],
    )
    before = store.document_text("document-docx")
    assert before["editableInPlace"] is True
    managed_path = Path(before["path"])
    managed_before = managed_path.read_bytes()

    content = (
        "# 项目背景\n\n"
        "这是**更新后的正文**。\n\n"
        "- 第一项\n"
        "- 第二项\n\n"
        "| 字段 | 内容 |\n"
        "| --- | --- |\n"
        "| 项目 | 日慈基金会 |"
    )
    saved = store.update_document_text(
        "document-docx",
        title="日慈项目资料",
        content=content,
        expected_version=before["storageVersion"],
        idempotency_key="save-docx-editor",
    )
    assert saved["storageVersion"] == before["storageVersion"] + 1
    assert saved["idempotentReplay"] is False
    assert saved["mediaType"] == store.DOCX_MEDIA_TYPE
    assert saved["fileName"] == "日慈项目资料.docx"
    assert Path(saved["path"]).name.endswith("-日慈项目资料.docx")
    # Editing an imported source creates a managed successor; the user's
    # original file remains untouched at its original path.
    assert managed_path.exists()
    assert hashlib.sha256(source.read_bytes()).hexdigest() == original_source_hash

    managed_path = Path(saved["path"])
    rendered = Document(managed_path)
    rendered_text = _all_text(rendered)
    assert "日慈项目资料" in rendered_text
    assert "项目背景" in rendered_text
    assert "更新后的正文" in rendered_text
    assert "第一项" in rendered_text
    assert "日慈基金会" in rendered_text
    assert rendered.sections[0].orientation == WD_ORIENT.LANDSCAPE
    assert abs(rendered.sections[0].top_margin - Inches(0.61)) < 1000
    assert "严格项目页眉" in rendered_text
    assert any(
        run.bold
        for paragraph in rendered.paragraphs
        for run in paragraph.runs
        if "更新后的正文" in run.text
    )

    replay = store.update_document_text(
        "document-docx",
        title="日慈项目资料",
        content=content,
        expected_version=before["storageVersion"],
        idempotency_key="save-docx-editor",
    )
    assert replay["idempotentReplay"] is True
    assert replay["storageVersion"] == saved["storageVersion"]

    before_stale = managed_path.read_bytes()
    with pytest.raises(LocalRuntimeError) as stale:
        store.update_document_text(
            "document-docx",
            title="不会覆盖",
            content="过期版本内容",
            expected_version=before["storageVersion"],
            idempotency_key="stale-docx-editor",
        )
    assert stale.value.code == "local_storage_version_conflict"
    assert managed_path.read_bytes() == before_stale

    stable_hash = hashlib.sha256(managed_path.read_bytes()).hexdigest()
    stable_version = store.document_text("document-docx")["storageVersion"]

    def fail_render(*args: object, **kwargs: object) -> bytes:
        raise LocalRuntimeError(
            422,
            "docx_roundtrip_failed",
            "模拟生成失败",
        )

    monkeypatch.setattr(store, "_render_docx_roundtrip", fail_render)
    with pytest.raises(LocalRuntimeError) as failed:
        store.update_document_text(
            "document-docx",
            title="失败不覆盖",
            content="这次生成会失败",
            expected_version=stable_version,
            idempotency_key="failed-docx-editor",
        )
    assert failed.value.code == "docx_roundtrip_failed"
    assert hashlib.sha256(managed_path.read_bytes()).hexdigest() == stable_hash
    assert store.document_text("document-docx")["storageVersion"] == (
        stable_version
    )

    restarted = _restart(runtime.database_path, sandbox_id)
    restored = restarted.document_text("document-docx")
    assert restored["editableInPlace"] is True
    assert restored["storageVersion"] == saved["storageVersion"]
    assert "更新后的正文" in restored["content"]
    assert managed_before != managed_path.read_bytes()


def test_docx_editor_rejects_corrupt_managed_source_without_overwrite(
    tmp_path: Path,
) -> None:
    _, store, sandbox_id = _store(tmp_path)
    broken = tmp_path / "broken-source.docx"
    broken.write_bytes(b"not-a-docx")
    material = store.import_paths(
        project_id="project-docx",
        mode="file",
        paths=[str(broken)],
    )["materials"][0]
    store.bind_cloud_documents(
        project_id="project-docx",
        local_materials=[material],
        cloud_documents=[
            {
                "localSourceId": material["localSourceId"],
                "documentId": "broken-document",
            }
        ],
    )
    state = store._load_project_state("project-docx")
    entry = state["documents"]["broken-document"]
    path, row = store._source_path(entry)
    before = path.read_bytes()

    with pytest.raises(LocalRuntimeError) as invalid:
        store.update_document_text(
            "broken-document",
            title="损坏模板",
            content="不会写入",
            expected_version=int(row["version"]),
        )
    assert invalid.value.code == "local_document_format_invalid"
    assert path.read_bytes() == before
    assert store.runtime.local_storage_object_get(
        sandbox_id=sandbox_id,
        object_id=entry["localSourceId"]
    )["version"] == row["version"]
